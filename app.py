# Standard library imports
import io
import os
import re
import json
import time
import logging
import traceback
from datetime import datetime
from enum import Enum, auto
from typing import Dict, List, Any, Tuple, Optional, Pattern, Callable, Set
from dataclasses import dataclass
from functools import wraps
import requests
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from types import MappingProxyType
import threading

# Third-party imports
import gradio as gr
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.exceptions import InvalidFileFormatError
from colorama import init, Fore, Style
# from weasyprint import HTML  # PDF generation related import

# Constants
DEFAULT_PORT = 7860
DEFAULT_HOST = "0.0.0.0"
DEFAULT_LOG_FORMAT = "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
DEFAULT_LOG_LEVEL = "INFO"

# 1. Base Exception Classes
class DocumentCheckError(Exception):
    """Base exception for document checker errors."""
    pass

class ConfigurationError(DocumentCheckError):
    """Raised when configuration is invalid."""
    pass

class DocumentTypeError(DocumentCheckError):
    """Raised when document type is invalid."""
    pass

# 2. Configuration Classes
@dataclass
class PatternConfig:
    """Configuration for pattern matching."""
    pattern: str
    description: str
    is_error: bool
    replacement: Optional[str] = None
    keep_together: bool = False
    format_name: Optional[str] = None
    
    def compile(self) -> Pattern:
        """Compile the pattern."""
        try:
            return re.compile(self.pattern)
        except re.error as e:
            raise ConfigurationError(f"Invalid pattern '{self.pattern}': {e}")

class DocumentType(Enum):
    """Enumeration of supported document types."""
    ADVISORY_CIRCULAR = auto()
    AIRWORTHINESS_CRITERIA = auto()
    DEVIATION_MEMO = auto()
    EXEMPTION = auto()
    FEDERAL_REGISTER_NOTICE = auto()
    ORDER = auto()
    POLICY_STATEMENT = auto()
    RULE = auto()
    SPECIAL_CONDITION = auto()
    TECHNICAL_STANDARD_ORDER = auto()
    OTHER = auto()

    @classmethod
    def from_string(cls, doc_type: str) -> 'DocumentType':
        """Convert string to DocumentType, case-insensitive."""
        try:
            return cls[doc_type.upper().replace(" ", "_")]
        except KeyError:
            raise DocumentTypeError(f"Unsupported document type: {doc_type}")

# 3. Utility Classes
@dataclass
class TextNormalization:
    """Text normalization utilities."""
    
    @staticmethod
    def normalize_heading(text: str) -> str:
        """Normalize heading text for consistent comparison."""
        # Remove excess whitespace
        text = ' '.join(text.split())
        # Normalize periods (convert multiple periods to single period)
        text = re.sub(r'\.+$', '.', text.strip())
        # Remove any whitespace before the period
        text = re.sub(r'\s+\.$', '.', text)
        return text
    
    @staticmethod
    def normalize_document_type(doc_type: str) -> str:
        """Normalize document type string."""
        return ' '.join(word.capitalize() for word in doc_type.lower().split())

# 4. Result Class
@dataclass
class DocumentCheckResult:
    success: bool
    issues: List[Dict[str, Any]]
    details: Optional[Dict[str, Any]] = None

# 5. Base Document Checker
class DocumentChecker:
    def __init__(self, config_path: Optional[str] = None):
        self.config_manager = DocumentCheckerConfig(config_path)
        self.logger = self.config_manager.logger

    @classmethod
    def extract_paragraphs(cls, doc_path: str) -> List[str]:
        """Extract paragraphs from a document."""
        try:
            doc = Document(doc_path)
            if not doc.paragraphs:
                raise DocumentCheckError("Document appears to be empty or invalid")
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    paragraphs.append(para.text.strip())
            
            if not paragraphs:
                raise DocumentCheckError("No text content found in document")
            
            return paragraphs
        except PackageNotFoundError as e:
            raise DocumentCheckError(f"Invalid document format: {str(e)}")
        except Exception as e:
            raise DocumentCheckError(f"Error reading document: {str(e)}")

    @staticmethod
    def validate_input(doc: List[str]) -> bool:
        return doc is not None and isinstance(doc, list) and len(doc) > 0

# 6. Configuration Manager
class DocumentCheckerConfig:
    
    REQUIRED_CONFIG_KEYS = {'logging', 'checks', 'document_types'}
    REQUIRED_LOGGING_KEYS = {'level', 'format'}
    REQUIRED_CHECKS_KEYS = {'acronyms', 'terminology_check', 'headings'}
    
    def __init__(self, config_path: Optional[str] = None):
        """Initialize configuration with optional config file."""
        self.default_config = {
            "logging": {
                "level": "INFO",
                "format": "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
            },
            "checks": {
                "acronyms": True,
                "terminology_check": True,
                "headings": True
            },
            "document_types": {
                "Advisory Circular": {
                    "required_headings": [
                        "Purpose.",
                        "Applicability.",
                        "Cancellation.",
                        "Related Material.",
                        "Definition of Key Terms."
                    ],
                    "skip_title_check": False
                },
                "Federal Register Notice": {
                    "required_headings": [
                        "Purpose of This Notice",
                        "Audience",
                        "Where can I Find This Notice"
                    ],
                    "skip_title_check": False
                },
                "Order": {
                    "required_headings": [
                        "Purpose of This Order.",
                        "Audience.",
                        "Where to Find This Order."
                    ],
                    "skip_title_check": False
                },
                "Policy Statement": {
                    "required_headings": [
                        "SUMMARY",
                        "CURRENT REGULATORY AND ADVISORY MATERIAL",
                        "RELEVANT PAST PRACTICE",
                        "POLICY",
                        "EFFECT OF POLICY",
                        "CONCLUSION"
                    ],
                    "skip_title_check": False
                },
                "Technical Standard Order": {
                    "required_headings": [
                        "PURPOSE.",
                        "APPLICABILITY.",
                        "REQUIREMENTS.",
                        "MARKING.",
                        "APPLICATION DATA REQUIREMENTS.",
                        "MANUFACTURER DATA REQUIREMENTS.",
                        "FURNISHED DATA REQUIREMENTS.",
                        "HOW TO GET REFERENCED DOCUMENTS."
                    ],
                    "skip_title_check": False
                },
                "Airworthiness Criteria": {
                    "required_headings": [],
                    "skip_title_check": True
                },
                "Deviation Memo": {
                    "required_headings": [],
                    "skip_title_check": True
                },
                "Exemption": {
                    "required_headings": [],
                    "skip_title_check": True
                },
                "Rule": {
                    "required_headings": [],
                    "skip_title_check": True
                },
                "Special Condition": {
                    "required_headings": [],
                    "skip_title_check": True
                },
                "Other": {
                    "required_headings": [],
                    "skip_title_check": True
                }        
            }
        }

        self.config = self._load_config(config_path)
        self._validate_config(self.config)
        self.logger = self._setup_logger()
        self.pattern_registry = self._setup_patterns()

    def _load_config(self, config_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Load configuration from JSON file or use default settings.
        
        Args:
            config_path (str, optional): Path to configuration file.
            
        Returns:
            Dict[str, Any]: Loaded configuration dictionary.
        """
        if config_path and os.path.exists(config_path):
            try:
                with open(config_path, 'r') as f:
                    user_config = json.load(f)
                    # Deep merge default and user config
                    return self._deep_merge(self.default_config.copy(), user_config)
            except (json.JSONDecodeError, IOError) as e:
                logging.warning(f"Error loading config: {e}. Using default config.")
                return self.default_config.copy()
        return self.default_config.copy()

    def _validate_config(self, config: Dict[str, Any]) -> None:
        """Validate the configuration structure."""
        # Check for required top-level keys
        missing_keys = self.REQUIRED_CONFIG_KEYS - set(config.keys())
        if missing_keys:
            raise ConfigurationError(f"Missing required configuration keys: {missing_keys}")
            
        # Validate logging configuration
        if 'logging' in config:
            missing_logging_keys = self.REQUIRED_LOGGING_KEYS - set(config['logging'].keys())
            if missing_logging_keys:
                raise ConfigurationError(f"Missing required logging configuration keys: {missing_logging_keys}")
                
        # Validate checks configuration
        if 'checks' in config:
            missing_checks_keys = self.REQUIRED_CHECKS_KEYS - set(config['checks'].keys())
            if missing_checks_keys:
                raise ConfigurationError(f"Missing required checks configuration keys: {missing_checks_keys}")
                
        # Validate document types configuration
        if 'document_types' not in config:
            raise ConfigurationError("Missing document_types configuration")
            
        # Validate patterns configuration if present
        if 'patterns' in config:
            if not isinstance(config['patterns'], dict):
                raise ConfigurationError("Patterns configuration must be a dictionary")
                
            # Validate boilerplate patterns if present
            if 'boilerplate' in config['patterns']:
                if not isinstance(config['patterns']['boilerplate'], dict):
                    raise ConfigurationError("Boilerplate patterns must be a dictionary")
                for doc_type, patterns in config['patterns']['boilerplate'].items():
                    if not isinstance(patterns, list):
                        raise ConfigurationError(f"Boilerplate patterns for {doc_type} must be a list")
                        
            # Validate required language patterns if present
            if 'required_language' in config['patterns']:
                if not isinstance(config['patterns']['required_language'], dict):
                    raise ConfigurationError("Required language patterns must be a dictionary")
                for doc_type, patterns in config['patterns']['required_language'].items():
                    if not isinstance(patterns, list):
                        raise ConfigurationError(f"Required language patterns for {doc_type} must be a list")

    def _deep_merge(self, base: Dict[str, Any], update: Dict[str, Any]) -> Dict[str, Any]:
        """
        Recursively merge two dictionaries.

        Args:
            base (Dict): Base dictionary to merge into.
            update (Dict): Dictionary to merge from.

        Returns:
            Dict: Merged dictionary.
        """
        for key, value in update.items():
            if isinstance(value, dict) and key in base and isinstance(base[key], dict):
                self._deep_merge(base[key], value)
            else:
                base[key] = value
        return base

    def _setup_logger(self) -> logging.Logger:
        """
        Set up and configure logging based on configuration.

        Returns:
            logging.Logger: Configured logger instance.
        """
        logger = logging.getLogger(__name__)
        log_level = getattr(logging, self.config['logging']['level'].upper())

        formatter = logging.Formatter(self.config['logging']['format'])

        # Console Handler
        console_handler = logging.StreamHandler()
        console_handler.setFormatter(formatter)
        console_handler.setLevel(log_level)

        logger.addHandler(console_handler)
        logger.setLevel(log_level)

        return logger
    
    def _setup_patterns(self) -> Dict[str, List[PatternConfig]]:
        """Set up pattern configurations for document checking."""
        patterns = {}
        
        # Load patterns from JSON file
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            patterns_file = os.path.join(current_dir, 'patterns.json')
            
            with open(patterns_file, 'r') as f:
                patterns_data = json.load(f)
                
            # Convert JSON data to PatternConfig objects
            for category, pattern_list in patterns_data.items():
                if category in ['boilerplate', 'required_language']:
                    # Handle document type specific patterns
                    patterns[category] = []
                    for doc_type, type_patterns in pattern_list.items():
                        for pattern in type_patterns:
                            patterns[category].append(
                                PatternConfig(
                                    pattern=pattern,
                                    description=f"{category} text for {doc_type}",
                                    is_error=category == 'required_language',
                                    keep_together=True
                                )
                            )
                else:
                    # Handle regular patterns
                    patterns[category] = [
                        PatternConfig(
                            pattern=p['pattern'],
                            description=p['description'],
                            is_error=p.get('is_error', False),
                            replacement=p.get('replacement'),
                            keep_together=p.get('keep_together', False)
                        ) for p in pattern_list
                    ]
                    
        except (json.JSONDecodeError, IOError) as e:
            self.logger.error(f"Error loading patterns: {e}")
            # Return empty patterns dictionary if file loading fails
            return {}
            
        return patterns

def profile_performance(func):
    """Decorator to profile function performance."""
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        # Get logger from the class instance (first argument)
        logger = args[0].logger
        logger.info(
            f"Performance: {func.__name__} took {end_time - start_time:.4f} seconds"
        )
        return result
    return wrapper

# 8. FAA Document Checker
class FAADocumentChecker(DocumentChecker):
    """Document checker implementation for FAA documents."""
    
    # Class Constants
    PERIOD_REQUIRED = MappingProxyType({
        DocumentType.ADVISORY_CIRCULAR: True,
        DocumentType.AIRWORTHINESS_CRITERIA: False,
        DocumentType.DEVIATION_MEMO: False,
        DocumentType.EXEMPTION: False,
        DocumentType.FEDERAL_REGISTER_NOTICE: False,
        DocumentType.ORDER: True,
        DocumentType.POLICY_STATEMENT: False,
        DocumentType.RULE: False,
        DocumentType.SPECIAL_CONDITION: False,
        DocumentType.TECHNICAL_STANDARD_ORDER: True,
        DocumentType.OTHER: False
    })
    
    HEADING_WORDS = frozenset({
        'APPLICABILITY', 'APPENDIX', 'AUTHORITY', 'BACKGROUND', 'CANCELLATION', 'CAUTION',
        'CHAPTER', 'CONCLUSION', 'DEPARTMENT', 'DEFINITION', 'DEFINITIONS', 'DISCUSSION',
        'DISTRIBUTION', 'EXCEPTION', 'EXPLANATION', 'FIGURE', 'GENERAL', 'GROUPS', 
        'INFORMATION', 'INSERT', 'INTRODUCTION', 'MATERIAL', 'NOTE', 'PARTS', 'PAST', 
        'POLICY', 'PRACTICE', 'PROCEDURES', 'PURPOSE', 'RELEVANT', 'RELATED', 
        'REQUIREMENTS', 'REPORT', 'SCOPE', 'SECTION', 'SUMMARY', 'TABLE', 'WARNING'
    })
    
    PREDEFINED_ACRONYMS = frozenset({
        'AGC', 'AIR', 'CFR', 'DC', 'DOT', 'FAA IR-M', 'FAQ', 'i.e.', 'e.g.', 'MA',
        'MD', 'MIL', 'MO', 'No.', 'PDF', 'SAE', 'SSN', 'TX', 'U.S.', 'U.S.C.', 'USA', 'US', 
        'WA', 'XX', 'ZIP', 'ACO', 'RGL'
    })

    # Boilerplate text patterns to ignore in length checks
    BOILERPLATE_PATTERNS = MappingProxyType({
        DocumentType.ADVISORY_CIRCULAR: (
            r"This is a guidance document\. Its content is not legally binding in its own right",
            r"The guidance provided in this AC is for manufacturers, modifiers, foreign regulatory authorities",
            r"The contents of this AC do not have the force and effect of law",
            r"This material does not change or create any additional regulatory requirements",
            r"If you find an error in this Advisory Circular",
            r"For your convenience, the AC Feedback Form is the last page of this AC"
        ),
        DocumentType.POLICY_STATEMENT: (
            r"The contents of this policy statement do not have the force and effect of law",
            r"This policy statement does not constitute a new regulation",
            r"If a proposed method of compliance appears to differ from the guidance expressed in this policy statement",
            r"Additional information on the effect of FAA policy statements may be found in FAA Order IR 8100\.16"
        ),
        DocumentType.FEDERAL_REGISTER_NOTICE: (
            r"Except for Confidential Business Information \(CBI\) as described in the following paragraph",
            r"Confidential Business Information \(CBI\) is commercial or financial information",
            r"Paperwork Reduction Act Burden Statement",
            r"The FAA invites interested people to take part in this rulemaking",
            r"The FAA will consider all comments received by the closing date for comments"
        ),
        DocumentType.SPECIAL_CONDITION: (
            r"Except for Confidential Business Information \(CBI\) as described in the following paragraph",
            r"Confidential Business Information \(CBI\) is commercial or financial information",
            r"Paperwork Reduction Act Burden Statement",
            r"The FAA invites interested people to take part in this rulemaking",
            r"The FAA will consider all comments received by the closing date for comments"
        )
    })

    # Required language patterns for different document types
    REQUIRED_LANGUAGE_PATTERNS = {
        DocumentType.ADVISORY_CIRCULAR: [
            r"This is a guidance document\. Its content is not legally binding in its own right",
            r"The guidance provided in this AC is for manufacturers, modifiers, foreign regulatory authorities",
            r"The contents of this AC do not have the force and effect of law",
            r"This material does not change or create any additional regulatory requirements",
            r"If you find an error in this Advisory Circular",
            r"For your convenience, the AC Feedback Form is the last page of this AC"
        ],
        DocumentType.POLICY_STATEMENT: [
            r"The contents of this policy statement do not have the force and effect of law",
            r"This policy statement does not constitute a new regulation",
            r"If a proposed method of compliance appears to differ from the guidance expressed in this policy statement",
            r"Additional information on the effect of FAA policy statements may be found in FAA Order IR 8100\.16"
        ],
        DocumentType.FEDERAL_REGISTER_NOTICE: [
            r"Except for Confidential Business Information \(CBI\) as described in the following paragraph",
            r"Confidential Business Information \(CBI\) is commercial or financial information",
            r"Paperwork Reduction Act Burden Statement",
            r"The FAA invites interested people to take part in this rulemaking",
            r"The FAA will consider all comments received by the closing date for comments"
        ],
        DocumentType.SPECIAL_CONDITION: [
            r"Except for Confidential Business Information \(CBI\) as described in the following paragraph",
            r"Confidential Business Information \(CBI\) is commercial or financial information",
            r"Paperwork Reduction Act Burden Statement",
            r"The FAA invites interested people to take part in this rulemaking",
            r"The FAA will consider all comments received by the closing date for comments"
        ]
    }

    # Constructor
    def __init__(self, config_path: Optional[str] = None):
        super().__init__(config_path)
        self.pattern_cache = PatternCache()
        
    def _compile_patterns(self, pattern_type: str) -> Dict[DocumentType, List[Pattern]]:
        """Compile patterns for a specific type."""
        patterns = {}
        for doc_type, configs in self.config_manager.pattern_registry.get(pattern_type, {}).items():
            patterns[doc_type] = [self.pattern_cache.get_pattern(config.pattern) for config in configs]
        return patterns

    def _get_doc_type_config(self, doc_type: str) -> Tuple[Dict[str, Any], bool]:
        """
        Get document type configuration and validate document type.
        
        Args:
            doc_type: Type of document being checked
            
        Returns:
            Tuple containing:
                - Document type configuration dictionary
                - Boolean indicating if document type is valid
                
        Raises:
            DocumentTypeError: If document type is invalid
        """
        # Validate document type
        doc_type_config = self.config_manager.config['document_types'].get(doc_type)
        if not doc_type_config:
            self.logger.error(f"Unsupported document type: {doc_type}")
            raise DocumentTypeError(f'Unsupported document type: {doc_type}')
            
        return doc_type_config, True

    @profile_performance
    def heading_title_check(self, doc: List[str], doc_type: str) -> DocumentCheckResult:
        if not self.validate_input(doc):
            self.logger.error("Invalid document input for heading check")
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])

        try:
            # Use the new helper method
            doc_type_config, _ = self._get_doc_type_config(doc_type)
        except DocumentTypeError as e:
            return DocumentCheckResult(success=False, issues=[{'error': str(e)}])

        # Get configuration for document-specific headings
        required_headings = doc_type_config.get('required_headings', [])

        if not required_headings:
            return DocumentCheckResult(
                success=True, 
                issues=[], 
                details={'message': f'No required headings defined for {doc_type}'}
            )

        headings_found = []
        required_headings_set = set(required_headings)

        # Precompile a regex pattern to match headings at the start of the paragraph
        # Escape special characters in headings and allow for optional spaces and periods
        heading_patterns = []
        for heading in required_headings:
            escaped_heading = re.escape(heading.rstrip('.'))
            pattern = rf'^\s*{escaped_heading}\.?\s*'
            heading_patterns.append(pattern)
        combined_pattern = re.compile('|'.join(heading_patterns), re.IGNORECASE)

        for para in doc:
            para_strip = para.strip()
            # Check if paragraph starts with any of the required headings
            match = combined_pattern.match(para_strip)
            if match:
                # Extract the matched heading
                matched_heading = match.group().strip()
                # Normalize the matched heading to compare with required headings
                matched_heading_base = matched_heading.rstrip('.').strip()
                # Find the exact heading from required headings (case-insensitive)
                for required_heading in required_headings:
                    if matched_heading_base.lower() == required_heading.rstrip('.').lower():
                        headings_found.append(required_heading)
                        break

        # Check if all required headings are found
        found_headings_set = set(headings_found)
        missing_headings = required_headings_set - found_headings_set
        unexpected_headings = found_headings_set - required_headings_set

        success = len(missing_headings) == 0
        issues = []
        
        if not success:
            issues.append({
                'type': 'missing_headings',
                'missing': list(missing_headings)
            })

        if unexpected_headings:
            issues.append({
                'type': 'unexpected_headings',
                'unexpected': list(unexpected_headings)
            })

        details = {
            'found_headings': list(found_headings_set),
            'required_headings': required_headings,
            'document_type': doc_type,
            'missing_count': len(missing_headings),
            'unexpected_count': len(unexpected_headings)
        }

        return DocumentCheckResult(success=success, issues=issues, details=details)

    @profile_performance
    def heading_title_period_check(self, doc: List[str], doc_type: str) -> DocumentCheckResult:
        """
        Check if headings have the correct period usage based on document type.
        """
        if not self.validate_input(doc):
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])

        issues = []
        doc_type_enum = DocumentType.from_string(doc_type)
        requires_period = self.PERIOD_REQUIRED.get(doc_type_enum, False)

        for paragraph in doc:
            # Skip empty paragraphs
            if not paragraph.strip():
                continue

            # Check if this is a heading (starts with a number followed by a period)
            match = re.match(r'^(\d+\.)\s*(.*)', paragraph.strip())
            if match:
                number, heading_text = match.groups()
                heading_text = heading_text.strip()

                # Check if heading ends with a period
                has_period = heading_text.endswith('.')
                
                if requires_period and not has_period:
                    issues.append({
                        'heading': heading_text,
                        'message': f"Heading '{heading_text}' should end with a period"
                    })
                elif not requires_period and has_period:
                    issues.append({
                        'heading': heading_text,
                        'message': f"Heading '{heading_text}' should not end with a period"
                    })

        return DocumentCheckResult(success=len(issues) == 0, issues=issues)

    @profile_performance
    def spacing_check(self, doc: List[str]) -> DocumentCheckResult:
        """Check for correct spacing in the document."""
        if not self.validate_input(doc):
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])

        issues = []
        for i, paragraph in enumerate(doc):
            # Check for multiple spaces
            if '  ' in paragraph:
                issues.append({
                    'issue_type': 'multiple_spaces',
                    'paragraph': i + 1,
                    'text': paragraph
                })
            
            # Check for space before punctuation
            if re.search(r'\s+[.,!?]', paragraph):
                issues.append({
                    'issue_type': 'space_before_punctuation',
                    'paragraph': i + 1,
                    'text': paragraph
                })
            
            # Check for missing space after punctuation
            if re.search(r'[.,!?][^ \n]', paragraph):
                issues.append({
                    'issue_type': 'missing_space_after_punctuation',
                    'paragraph': i + 1,
                    'text': paragraph
                })

        success = len(issues) == 0
        return DocumentCheckResult(success=success, issues=issues)

    @profile_performance
    def check_parentheses(self, doc: List[str]) -> DocumentCheckResult:
        """
        Check for matching parentheses in the document.

        Args:
            doc (List[str]): List of document paragraphs

        Returns:
            DocumentCheckResult: Result containing any mismatched parentheses issues
        """
        if not self.validate_input(doc):
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])

        issues = []
        
        for i, paragraph in enumerate(doc, 1):
            if not paragraph.strip():  # Skip empty paragraphs
                continue
            
            stack = []  # Track unmatched opening parentheses
            sentences = re.split(r'(?<=[.!?])\s+', paragraph)  # Split paragraph into sentences
            for sentence in sentences:
                for j, char in enumerate(sentence):
                    if char == '(':
                        stack.append((sentence, j))  # Store sentence and character position
                    elif char == ')':
                        if stack:
                            stack.pop()  # Remove matching opening parenthesis
                        else:
                            # No matching opening parenthesis
                            issues.append({
                                'type': 'missing_opening',
                                'paragraph': i,  # Still tracked but not included in the message
                                'position': j,
                                'text': sentence,
                                'message': f"Add an opening parenthesis to the sentence: \"{sentence.strip()}\""
                            })

            # Check for any unmatched opening parentheses left in the stack
            for unmatched in stack:
                sentence, pos = unmatched
                issues.append({
                    'type': 'missing_closing',
                    'paragraph': i,  # Still tracked but not included in the message
                    'position': pos,
                    'text': sentence,
                    'message': f"Add a closing parenthesis to the sentence: \"{sentence.strip()}\""
                })

        return DocumentCheckResult(success=len(issues) == 0, issues=issues)

    @profile_performance
    def check_abbreviation_usage(self, doc: List[str]) -> DocumentCheckResult:
        """Check for abbreviation consistency after first definition."""
        if not self.validate_input(doc):
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])

        # Track abbreviations and their usage
        abbreviations = {}  # Store defined abbreviations
        inconsistent_uses = []  # Track full term usage after definition

        def process_sentence(sentence: str) -> None:
            """Process a single sentence for abbreviation usage."""
            for acronym, data in abbreviations.items():
                full_term = data["full_term"]
                if full_term not in sentence:
                    continue
                    
                # Skip if this is the definition sentence
                if sentence.strip() == data["first_occurrence"]:
                    continue
                    
                # Track inconsistent usage
                if not data["defined"]:
                    inconsistent_uses.append({
                        'issue_type': 'full_term_after_acronym',
                        'full_term': full_term,
                        'acronym': acronym,
                        'sentence': sentence.strip(),
                        'definition_context': data["first_occurrence"]
                    })
                data["defined"] = False  # Mark as used

        # Process each paragraph
        for paragraph in doc:
            sentences = re.split(r'(?<=[.!?])\s+', paragraph)
            for sentence in sentences:
                process_sentence(sentence.strip())

        success = len(inconsistent_uses) == 0
        return DocumentCheckResult(success=success, issues=inconsistent_uses)

    @profile_performance
    def check_date_format_usage(self, doc: List[str]) -> DocumentCheckResult:
        """
        Check document for date format usage issues.
        
        Args:
            doc: List of document paragraphs
            
        Returns:
            DocumentCheckResult containing any date format issues found
        """
        # Define patterns to ignore (e.g., AC references)
        ignore_patterns = [
            r'AC \d{2}-\d{2,3}[A-Z]?',  # AC 20-100, AC 20-100A
            r'AC \d{2}-\d{2,3}[A-Z]?-\d{1,2}',  # AC 20-100-1
            r'AC \d{2}-\d{2,3}[A-Z]?-\d{1,2}[A-Z]?',  # AC 20-100-1A
        ]
        
        # Compile ignore patterns once
        compiled_ignore_patterns = [self.pattern_cache.get_pattern(p) for p in ignore_patterns]
        
        # Track unique issues
        date_format_issues = set()
        
        for paragraph in doc:
            # Temporarily remove ignored text
            temp_paragraph = paragraph
            for pattern in compiled_ignore_patterns:
                temp_paragraph = pattern.sub('', temp_paragraph)
            
            # Check for date formats
            date_patterns = self.config_manager.pattern_registry.get('date_formats', [])
            for pattern_config in date_patterns:
                pattern = self.pattern_cache.get_pattern(pattern_config.pattern)
                for match in pattern.finditer(temp_paragraph):
                    date_format_issues.add(match.group())

        # Compile issues
        issues = []
        if date_format_issues:
            issues.append({
                'issue_type': 'date_format',
                'description': 'Inconsistent date format usage',
                'occurrences': list(date_format_issues)
            })

        details = {
            'total_date_formats': len(date_format_issues),
            'date_formats': list(date_format_issues)
        }

        return DocumentCheckResult(success=len(issues) == 0, issues=issues, details=details)

    @profile_performance
    def check_placeholder_usage(self, doc: List[str]) -> DocumentCheckResult:
        """
        Check document for placeholder usage issues.
        
        Args:
            doc: List of document paragraphs
            
        Returns:
            DocumentCheckResult containing any placeholder issues found
        """
        def process_placeholders(doc: List[str], patterns: List[PatternConfig]) -> DocumentCheckResult:
            # Track unique issues
            tbd_placeholders = set()
            to_be_determined_placeholders = set()
            to_be_added_placeholders = set()
            
            # Compile patterns once
            tbd_pattern = self.pattern_cache.get_pattern(r'\bTBD\b')
            to_be_determined_pattern = self.pattern_cache.get_pattern(r'\bTo be determined\b')
            to_be_added_pattern = self.pattern_cache.get_pattern(r'\bTo be added\b')
            
            for paragraph in doc:
                # Check for TBD
                for match in tbd_pattern.finditer(paragraph):
                    tbd_placeholders.add(match.group())
                
                # Check for "To be determined"
                for match in to_be_determined_pattern.finditer(paragraph):
                    to_be_determined_placeholders.add(match.group())
                
                # Check for "To be added"
                for match in to_be_added_pattern.finditer(paragraph):
                    to_be_added_placeholders.add(match.group())

            # Compile issues
            issues = []
            if tbd_placeholders:
                issues.append({
                    'issue_type': 'tbd_placeholder',
                    'description': 'Remove TBD placeholder',
                    'occurrences': list(tbd_placeholders)
                })
                
            if to_be_determined_placeholders:
                issues.append({
                    'issue_type': 'to_be_determined_placeholder',
                    'description': "Remove 'To be determined' placeholder",
                    'occurrences': list(to_be_determined_placeholders)
                })
                
            if to_be_added_placeholders:
                issues.append({
                    'issue_type': 'to_be_added_placeholder',
                    'description': "Remove 'To be added' placeholder",
                    'occurrences': list(to_be_added_placeholders)
                })

            details = {
                'total_placeholders': len(tbd_placeholders) + 
                                    len(to_be_determined_placeholders) + 
                                    len(to_be_added_placeholders),
                'placeholder_types': {
                    'TBD': len(tbd_placeholders),
                    'To be determined': len(to_be_determined_placeholders),
                    'To be added': len(to_be_added_placeholders)
                }
            }

            return DocumentCheckResult(success=len(issues) == 0, issues=issues, details=details)

        return self._process_patterns(doc, 'placeholders', process_placeholders)

    @profile_performance
    def check_paragraph_length(self, doc: List[str]) -> DocumentCheckResult:
        """Check paragraph lengths, ignoring boilerplate text."""
        if not self.validate_input(doc):
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])
            
        issues = []
        doc_type = self._get_doc_type(doc)
        boilerplate_patterns = self.boilerplate_patterns.get(doc_type, [])
        
        for i, paragraph in enumerate(doc, 1):
            # Skip if paragraph matches any boilerplate pattern
            if any(pattern.search(paragraph) for pattern in boilerplate_patterns):
                continue
                
            # Count sentences (split on period, exclamation, question mark followed by space)
            sentences = re.split(r'(?<=[.!?])\s+', paragraph.strip())
            # Count lines (split on newlines or when length exceeds ~80 characters)
            lines = []
            current_line = ""
            
            for word in paragraph.split():
                if len(current_line) + len(word) + 1 <= 80:
                    current_line += " " + word if current_line else word
                else:
                    lines.append(current_line)
                    current_line = word
            if current_line:
                lines.append(current_line)
            
            # Check if paragraph exceeds either threshold
            if len(sentences) > 6 or len(lines) > 8:
                # Get first sentence for context
                first_sentence = sentences[0].strip()
                issues.append(f"Review the paragraph that starts with: \"{first_sentence}\"")
        
        return DocumentCheckResult(success=len(issues) == 0, issues=issues)
    
    @profile_performance
    def check_sentence_length(self, doc: List[str]) -> DocumentCheckResult:
        """Check sentence lengths, ignoring boilerplate text."""
        if not self.validate_input(doc):
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])
            
        issues = []
        doc_type = self._get_doc_type(doc)
        boilerplate_patterns = self.boilerplate_patterns.get(doc_type, [])
        
        for i, (sentence, paragraph) in enumerate(self._process_sentences(doc), 1):
            # Skip if sentence matches any boilerplate pattern
            if any(pattern.search(sentence) for pattern in boilerplate_patterns):
                continue
                
            # Count words (splitting on whitespace)
            words = sentence.split()
            word_count = len(words)
            
            if word_count > 35:
                issues.append({
                    'sentence': sentence,
                    'word_count': word_count
                })
        
        return DocumentCheckResult(
            success=len(issues) == 0,
            issues=issues,
            details={
                'total_sentences': len(issues),
                'long_sentences': len(issues),
                'max_length': max(issue['word_count'] for issue in issues) if issues else 0,
                'avg_length': sum(issue['word_count'] for issue in issues) / len(issues) if issues else 0
            }
        )
    
    @profile_performance
    def check_section_508_compliance(self, doc_path: str) -> DocumentCheckResult:
        """
        Check if the document complies with Section 508 accessibility requirements.
        
        Args:
            doc_path (str): Path to the document file to be checked.
            
        Returns:
            DocumentCheckResult: A result object containing:
                - success (bool): True if the document passes all checks
                - issues (List[Dict[str, Any]]): List of accessibility issues found
                - details (Optional[Dict[str, Any]]): Additional details about the check results
                
        Raises:
            IOError: If there are issues reading the document file
            OSError: If there are system-level issues accessing the file
            ValueError: If the document format is invalid
            TypeError: If the document content is not in the expected format
        """
        try:
            doc = Document(doc_path)
            issues = []
            images_with_alt = 0
            heading_structure = {}
            heading_issues = []  # Separate list for heading-specific issues
            hyperlink_issues = []  # New list for hyperlink issues

            # Image alt text check
            for shape in doc.inline_shapes:
                alt_text = None
                if hasattr(shape, '_inline') and hasattr(shape._inline, 'docPr'):
                    docPr = shape._inline.docPr
                    alt_text = docPr.get('descr') or docPr.get('title')

                if alt_text:
                    images_with_alt += 1
                else:
                    issues.append({
                        'category': 'image_alt_text',
                        'message': 'Image is missing descriptive alt text.',
                        'context': 'Ensure all images have descriptive alt text.'
                    })

            # Enhanced heading structure check
            headings = []
            
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    try:
                        level = int(paragraph.style.name.split()[-1])
                        text = paragraph.text.strip()
                        
                        if not text:
                            continue
                            
                        headings.append((text, level))
                        heading_structure[level] = heading_structure.get(level, 0) + 1
                        
                    except ValueError:
                        continue

            # Check heading hierarchy
            if headings:
                min_level = min(level for _, level in headings)
                
                if min_level > 1:
                    heading_issues.append({
                        'severity': 'error',
                        'type': 'missing_h1',
                        'message': 'Document should start with a Heading 1',
                        'context': f"First heading found is level {headings[0][1]}: '{headings[0][0]}'",
                        'recommendation': 'Add a Heading 1 at the start of the document'
                    })

                # Check for skipped levels
                previous_heading = None
                for text, level in headings:
                    if previous_heading:
                        prev_text, prev_level = previous_heading
                        
                        # Only check for skipped levels when going deeper
                        if level > prev_level + 1:
                            missing_levels = list(range(prev_level + 1, level))
                            heading_issues.append({
                                'severity': 'error',
                                'type': 'skipped_levels',
                                'message': f"Skipped heading level(s) {', '.join(map(str, missing_levels))} - Found H{level} '{text}' after H{prev_level} '{prev_text}'. Add H{prev_level + 1} before this section.",
                            })
                        
                    previous_heading = (text, level)

            # Enhanced Hyperlink Accessibility Check
            for paragraph in doc.paragraphs:
                # Check both hyperlink fields and runs with hyperlink formatting
                hyperlinks = []
                
                # Method 1: Check for hyperlink fields
                if hasattr(paragraph, '_element') and hasattr(paragraph._element, 'xpath'):
                    hyperlinks.extend(paragraph._element.xpath('.//w:hyperlink'))
                
                # Method 2: Check for hyperlink style runs
                for run in paragraph.runs:
                    if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                        if run._element.rPr is not None:
                            if run._element.rPr.xpath('.//w:rStyle[@w:val="Hyperlink"]'):
                                hyperlinks.append(run)
                    
                    # Method 3: Check for direct hyperlink elements
                    if hasattr(run, '_r'):
                        if run._r.xpath('.//w:hyperlink'):
                            hyperlinks.append(run)

                # Process found hyperlinks
                for hyperlink in hyperlinks:
                    # Extract link text based on element type
                    if hasattr(hyperlink, 'text'):  # For run objects
                        link_text = hyperlink.text.strip()
                    else:  # For hyperlink elements
                        link_text = ''.join([t.text for t in hyperlink.xpath('.//w:t')])
                    
                    if not link_text:  # Skip empty links
                        continue

                    # Check for accessibility issues
                    non_descriptive = [
                        'click here', 'here', 'link', 'this link', 'more', 
                        'read more', 'learn more', 'click', 'see this', 
                        'see here', 'go', 'url', 'this', 'page'
                    ]
                    
                    if any(phrase == link_text.lower() for phrase in non_descriptive):
                        hyperlink_issues.append({
                            'category': 'hyperlink_accessibility',
                            'severity': 'warning',
                            'message': 'Non-descriptive hyperlink text detected',
                            'context': f'Link text: "{link_text}"',
                            'recommendation': 'Replace with descriptive text that indicates the link destination',
                            'user_message': f'Replace non-descriptive link text "{link_text}" with text that clearly indicates where the link will take the user'
                        })
                    elif len(link_text.strip()) < 4:  # Check for very short link text
                        hyperlink_issues.append({
                            'category': 'hyperlink_accessibility',
                            'severity': 'warning',
                            'message': 'Hyperlink text may be too short to be meaningful',
                            'context': f'Link text: "{link_text}"',
                            'recommendation': 'Use longer, more descriptive text that indicates the link destination',
                            'user_message': f'Link text "{link_text}" is too short - use descriptive text that clearly indicates the link destination'
                        })
                    elif link_text.lower().startswith(('http', 'www', 'ftp')):
                        hyperlink_issues.append({
                            'category': 'hyperlink_accessibility',
                            'severity': 'warning',
                            'message': 'Raw URL used as link text',
                            'context': f'Link text: "{link_text}"',
                            'recommendation': 'Replace the URL with descriptive text that indicates the link destination',
                            'user_message': f'Replace the URL "{link_text}" with meaningful text that describes the link destination'
                        })

            # Add hyperlink issues to main issues list
            if hyperlink_issues:
                issues.extend(hyperlink_issues)

            # Combine all issues
            if heading_issues:
                issues.extend([{
                    'category': '508_compliance_heading_structure',
                    **issue
                } for issue in heading_issues])

            # Enhanced details with heading structure information
            details = {
                'total_images': len(doc.inline_shapes),
                'images_with_alt': images_with_alt,
                'heading_structure': {
                    'total_headings': len(headings),
                    'levels_found': dict(sorted(heading_structure.items())),
                    'hierarchy_depth': max(heading_structure.keys()) if heading_structure else 0,
                    'heading_sequence': [(text[:50] + '...' if len(text) > 50 else text, level) 
                                       for text, level in headings],
                    'issues_found': len(heading_issues)
                },
                'hyperlink_accessibility': {  # New details section
                    'total_issues': len(hyperlink_issues),
                    'non_descriptive_links': sum(1 for issue in hyperlink_issues 
                                               if 'Non-descriptive' in issue['message']),
                    'raw_urls': sum(1 for issue in hyperlink_issues 
                                  if 'Raw URL' in issue['message'])
                }
            }

            return DocumentCheckResult(
                success=len(issues) == 0,
                issues=issues,
                details=details
            )
            
        except (IOError, OSError) as e:
            self.logger.error(f"File system error during 508 compliance check: {str(e)}")
            return DocumentCheckResult(
                success=False,
                issues=[{
                    'category': 'error',
                    'message': f'File system error during 508 compliance check: {str(e)}'
                }]
            )
        except (ValueError, TypeError) as e:
            self.logger.error(f"Data processing error during 508 compliance check: {str(e)}")
            return DocumentCheckResult(
                success=False,
                issues=[{
                    'category': 'error',
                    'message': f'Data processing error during 508 compliance check: {str(e)}'
                }]
            )
        except Exception as e:
            self.logger.error(f"Unexpected error during 508 compliance check: {str(e)}")
            return DocumentCheckResult(
                success=False,
                issues=[{
                    'category': 'error',
                    'message': f'Unexpected error during 508 compliance check: {str(e)}'
                }]
            )

    def _format_compliance_issues(self, result: DocumentCheckResult) -> List[str]:
        """Format compliance issues with clear, user-friendly descriptions."""
        formatted_issues = []
        
        for issue in result.issues:
            if issue.get('category') == '508_compliance_heading_structure':
                # Existing heading structure formatting...
                message = issue.get('message', 'No description provided')
                context = issue.get('context', 'No context provided').strip()
                recommendation = issue.get('recommendation', 'No recommendation provided').strip()
                formatted_issues.append(
                    f"    • {message}. Context: {context}. Recommendation: {recommendation}"
                )
            elif issue.get('category') == 'image_alt_text':
                # Existing alt text formatting...
                formatted_issues.append(
                    f"    • {issue.get('message', 'No description provided')}. {issue.get('context', '')}"
                )
            elif issue.get('category') == 'hyperlink_accessibility':
                # Use the new user-friendly message
                formatted_issues.append(
                    f"    • {issue.get('user_message', issue.get('message', 'No description provided'))}"
                )
            elif 'context' in issue and issue['context'].startswith('Link text:'):
                # This catches the hyperlink issues that might not have the category set
                link_text = issue['context'].replace('Link text:', '').strip().strip('"')
                if any(phrase == link_text.lower() for phrase in ['here', 'click here', 'more', 'link']):
                    formatted_issues.append(
                        f"    • Replace non-descriptive link text \"{link_text}\" with text that clearly indicates where the link will take the user"
                    )
                elif link_text.lower().startswith(('http', 'www', 'ftp')):
                    formatted_issues.append(
                        f"    • Replace the URL \"{link_text}\" with meaningful text that describes the link destination"
                    )
                elif len(link_text) < 4:
                    formatted_issues.append(
                        f"    • Link text \"{link_text}\" is too short - use descriptive text that clearly indicates the link destination"
                    )
                else:
                    formatted_issues.append(f"    • {issue.get('message', 'No description provided')} {issue['context']}")
            else:
                # Generic formatting for other issues
                message = issue.get('message', 'No description provided')
                context = issue.get('context', '').strip()
                formatted_issues.append(
                    f"    • {message} {context}"
                )

        return formatted_issues

    @profile_performance
    def check_hyperlink_usage(self, doc: List[str]) -> DocumentCheckResult:
        """
        Check document for hyperlink usage issues.
        
        Args:
            doc: List of document paragraphs
            
        Returns:
            DocumentCheckResult containing any hyperlink issues found
        """
        if not self.validate_input(doc):
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])

        issues = []
        checked_urls = set()
        
        # URL pattern - matches http/https URLs
        url_pattern = re.compile(
            r'https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b[-a-zA-Z0-9()@:%_\+.~#?&//=]*'
        )
        
        # Helper function to check a single URL
        def check_url(url: str) -> Optional[Dict[str, str]]:
            """
            Check if a URL is accessible and return any issues found.
            
            Args:
                url: The URL to check for accessibility.
                
            Returns:
                Optional[Dict[str, str]]: A dictionary containing the URL and error message if the URL is
                inaccessible, or None if the URL is accessible.
            """
            try:
                response = requests.head(url, timeout=5, allow_redirects=True, headers={'User-Agent': 'CheckerTool/1.0'})
                if response.status_code >= 400:
                    return {
                        'url': url,
                        'message': f"Broken link: {url} (HTTP {response.status_code})"
                    }
            except requests.RequestException:
                return {
                    'url': url,
                    'message': f"Failed to access URL: {url}"
                }
            return None

        # Extract and deduplicate URLs
        for paragraph in doc:
            urls = {match.group() for match in url_pattern.finditer(paragraph)}
            checked_urls.update(urls)
        
        # Concurrently check URLs
        with ThreadPoolExecutor(max_workers=10) as executor:
            future_to_url = {executor.submit(check_url, url): url for url in checked_urls}
            for future in as_completed(future_to_url):
                issue = future.result()
                if issue:
                    issues.append(issue)

        return DocumentCheckResult(
            success=len(issues) == 0,
            issues=issues,
            details={
                'total_urls_checked': len(checked_urls),
                'broken_urls': len(issues)
            }
        )

    def _load_valid_words(self) -> Set[str]:
        """
        Load valid English words from valid_words.txt file.
        
        Returns:
            Set[str]: Set of valid English words in lowercase
        """
        try:
            # Get the directory containing the current file
            current_dir = os.path.dirname(os.path.abspath(__file__))
            words_file = os.path.join(current_dir, 'valid_words.txt')
            
            # Load words from file
            with open(words_file, 'r') as f:
                words = {line.strip().lower() for line in f if line.strip()}
                
            return words
            
        except Exception as e:
            self.logger.warning(f"Error loading word list: {e}")
            return set()  # Return empty set as fallback
    
    @profile_performance
    def check_cross_reference_usage(self, doc_path: str) -> DocumentCheckResult:
        """
        Check document for cross-reference usage issues.
        
        Args:
            doc_path: Path to the document to check
            
        Returns:
            DocumentCheckResult containing any cross-reference issues found
        """
        try:
            doc = Document(doc_path)
        except Exception as e:
            self.logger.error(f"Error reading the document: {e}")
            return DocumentCheckResult(success=False, issues=[{'error': str(e)}], details={})

        heading_structure = self._extract_paragraph_numbering(doc)
        valid_sections = {number for number, _ in heading_structure}
        tables = set()
        figures = set()
        issues = []

        # Skip patterns for external references
        skip_patterns = [
            r'(?:U\.S\.C\.|USC)\s+(?:§+\s*)?(?:Section|section)?\s*\d+',
            r'Section\s+\d+(?:\([a-z]\))*\s+of\s+(?:the\s+)?(?:United States Code|U\.S\.C\.)',
            r'Section\s+\d+(?:\([a-z]\))*\s+of\s+Title\s+\d+',
            r'(?:Section|§)\s*\d+(?:\([a-z]\))*\s+of\s+the\s+Act',
            r'Section\s+\d+\([a-z]\)',
            r'§\s*\d+\([a-z]\)',
            r'\d+\s*(?:CFR|C\.F\.R\.)',
            r'Part\s+\d+(?:\.[0-9]+)*\s+of\s+Title\s+\d+',
            r'Public\s+Law\s+\d+[-–]\d+',
            r'Title\s+\d+,\s+Section\s+\d+(?:\([a-z]\))*',
            r'\d+\s+U\.S\.C\.\s+\d+(?:\([a-z]\))*',
        ]
        skip_regex = re.compile('|'.join(skip_patterns), re.IGNORECASE)

        try:
            # Extract tables and figures
            for para in doc.paragraphs:
                text = para.text.strip() if hasattr(para, 'text') else ''
                
                # Table extraction
                if text.lower().startswith('table'):
                    matches = [
                        re.match(r'^table\s+(\d{1,2}(?:-\d+)?)\b', text, re.IGNORECASE),
                        re.match(r'^table\s+(\d{1,2}(?:\.\d+)?)\b', text, re.IGNORECASE)
                    ]
                    for match in matches:
                        if match:
                            tables.add(match.group(1))

                # Figure extraction
                if text.lower().startswith('figure'):
                    matches = [
                        re.match(r'^figure\s+(\d{1,2}(?:-\d+)?)\b', text, re.IGNORECASE),
                        re.match(r'^figure\s+(\d{1,2}(?:\.\d+)?)\b', text, re.IGNORECASE)
                    ]
                    for match in matches:
                        if match:
                            figures.add(match.group(1))

            # Check references
            for para in doc.paragraphs:
                para_text = para.text.strip() if hasattr(para, 'text') else ''
                if not para_text or skip_regex.search(para_text):
                    continue

                # Table reference check
                table_refs = re.finditer(
                    r'(?:see|in|refer to)?\s*(?:table|Table)\s+(\d{1,2}(?:[-\.]\d+)?)\b', 
                    para_text
                )
                for match in table_refs:
                    ref = match.group(1)
                    if ref not in tables:
                        issues.append({
                            'type': 'Table',
                            'reference': ref,
                            'context': para_text,
                            'message': f"Referenced Table {ref} not found in document"
                        })

                # Figure reference check
                figure_refs = re.finditer(
                    r'(?:see|in|refer to)?\s*(?:figure|Figure)\s+(\d{1,2}(?:[-\.]\d+)?)\b', 
                    para_text
                )
                for match in figure_refs:
                    ref = match.group(1)
                    if ref not in figures:
                        issues.append({
                            'type': 'Figure',
                            'reference': ref,
                            'context': para_text,
                            'message': f"Referenced Figure {ref} not found in document"
                        })

                # Section/paragraph reference check
                section_refs = re.finditer(
                    r'(?:paragraph|section|appendix)\s+([A-Z]?\.?\d+(?:\.\d+)*)',
                    para_text,
                    re.IGNORECASE
                )

                for match in section_refs:
                    ref = match.group(1).strip('.')
                    if not skip_regex.search(para_text):
                        if ref not in valid_sections:
                            found = False
                            for valid_section in valid_sections:
                                if valid_section.strip('.') == ref.strip('.'):
                                    found = True
                                    break
                            
                            if not found:
                                issues.append({
                                    'type': 'Paragraph',
                                    'reference': ref,
                                    'context': para_text,
                                    'message': f"Confirm paragraph {ref} referenced in '{para_text}' exists in the document"
                                })

        except Exception as e:
            self.logger.error(f"Error processing cross references: {str(e)}")
            return DocumentCheckResult(
                success=False,
                issues=[{'type': 'error', 'message': f"Error processing cross references: {str(e)}"}],
                details={}
            )

        return DocumentCheckResult(
            success=len(issues) == 0,
            issues=issues,
            details={
                'total_tables': len(tables),
                'total_figures': len(figures),
                'found_tables': sorted(list(tables)),
                'found_figures': sorted(list(figures)),
                'heading_structure': heading_structure,
                'valid_sections': sorted(list(valid_sections))
            }
        )

    def _extract_paragraph_numbering(self, doc: Document, in_appendix: bool = False) -> List[Tuple[str, str]]:
        """
        Extract paragraph numbers from document headings.
        """
        numbered_paragraphs = []
        
        try:
            # Track heading hierarchy (limit to 6 levels as per standard heading styles)
            current_numbers = {
                1: 0,  # Heading 1: 1, 2, 3, ...
                2: 0,  # Heading 2: 1.1, 1.2, 1.3, ...
                3: 0,  # Heading 3: 1.1.1, 1.1.2, ...
                4: 0,
                5: 0,
                6: 0
            }
            current_parent = {
                2: 0,  # Parent number for level 2
                3: 0,  # Parent number for level 3
                4: 0,
                5: 0,
                6: 0
            }
            last_level = {
                1: 0,  # Last number used at each level
                2: 0,
                3: 0,
                4: 0,
                5: 0,
                6: 0
            }
            
            for para in doc.paragraphs:
                style_name = para.style.name if hasattr(para, 'style') and hasattr(para.style, 'name') else ''
                text = para.text.strip() if hasattr(para, 'text') else ''
                
                # Only process if it's a heading style
                if style_name.startswith('Heading'):
                    try:
                        heading_level = int(style_name.replace('Heading ', ''))
                        
                        # Skip if heading level is beyond our supported range
                        if heading_level > 6:
                            continue
                            
                        if heading_level == 1:
                            # For Heading 1, simply increment
                            current_numbers[1] += 1
                            last_level[1] = current_numbers[1]
                            # Reset all lower levels
                            for level in range(2, 7):  # Changed from 8 to 7
                                current_numbers[level] = 0
                                current_parent[level] = current_numbers[1]
                                last_level[level] = 0
                        else:
                            # Check if we're still in the same parent section
                            parent_changed = current_parent[heading_level] != current_numbers[heading_level - 1]
                            
                            if parent_changed:
                                # Parent section changed
                                current_numbers[heading_level] = 1
                                current_parent[heading_level] = current_numbers[heading_level - 1]
                            else:
                                # Same parent, increment this level
                                current_numbers[heading_level] += 1
                            
                            last_level[heading_level] = current_numbers[heading_level]
                            
                            # Reset all lower levels
                            for level in range(heading_level + 1, 7):  # Changed from 8 to 7
                                current_numbers[level] = 0
                                current_parent[level] = 0
                                last_level[level] = 0
                        
                        # Build section number
                        section_parts = []
                        for level in range(1, heading_level + 1):
                            if level == 1:
                                section_parts.append(str(current_numbers[1]))
                            else:
                                if current_numbers[level] > 0:
                                    section_parts.append(str(current_numbers[level]))
                        
                        section_number = '.'.join(section_parts)
                        
                        if text:
                            numbered_paragraphs.append((section_number, text))
                            
                    except ValueError:
                        continue
            
        except Exception as e:
            self.logger.error(f"Error processing document structure: {str(e)}, Type: {type(e)}, Details: {repr(e)}")
            return []
        
        return numbered_paragraphs

    def _check_heading_sequence(self, current_level: int, previous_level: int) -> Optional[str]:
        """
        Check if heading sequence is valid.
        Returns error message if invalid, None if valid.
        
        Rules:
        - Can go from any level to H1 or H2 (restart numbering)
        - When going deeper, can only go one level at a time (e.g., H1 to H2, H2 to H3)
        - Can freely go to any higher level (e.g., H3 to H1, H4 to H2)
        """
        # When going to a deeper level, only allow one level at a time
        if current_level > previous_level:
            if current_level != previous_level + 1:
                return f"Skipped heading level(s) {previous_level + 1} - Found H{current_level} after H{previous_level}. Add H{previous_level + 1} before this section."
            
        # All other cases are valid:
        # - Going to H1 (restart numbering)
        # - Going to any higher level (e.g., H3 to H1)
        return None

    def _check_heading_structure(self, doc: Document) -> List[Dict[str, str]]:
        """Check document heading structure."""
        issues = []
        previous_level = 0
        previous_heading = ""
        
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                try:
                    current_level = int(para.style.name.replace('Heading ', ''))
                    
                    # Check sequence
                    error = self._check_heading_sequence(current_level, previous_level)
                    if error:
                        issues.append({
                            'category': '508_compliance_heading_structure',
                            'message': error,
                            'context': f"'{para.text}'",
                            'recommendation': f"Ensure heading levels follow a logical sequence."
                        })
                    
                    previous_level = current_level
                    previous_heading = para.text
                    
                except ValueError:
                    continue
        
        return issues
    
    @profile_performance
    def check_readability(self, doc: List[str]) -> DocumentCheckResult:
        """
        Check document readability using multiple metrics and plain language standards.
        
        Args:
            doc (List[str]): List of document paragraphs
            
        Returns:
            DocumentCheckResult: Results including readability scores and identified issues
        """
        if not self.validate_input(doc):
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])

        issues = []
        text_stats = {
            'total_words': 0,
            'total_syllables': 0,
            'total_sentences': 0,
            'complex_words': 0,
            'passive_voice_count': 0
        }
        
        # Patterns for identifying passive voice
        passive_patterns = [
            r'\b(?:am|is|are|was|were|be|been|being)\s+\w+ed\b',
            r'\b(?:am|is|are|was|were|be|been|being)\s+\w+en\b',
            r'\b(?:has|have|had)\s+been\s+\w+ed\b',
            r'\b(?:has|have|had)\s+been\s+\w+en\b'
        ]
        passive_regex = re.compile('|'.join(passive_patterns), re.IGNORECASE)

        def count_syllables(word: str) -> int:
            """
            Count the number of syllables in a word using basic vowel counting rules.
            
            Args:
                word: The word to count syllables in.
                
            Returns:
                int: The number of syllables in the word.
                
            Note:
                This is a simplified syllable counter that may not be accurate for all words.
                It counts vowel groups and handles some common exceptions.
            """
            word = word.lower()
            count = 0
            vowels = 'aeiouy'
            on_vowel = False
            
            for char in word:
                is_vowel = char in vowels
                if is_vowel and not on_vowel:
                    count += 1
                on_vowel = is_vowel
                
            if word.endswith('e'):
                count -= 1
            if word.endswith('le') and len(word) > 2 and word[-3] not in vowels:
                count += 1
            if count == 0:
                count = 1
                
            return count

        # Process each paragraph
        for paragraph in doc:
            if not paragraph.strip():
                continue
                
            # Split into sentences
            sentences = re.split(r'(?<=[.!?])\s+', paragraph.strip())
            text_stats['total_sentences'] += len(sentences)
            
            # Check each sentence
            for sentence in sentences:
                # Count passive voice instances
                if passive_regex.search(sentence):
                    text_stats['passive_voice_count'] += 1
                    
                # Process words
                words = sentence.split()
                text_stats['total_words'] += len(words)
                
                for word in words:
                    word = re.sub(r'[^\w\s]', '', word.lower())
                    if not word:
                        continue
                        
                    syllables = count_syllables(word)
                    text_stats['total_syllables'] += syllables
                    
                    if syllables >= 3:
                        text_stats['complex_words'] += 1

        # Calculate readability metrics
        try:
            # Flesch Reading Ease
            flesch_ease = 206.835 - 1.015 * (text_stats['total_words'] / text_stats['total_sentences']) - 84.6 * (text_stats['total_syllables'] / text_stats['total_words'])
            
            # Flesch-Kincaid Grade Level
            flesch_grade = 0.39 * (text_stats['total_words'] / text_stats['total_sentences']) + 11.8 * (text_stats['total_syllables'] / text_stats['total_words']) - 15.59
            
            # Gunning Fog Index
            fog_index = 0.4 * ((text_stats['total_words'] / text_stats['total_sentences']) + 100 * (text_stats['complex_words'] / text_stats['total_words']))
            
            # Calculate passive voice percentage
            passive_percentage = (text_stats['passive_voice_count'] / text_stats['total_sentences']) * 100 if text_stats['total_sentences'] > 0 else 0
            
            # Add readability summary with high-level guidance and specific issues
            issues = []
            
            if flesch_ease < 50:
                issues.append({
                    'type': 'readability_score',
                    'metric': 'Flesch Reading Ease',
                    'score': round(flesch_ease, 1),
                    'message': 'Document may be too difficult for general audience. Consider simplifying language.'
                })
                
            if flesch_grade > 12:
                issues.append({
                    'type': 'readability_score',
                    'metric': 'Flesch-Kincaid Grade Level',
                    'score': round(flesch_grade, 1),
                    'message': 'Reading level is above 12th grade. Consider simplifying for broader accessibility.'
                })
                
            if fog_index > 12:
                issues.append({
                    'type': 'readability_score',
                    'metric': 'Gunning Fog Index',
                    'score': round(fog_index, 1),
                    'message': 'Text complexity may be too high. Consider using simpler words and shorter sentences.'
                })
                
            if passive_percentage > 10:
                issues.append({
                    'type': 'passive_voice',
                    'percentage': round(passive_percentage, 1),
                    'message': f'Document uses {round(passive_percentage, 1)}% passive voice (target: less than 10%). Consider using more active voice.'
                })
            
            details = {
                'metrics': {
                    'flesch_reading_ease': round(flesch_ease, 1),
                    'flesch_kincaid_grade': round(flesch_grade, 1),
                    'gunning_fog_index': round(fog_index, 1),
                    'passive_voice_percentage': round(passive_percentage, 1)
                }
            }
            
            return DocumentCheckResult(
                success=len(issues) == 0,
                issues=issues,
                details=details
            )
            
        except Exception as e:
            self.logger.error(f"Error calculating readability metrics: {str(e)}")
            return DocumentCheckResult(
                success=False,
                issues=[{'error': f'Error calculating readability metrics: {str(e)}'}]
            )

    @profile_performance
    def check_required_language(self, doc: List[str]) -> DocumentCheckResult:
        """
        Check if the document contains all required language based on its type.
        
        Args:
            doc (List[str]): List of document paragraphs
            
        Returns:
            DocumentCheckResult: Results of required language check
        """
        if not self.validate_input(doc):
            return DocumentCheckResult(success=False, issues=[{'error': 'Invalid document input'}])
            
        issues = []
        doc_type = self._get_doc_type(doc)
        required_patterns = self.required_language_patterns.get(doc_type, [])
        
        if not required_patterns:
            return DocumentCheckResult(success=True, issues=[])
            
        # Combine all paragraphs for searching
        full_text = ' '.join(doc)
        
        # Check for each required pattern
        for pattern in required_patterns:
            if not pattern.search(full_text):
                # Extract the first few words of the pattern for context
                context = pattern.pattern.split('\\')[0].strip()
                issues.append(f"Required language not found: \"{context}...\"")
        
        return DocumentCheckResult(success=len(issues) == 0, issues=issues)

    def _get_doc_type(self, doc: List[str]) -> DocumentType:
        """
        Determine the document type from the first paragraph.
        
        Args:
            doc (List[str]): List of document paragraphs
            
        Returns:
            DocumentType: The determined document type
            
        Raises:
            DocumentTypeError: If document type cannot be determined
        """
        if not doc:
            raise DocumentTypeError("Empty document")
            
        first_paragraph = doc[0].strip().lower()
        
        # Check for document type indicators
        if "advisory circular" in first_paragraph:
            return DocumentType.ADVISORY_CIRCULAR
        elif "policy statement" in first_paragraph:
            return DocumentType.POLICY_STATEMENT
        elif "federal register" in first_paragraph:
            return DocumentType.FEDERAL_REGISTER_NOTICE
        elif "special condition" in first_paragraph:
            return DocumentType.SPECIAL_CONDITION
        elif "airworthiness criteria" in first_paragraph:
            return DocumentType.AIRWORTHINESS_CRITERIA
        elif "deviation memo" in first_paragraph:
            return DocumentType.DEVIATION_MEMO
        elif "exemption" in first_paragraph:
            return DocumentType.EXEMPTION
        elif "order" in first_paragraph:
            return DocumentType.ORDER
        elif "rule" in first_paragraph:
            return DocumentType.RULE
        elif "technical standard order" in first_paragraph:
            return DocumentType.TECHNICAL_STANDARD_ORDER
        else:
            return DocumentType.OTHER

    @profile_performance
    def check_phone_number_format_usage(self, doc: List[str]) -> DocumentCheckResult:
        """
        Check document for phone number format usage issues.
        
        Args:
            doc: List of document paragraphs
            
        Returns:
            DocumentCheckResult containing any phone number format issues found
        """
        issues = []
        phone_formats = set()
        phone_numbers = []
        
        # Compile all phone number patterns
        patterns = self.patterns.get('phone_numbers', [])
        compiled_patterns = [(re.compile(p.pattern), p.format_name, p.description) for p in patterns]
        
        # Log the patterns being used
        self.logger.debug(f"Checking phone numbers with {len(patterns)} patterns")
        
        # Find all phone numbers and their formats
        for paragraph_num, paragraph in enumerate(doc, 1):
            paragraph_has_phone = False
            for pattern, format_name, description in compiled_patterns:
                matches = list(pattern.finditer(paragraph))
                if matches:
                    paragraph_has_phone = True
                    for match in matches:
                        phone_number = match.group(0)
                        phone_numbers.append((phone_number, format_name, description, paragraph_num))
                        phone_formats.add(format_name)
                        self.logger.debug(f"Found phone number '{phone_number}' in format '{format_name}' in paragraph {paragraph_num}")
            
            if paragraph_has_phone:
                self.logger.debug(f"Paragraph {paragraph_num} contains phone numbers")
        
        # Log summary of found phone numbers
        self.logger.debug(f"Found {len(phone_numbers)} phone numbers in {len(phone_formats)} different formats")
        
        # If we found phone numbers and more than one format is used
        if phone_numbers and len(phone_formats) > 1:
            format_counts = {}
            format_examples = {}
            for number, format_name, description, _ in phone_numbers:
                format_counts[format_name] = format_counts.get(format_name, 0) + 1
                if format_name not in format_examples:
                    format_examples[format_name] = number
            
            # Create a detailed message showing the inconsistency
            format_details = []
            for format_name, count in format_counts.items():
                example = format_examples[format_name]
                format_details.append(f"{count} in {format_name} format (e.g., {example})")
            
            # Add paragraph numbers to the details
            paragraph_numbers = sorted(set(p_num for _, _, _, p_num in phone_numbers))
            
            issues.append({
                'type': 'phone_format_inconsistency',
                'message': f"Inconsistent phone number formats found: {', '.join(format_details)}. Please use a single consistent format throughout the document.",
                'details': {
                    'phone_numbers': [num for num, _, _, _ in phone_numbers],
                    'formats_used': list(phone_formats),
                    'format_examples': format_examples,
                    'paragraph_numbers': paragraph_numbers
                }
            })
            
            self.logger.info(f"Found inconsistent phone number formats in paragraphs {paragraph_numbers}")
        
        return DocumentCheckResult(
            success=len(issues) == 0,
            issues=issues
        )

    def _format_phone_format_issues(self, result: DocumentCheckResult) -> List[str]:
        """Format phone number format consistency issues."""
        output = []
        for issue in result.issues:
            if issue['type'] == 'phone_format_inconsistency':
                output.append(f"    • {issue['message']}")
                if 'details' in issue and 'phone_numbers' in issue['details']:
                    output.append("      Found phone numbers:")
                    for phone in issue['details']['phone_numbers']:
                        output.append(f"        • {phone}")
                    output.append("      Please choose one format and use it consistently throughout the document.")
                    if 'format_examples' in issue['details']:
                        output.append("      Example formats found:")
                        for format_name, example in issue['details']['format_examples'].items():
                            output.append(f"        • {format_name}: {example}")
        return output

    def format_results(self, results: Dict[str, Any], doc_type: str) -> str:
        """Format the check results into a readable string."""
        output = []
        
        # Add header
        output.append(f"Document Check Results for {doc_type}")
        output.append("=" * 50)
        output.append("")
        
        # Format each result
        for check_name, result in results.items():
            if not result.success:
                output.append(f"{check_name.replace('_', ' ').title()} Issues:")
                output.append("-" * 30)
                
                if check_name == 'heading_title':
                    output.extend(self._format_heading_issues(result, doc_type))
                elif check_name == 'heading_title_period':
                    output.extend(self._format_period_issues(result))
                elif check_name == 'caption':
                    output.extend(self._format_caption_issues(result.issues, doc_type))
                elif check_name == 'table_figure_reference':
                    output.extend(self._format_reference_issues(result))
                elif check_name == 'terminology':
                    output.extend(self._format_standard_issue(issue) for issue in result.issues)
                elif check_name == 'acronyms':
                    output.extend(self._format_unused_acronym_issues(result))
                elif check_name == 'parentheses':
                    output.extend(self._format_parentheses_issues(result))
                elif check_name == 'section_symbol':
                    output.extend(self._format_section_symbol_issues(result))
                elif check_name == 'spacing':
                    output.extend(self._format_spacing_issues(result))
                elif check_name == 'readability':
                    output.extend(self._format_readability_issues(result))
                elif check_name == 'phone_format_check':
                    output.extend(self._format_phone_format_issues(result))
                else:
                    output.extend(self._format_standard_issue(issue) for issue in result.issues)
                
                output.append("")
        
        return '\n'.join(output)

class DocumentCheckResultsFormatter:
    
    def __init__(self):
        init()  # Initialize colorama
        self.logger = logging.getLogger(__name__)
        
        # Enhanced issue categories with examples and specific fixes
        self.issue_categories = {
            'heading_title_check': {
                'title': 'Required Headings Check',
                'description': 'Verifies that your document includes all mandatory section headings, with requirements varying by document type. For example, long-template Advisory Circulars require headings like "Purpose." and "Applicability." with initial caps and periods, while Federal Register Notices use ALL CAPS headings like "SUMMARY" and "BACKGROUND" without periods. This check ensures both the presence of required headings and their correct capitalization format based on document type.',
                'solution': 'Add all required headings in the correct order using the correct capitalization format.',
                'example_fix': {
                    'before': 'Missing required heading "PURPOSE."',
                    'after': 'Added heading "PURPOSE." at the beginning of the document'
                }
            },
            'heading_title_period_check': {
                'title': 'Heading Period Format',
                'description': 'Examines heading punctuation to ensure compliance with FAA document formatting standards. Some FAA documents (like Advisory Circulars and Orders) require periods at the end of headings, while others (like Federal Register Notices) don\'t.',
                'solution': 'Format heading periods according to document type requirements.',
                'example_fix': {
                    'before': 'Purpose',
                    'after': 'Purpose.' # For ACs and Orders
                }
            },
            'table_figure_reference_check': {
                'title': 'Table and Figure References',
                'description': 'Analyzes how tables and figures are referenced within your document text. Capitalize references at the beginning of sentences (e.g., "Table 2-1 shows...") and use lowercase references within sentences (e.g., "...as shown in table 2-1").',
                'solution': 'Capitalize references at start of sentences, use lowercase within sentences.',
                'example_fix': {
                    'before': 'The DTR values are specified in Table 3-1 and Figure 3-2.',
                    'after': 'The DTR values are specified in table 3-1 and figure 3-2.'
                }
            },
            'acronym_check': {
                'title': 'Acronym Definition Issues',
                'description': 'Ensures every acronym is properly introduced with its full term at first use. The check identifies undefined acronyms while recognizing common exceptions (like U.S.) that don\'t require definition.',
                'solution': 'Define each acronym at its first use, e.g., "Federal Aviation Administration (FAA)".',
                'example_fix': {
                    'before': 'This order establishes general FAA organizational policies.',
                    'after': 'This order establishes general Federal Aviation Administration (FAA) organizational policies.'
                }
            },
            'acronym_usage_check': {
                'title': 'Unused Acronym Definitions',
                'description': 'Ensures that all acronyms defined in the document are actually used later. If an acronym is defined but never referenced, the definition should be removed to avoid confusion or unnecessary clutter.',
                'solution': 'Identify acronyms that are defined but not used later in the document and remove their definitions.',
                'example_fix': {
                    'before': 'Operators must comply with airworthiness directives (AD) to ensure aircraft safety and regulatory compliance.',
                    'after': 'Operators must comply with airworthiness directives to ensure aircraft safety and regulatory compliance.'
                }
            },
            'terminology_check': {
                'title': 'Incorrect Terminology',
                'description': 'Evaluates document text against the various style manuals and orders to identify non-compliant terminology, ambiguous references, and outdated phrases. This includes checking for prohibited relative references (like "above" or "below"), proper legal terminology (like "must" instead of "shall"), and consistent formatting of regulatory citations.',
                'solution': 'Use explicit references to paragraphs, sections, tables, and figures.',
                'example_fix': {
                    'before': 'Operators shall comply with ADs to ensure aircraft safety and regulatory compliance',
                    'after': 'Operators must comply with ADs to ensure aircraft safety and regulatory compliance.'
                }
            },
            'section_symbol_usage_check': {
                'title': 'Section Symbol (§) Format Issues',
                'description': 'Examines the usage of section symbols (§) throughout your document. This includes verifying proper symbol placement in regulatory references, ensuring sections aren\'t started with the symbol, checking consistency in multiple-section citations, and validating proper CFR citations. For ACs, see FAA Order 1320.46.',
                'solution': 'Format section symbols correctly and never start sentences with them.',
                'example_fix': {
                    'before': '§ 23.3 establishes design criteria.',
                    'after': 'Section 23.3 establishes design criteria.'
                }
            },
            'double_period_check': {
                'title': 'Multiple Period Issues',
                'description': 'Examines sentences for accidental double periods that often occur during document editing and revision. While double periods are sometimes found in ellipses (...) or web addresses, they should never appear at the end of standard sentences in FAA documentation.',
                'solution': 'Remove multiple periods that end sentences.',
                'example_fix': {
                    'before': 'The following ACs are related to the guidance in this document..',
                    'after': 'The following ACs are related to the guidance in this document.'
                }
            },
            'spacing_check': {
                'title': 'Spacing Issues',
                'description': 'Analyzes document spacing patterns to ensure compliance with FAA formatting standards. This includes checking for proper spacing around regulatory references (like "AC 25-1" not "AC25-1"), section symbols (§ 25.1), paragraph references, and multiple spaces between words.',
                'solution': 'Fix spacing issues: remove any missing spaces, double spaces, or inadvertent tabs.',
                'example_fix': {
                    'before': 'AC25.25 states that  SFAR88 and §25.981 require...',
                    'after': 'AC 25.25 states that SFAR 88 and § 25.981 require...'
                }
            },
            'date_formats_check': {
                'title': 'Date Format Issues',
                'description': 'Examines all date references in your document. The check automatically excludes technical reference numbers that may look like dates to ensure accurate validation of true date references. Note, though, there might be instances in the heading of the document where the date is formatted as "MM/DD/YYYY", which is acceptable. This applies mostly to date formats within the document body.',
                'solution': 'Use the format "Month Day, Year" where appropriate.',
                'example_fix': {
                    'before': 'This policy statement cancels Policy Statement PS-AIR100-2006-MMPDS, dated 7/25/2006.',
                    'after': 'This policy statement cancels Policy Statement PS-AIR100-2006-MMPDS, dated July 25, 2006.'
                }
            },
            'placeholders_check': {
                'title': 'Placeholder Content',
                'description': 'Identifies incomplete content and temporary placeholders that must be finalized before document publication. This includes common placeholder text (like "TBD" or "To be determined"), draft markers, and incomplete sections.',
                'solution': 'Replace all placeholder content with actual content.',
                'example_fix': {
                    'before': 'Pilots must submit the [Insert text] form to the FAA for approval.',
                    'after': 'Pilots must submit the Report of Eye Evaluation form 8500-7 to the FAA for approval.'
                }
            },
            'parentheses_check': {
                'title': 'Parentheses Balance Check',
                'description': 'Ensures that all parentheses in the document are properly paired with matching opening and closing characters.',
                'solution': 'Add missing opening or closing parentheses where indicated.',
                'example_fix': {
                    'before': 'The system (as defined in AC 25-11B performs...',
                    'after': 'The system (as defined in AC 25-11B) performs...'
                }
            },
            'paragraph_length_check': {
                'title': 'Paragraph Length Issues',
                'description': 'Flags paragraphs exceeding 6 sentences or 8 lines to enhance readability and clarity. While concise paragraphs are encouraged, with each focusing on a single idea or related points, exceeding these limits doesn\'t necessarily indicate a problem. Some content may appropriately extend beyond 8 lines, especially if it includes necessary details. Boilerplate language or template text exceeding these limits is not subject to modification or division.',
                'solution': 'Where possible, split long paragraphs into smaller sections, ensuring each focuses on one primary idea. If restructuring is not feasible or the content is boilerplate text, no changes are needed.',
                'example_fix': {
                    'before': 'A very long paragraph covering multiple topics and spanning many lines...',
                    'after': 'Multiple shorter paragraphs or restructured paragraphs, each focused on a single topic or related points.'
                }
            },
            'sentence_length_check': {
                'title': 'Sentence Length Issues',
                'description': 'Analyzes sentence length to ensure readability. While the ideal length varies with content complexity, sentences over 35 words often become difficult to follow. Technical content, regulatory references, notes, warnings, and list items are excluded from this check.',
                'solution': 'Break long sentences into smaller ones where possible, focusing on one main point per sentence. Consider using lists for complex items.',
                'example_fix': {
                    'before': 'The operator must ensure that all required maintenance procedures are performed in accordance with the manufacturer\'s specifications and that proper documentation is maintained throughout the entire process to demonstrate compliance with regulatory requirements.',
                    'after': 'The operator must ensure all required maintenance procedures are performed according to manufacturer specifications. Additionally, proper documentation must be maintained to demonstrate regulatory compliance.'
                }
            },
            'document_title_check': {
                'title': 'Referenced Document Title Format Issues',
                'description': 'Checks document title formatting based on document type. Advisory Circulars require italics without quotes, while all other document types require quotes without italics.',
                'solution': 'Format document titles according to document type: use italics for Advisory Circulars, quotes for all other document types.',
                'example_fix': {
                    'before': 'See AC 25.1309-1B, System Design and Analysis, for information on X.',
                    'after': 'See AC 25.1309-1B, <i>System Design and Analysis</i>, for information on X.'
                }
            },
            '508_compliance_check': {
                'title': 'Section 508 Compliance Issues',
                'description': 'Checks document accessibility features required by Section 508 standards: Image alt text for screen readers, heading structure issues (missing heading 1, skipped heading levels, and out of sequence headings), and hyperlink accessibility (ensuring links have meaningful descriptive text).',
                'solution': 'Address each accessibility issue: add image alt text for screen readers, fix heading structure, and ensure hyperlinks have descriptive text that indicates their destination.',
                'example_fix': {
                    'before': [
                        'Image without alt text',
                        'Heading sequence: H1 → H2 → H4 (skipped H3)',
                        'Link text: "click here" or "www.example.com"'
                    ],
                    'after': [
                        'Image with descriptive alt text',
                        'Proper heading sequence: H1 → H2 → H3 → H4',
                        'Descriptive link text: "FAA Compliance Guidelines" or "Download the Safety Report"'
                    ]
                }
            },
            'hyperlink_check': {
                'title': 'Hyperlink Issues',
                'description': 'Checks for potentially broken or inaccessible URLs in the document. This includes checking response codes and connection issues.',
                'solution': 'Verify each flagged URL is correct and accessible.',
                'example_fix': {
                    'before': 'See https://broken-link.example.com for more details.',
                    'after': 'See https://www.faa.gov for more details.'
                }
            },
            'cross_references_check': {
                'title': 'Cross-Reference Issues',
                'description': 'Checks for missing or invalid cross-references to paragraphs, tables, figures, and appendices within the document.',
                'solution': 'Ensure that all referenced elements are present in the document and update or remove any incorrect references.',
                'example_fix': {
                    'before': 'See table 5-2 for more information. (there is no table 5-2)',
                    'after': 'Either update the table reference or add table 5-2 if missing'
                }
            },
            'readability_check': {
                'title': 'Readability Issues',
                'description': 'Analyzes document readability using multiple metrics including Flesch Reading Ease, Flesch-Kincaid Grade Level, and Gunning Fog Index. Also checks for passive voice usage and technical jargon.',
                'solution': 'Simplify language, reduce passive voice, and replace technical jargon with plain language alternatives.',
                'example_fix': {
                    'before': 'The implementation of the procedure was facilitated by technical personnel.',
                    'after': 'Technical staff helped start the procedure.'
                }
            },
            'phone_format_check': {
                'title': 'Phone Number Format',
                'description': 'Checks for consistent phone number formatting throughout the document. The checker looks for phone numbers in various formats and flags inconsistencies.',
                'solution': 'Choose one phone number format and use it consistently throughout the document. Common formats include (xxx) xxx-xxxx, xxx-xxx-xxxx, xxx.xxx.xxxx, or xxxxxxxxxx.',
                'example_fix': {
                    'before': 'Dial (603) 333-3456 and then dial 606-354-2456',
                    'after': 'Dial (603) 333-3456 and then dial (606) 354-2456'
                },
            }
        }

        # Add these two helper methods here, after __init__ and before other methods
    def _format_colored_text(self, text: str, color: str) -> str:
        """Helper method to format colored text with reset.
        
        Args:
            text: The text to be colored
            color: The color to apply (from colorama.Fore)
            
        Returns:
            str: The colored text with reset styling
        """
        return f"{color}{text}{Style.RESET_ALL}"
    
    def _format_example(self, example_fix: Dict[str, str]) -> List[str]:
        """Format example fixes consistently.
        
        Args:
            example_fix: Dictionary containing 'before' and 'after' examples
            
        Returns:
            List[str]: Formatted example lines
        """
        return [
            f"    ❌ Incorrect: {example_fix['before']}",
            f"    ✓ Correct: {example_fix['after']}"
        ]
    
    def _format_heading_issues(self, result: DocumentCheckResult, doc_type: str) -> List[str]:
        """Format heading check issues consistently."""
        output = []
        
        for issue in result.issues:
            if issue.get('type') == 'missing_headings':
                missing = sorted(issue['missing'])
                output.append(f"\n  Missing Required Headings for {doc_type}:")
                for heading in missing:
                    output.append(f"    • {heading}")
            elif issue.get('type') == 'unexpected_headings':
                unexpected = sorted(issue['unexpected'])
                output.append(f"\n  Unexpected Headings Found:")
                for heading in unexpected:
                    output.append(f"    • {heading}")
        
        return output

    def _format_period_issues(self, result: DocumentCheckResult) -> List[str]:
        """Format period check issues consistently."""
        output = []
        
        if result.issues:
            output.append(f"\n  Heading Period Format Issues:")
            for issue in result.issues:
                if 'message' in issue:
                    output.append(f"    • {issue['message']}")
        
        return output

    def _format_caption_issues(self, issues: List[Dict], doc_type: str) -> List[str]:
        """Format caption check issues with clear replacement instructions."""
        formatted_issues = []
        for issue in issues:
            if 'incorrect_caption' in issue:
                caption_parts = issue['incorrect_caption'].split()
                if len(caption_parts) >= 2:
                    caption_type = caption_parts[0]  # "Table" or "Figure"
                    number = caption_parts[1]
                    
                    # Determine correct format based on document type
                    if doc_type in ["Advisory Circular", "Order"]:
                        if '-' not in number:
                            correct_format = f"{caption_type} {number}-1"
                    else:
                        if '-' in number:
                            correct_format = f"{caption_type} {number.split('-')[0]}"
                        else:
                            correct_format = issue['incorrect_caption']

                    formatted_issues.append(
                        f"    • Replace '{issue['incorrect_caption']}' with '{correct_format}'"
                    )

        return formatted_issues

    def _format_reference_issues(self, result: DocumentCheckResult) -> List[str]:
        """Format reference issues with clear, concise descriptions."""
        formatted_issues = []
        
        for issue in result.issues:
            ref_type = issue.get('type', '')
            ref_num = issue.get('reference', '')
            context = issue.get('context', '').strip()
            
            if context:  # Only include context if it exists
                formatted_issues.append(
                    f"    • Confirm {ref_type} {ref_num} referenced in '{context}' exists in the document"
                )
            else:
                formatted_issues.append(
                    f"    • Confirm {ref_type} {ref_num} exists in the document"
                )

        return formatted_issues

    def _format_standard_issue(self, issue: Dict[str, Any]) -> str:
        """Format standard issues consistently."""
        if isinstance(issue, str):
            return f"    • {issue}"
        
        if 'incorrect' in issue and 'correct' in issue:
            return f"    • Replace '{issue['incorrect']}' with '{issue['correct']}'"
        
        if 'incorrect_term' in issue and 'correct_term' in issue:
            return f"    • Replace '{issue['incorrect_term']}' with '{issue['correct_term']}'"
        
        if 'sentence' in issue and 'word_count' in issue:  # For sentence length check
            return f"    • Review this sentence: \"{issue['sentence']}\""
        
        if 'sentence' in issue:
            return f"    • {issue['sentence']}"
        
        if 'description' in issue:
            return f"    • {issue['description']}"
        
        if 'type' in issue and issue['type'] == 'long_paragraph':
            return f"    • {issue['message']}"
        
        # Fallback for other issue formats
        return f"    • {str(issue)}"

    def _format_unused_acronym_issues(self, result: DocumentCheckResult) -> List[str]:
        """Format unused acronym issues with a simple, clear message.
        
        Args:
            result: DocumentCheckResult containing acronym issues
            
        Returns:
            List[str]: Formatted list of unused acronym issues
        """
        formatted_issues = []
        
        if result.issues:
            for issue in result.issues:
                if isinstance(issue, dict) and 'acronym' in issue:
                    formatted_issues.append(f"    • Acronym '{issue['acronym']}' was defined but never used.")
                elif isinstance(issue, str):
                    # Handle case where issue might be just the acronym
                    formatted_issues.append(f"    • Acronym '{issue}' was defined but never used.")
    
        return formatted_issues

    def _format_parentheses_issues(self, result: DocumentCheckResult) -> List[str]:
        """Format parentheses issues with clear instructions for fixing."""
        formatted_issues = []
        
        if result.issues:
            for issue in result.issues:
                formatted_issues.append(f"    • {issue['message']}")
        
        return formatted_issues

    def _format_section_symbol_issues(self, result: DocumentCheckResult) -> List[str]:
        """Format section symbol issues with clear replacement instructions."""
        formatted_issues = []
        
        if result.issues:
            for issue in result.issues:
                if 'incorrect' in issue and 'correct' in issue:
                    if issue.get('is_sentence_start'):
                        formatted_issues.append(
                            f"    • Do not begin sentences with the section symbol. "
                            f"Replace '{issue['incorrect']}' with '{issue['correct']}' at the start of the sentence"
                        )
                    else:
                        formatted_issues.append(
                            f"    • Replace '{issue['incorrect']}' with '{issue['correct']}'"
                        )
        
        return formatted_issues

    def _format_paragraph_length_issues(self, result: DocumentCheckResult) -> List[str]:
        """Format paragraph length issues with clear instructions for fixing.
        
        Args:
            result: DocumentCheckResult containing paragraph length issues
            
        Returns:
            List[str]: Formatted list of paragraph length issues
        """
        formatted_issues = []
        
        if result.issues:
            for issue in result.issues:
                if isinstance(issue, str):
                    formatted_issues.append(f"    • {issue}")
                elif isinstance(issue, dict) and 'message' in issue:
                    formatted_issues.append(f"    • {issue['message']}")
                else:
                    # Fallback for unexpected issue format
                    formatted_issues.append(f"    • Review paragraph for length issues: {str(issue)}")
        
        return formatted_issues

    def format_results(self, results: Dict[str, Any], doc_type: str) -> str:
        """
        Format check results into a detailed, user-friendly report.
        
        Args:
            results: Dictionary of check results
            doc_type: Type of document being checked
            
        Returns:
            str: Formatted report with consistent styling
        """
        # Determine caption format based on document type
        if doc_type in ["Advisory Circular", "Order"]:
            table_format = {
                'title': 'Table Caption Format Issues',
                'description': 'Analyzes table captions to ensure they follow the FAA\'s dual-numbering system, where tables must be numbered according to their chapter or appendix location (X-Y format). The first number (X) indicates the chapter number, while the second number (Y) provides sequential numbering within that chapter. For more information, see FAA Order 1320.46.',
                'solution': 'Use the format "Table X-Y" where X is the chapter or appendix number and Y is the sequence number',
                'example_fix': {
                    'before': 'Table 5. | Table A.',
                    'after': 'Table 5-1. | Table A-1.'
                }
            }
            figure_format = {
                'title': 'Figure Caption Format Issues',
                'description': 'Analyzes figure captions to ensure they follow the FAA\'s dual-numbering system, where figures must be numbered according to their chapter or appendix location (X-Y format). The first number (X) indicates the chapter number, while the second number (Y) provides sequential numbering within that chapter. For more information, see FAA Order 1320.46.',
                'solution': 'Use the format "Figure X-Y" where X is the chapter or appendix number and Y is the sequence number',
                'example_fix': {
                    'before': 'Figure 5. | Figure A.',
                    'after': 'Figure 5-1. | Figure A-1.'
                }
            }
        else:
            table_format = {
                'title': 'Table Caption Format Issues',
                'description': f'Analyzes table captions to ensure they follow the FAA\'s sequential numbering system for {doc_type}s. Tables must be numbered consecutively throughout the document using a single number format.',
                'solution': 'Use the format "Table X" where X is a sequential number',
                'example_fix': {
                    'before': 'Table 5-1. | Table A-1',
                    'after': 'Table 5. | Table 1.'
                }
            }
            figure_format = {
                'title': 'Figure Caption Format Issues',
                'description': f'Analyzes figure captions to ensure they follow the FAA\'s sequential numbering system for {doc_type}s. Figures must be numbered consecutively throughout the document using a single number format.',
                'solution': 'Use the format "Figure X" where X is a sequential number',
                'example_fix': {
                    'before': 'Figure 5-1. | Figure A-1.',
                    'after': 'Figure 5. | Figure 3.'
                }
            }

        # Update the issue categories with the correct format
        self.issue_categories['caption_check_table'] = table_format
        self.issue_categories['caption_check_figure'] = figure_format

        # Define formatting rules for different document types
        formatting_rules = {
            "italics_only": {
                "types": ["Advisory Circular"],
                "italics": True, 
                "quotes": False,
                "description": "For Advisory Circulars, referenced document titles should be italicized but not quoted. For more information, see FAA Order 1320.46.",
                "example": "See AC 25.1309-1B, <i>System Design and Analysis</i>, for information on X."
            },
            "quotes_only": {
                "types": [
                    "Airworthiness Criteria", "Deviation Memo", "Exemption", 
                    "Federal Register Notice", "Order", "Policy Statement", "Rule", 
                    "Special Condition", "Technical Standard Order", "Other"
                ],
                "italics": False, 
                "quotes": True,
                "description": "For this document type, referenced document titles should be in quotes without italics.",
                "example": 'See AC 25.1309-1B, "System Design and Analysis," for information on X.'
            }
        }

        # Find the formatting group for the current document type
        format_group = None
        for group, rules in formatting_rules.items():
            if doc_type in rules["types"]:
                format_group = rules
                break

        # Use quotes_only as default if document type not found
        if not format_group:
            format_group = formatting_rules["quotes_only"]

        # Update document title check category based on document type
        if doc_type == "Advisory Circular":
            self.issue_categories['document_title_check'] = {
                'title': 'Referenced Document Title Format Issues',
                'description': 'For Advisory Circulars, all referenced document titles must be italicized.',
                'solution': 'Format document titles using italics for Advisory Circulars.',
                'example_fix': {
                    'before': 'See AC 25.1309-1B, System Design and Analysis, for information on X.',
                    'after': 'See AC 25.1309-1B, <i>System Design and Analysis</i>, for information on X.'
                }
            }
        else:
            self.issue_categories['document_title_check'] = {
                'title': 'Referenced Document Title Format Issues',
                'description': f'For {doc_type}s, all referenced document titles must be enclosed in quotation marks.',
                'solution': 'Format document titles using quotation marks.',
                'example_fix': {
                    'before': 'See AC 25.1309-1B, System Design and Analysis, for information on X.',
                    'after': 'See AC 25.1309-1B, "System Design and Analysis," for information on X.'
                }
            }
        
        output = []
        
        self.issue_categories['acronym_usage_check'] = {
            'title': 'Unused Acronym Definitions',
            'description': 'Ensures that all acronyms defined in the document are actually used later. If an acronym is defined but never referenced, the definition should be removed to avoid confusion or unnecessary clutter.',
            'solution': 'Identify acronyms that are defined but not used later in the document and remove their definitions.',
            'example_fix': {
                'before': 'Operators must comply with airworthiness directives (AD) to ensure aircraft safety and regulatory compliance.',
                'after': 'Operators must comply with airworthiness directives to ensure aircraft safety and regulatory compliance.'
            }
        }

        output = []
        
        # Header
        output.append(f"\n{Fore.CYAN}{'='*80}")
        output.append(f"Document Check Results Summary")
        output.append(f"{'='*80}{Style.RESET_ALL}\n")
        
        # Count total issues
        total_issues = sum(1 for r in results.values() if not r.success)
        if total_issues == 0:
            output.append(f"{self._format_colored_text('✓ All checks passed successfully!', Fore.GREEN)}\n")
            return '\n'.join(output)
        
        output.append(f"{Fore.YELLOW}Found {total_issues} categories of issues that need attention:{Style.RESET_ALL}\n")
        
        # Process all check results consistently
        for check_name, result in results.items():
            if not result.success and check_name in self.issue_categories:
                category = self.issue_categories[check_name]
                
                output.append("\n")
                output.append(self._format_colored_text(f"■ {category['title']}", Fore.YELLOW))
                output.append(f"  {category['description']}")
                output.append(f"  {self._format_colored_text('How to fix: ' + category['solution'], Fore.GREEN)}")
                        
                output.append(f"\n  {self._format_colored_text('Example Fix:', Fore.CYAN)}")
                output.extend(self._format_example(category['example_fix']))
                output.append("")
                
                output.append(f"  {self._format_colored_text('Issues found in your document:', Fore.CYAN)}")
                
                # Special handling for date format issues
                if check_name == 'date_formats_check':
                    for issue in result.issues:
                        output.append(f"    • Replace '{issue['incorrect']}' with '{issue['correct']}'")
                # Handle other check types
                elif check_name == 'heading_title_check':
                    output.extend(self._format_heading_issues(result, doc_type))
                elif check_name == 'heading_title_period_check':
                    output.extend(self._format_period_issues(result))
                elif check_name == 'table_figure_reference_check':
                    output.extend(self._format_reference_issues(result))
                elif check_name in ['caption_check_table', 'caption_check_figure']:
                    output.extend(self._format_caption_issues(result.issues, doc_type))
                elif check_name == 'acronym_usage_check':
                    output.extend(self._format_unused_acronym_issues(result))
                elif check_name == 'section_symbol_usage_check':
                    output.extend(self._format_section_symbol_issues(result))
                elif check_name == 'parentheses_check':
                    output.extend(self._format_parentheses_issues(result))
                elif check_name == '508_compliance_check':
                    if not result.success:
                        # Combine all 508 compliance issues into a single list
                        for issue in result.issues:
                            if issue.get('category') == '508_compliance_heading_structure':
                                output.append(f"    • {issue['message']}")
                                if 'context' in issue:
                                    output.append(f"      Context: {issue['context']}")
                                if 'recommendation' in issue:
                                    output.append(f"      Recommendation: {issue['recommendation']}")
                            elif issue.get('category') == 'image_alt_text':
                                if 'context' in issue:
                                    output.append(f"    • {issue['context']}")
                            elif issue.get('category') == 'hyperlink_accessibility':
                                output.append(f"    • {issue.get('user_message', issue.get('message', 'No description provided'))}")
                elif check_name == 'hyperlink_check':
                    for issue in result.issues:
                        output.append(f"    • {issue['message']}")
                        if 'status_code' in issue:
                            output.append(f"      (HTTP Status: {issue['status_code']})")
                        elif 'error' in issue:
                            output.append(f"      (Error: {issue['error']})")
                elif check_name == 'cross_references_check':
                    for issue in result.issues:
                        output.append(f"    • Confirm {issue['type']} {issue['reference']} referenced in '{issue['context']}' exists in the document")
                elif check_name == 'readability_check':
                    output.extend(self._format_readability_issues(result))
                else:
                    formatted_issues = [self._format_standard_issue(issue) for issue in result.issues[:15]]
                    output.extend(formatted_issues)
                    
                    if len(result.issues) > 50:
                        output.append(f"\n    ... and {len(result.issues) - 50} more similar issues.")
        
        return '\n'.join(output)

    def save_report(self, results: Dict[str, Any], filepath: str, doc_type: str) -> None:
        """Save the formatted results to a file with proper formatting."""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                # Create a report without color codes
                report = self.format_results(results, doc_type)
                
                # Strip color codes
                for color in [Fore.CYAN, Fore.GREEN, Fore.YELLOW, Fore.RED, Style.RESET_ALL]:
                    report = report.replace(str(color), '')
                
                # Convert markdown-style italics to alternative formatting for plain text
                report = report.replace('*', '_')
                
                f.write(report)
        except Exception as e:
            self.logger.error(f"Error saving report: {e}")

    def _format_readability_issues(self, result: DocumentCheckResult) -> List[str]:
        """Format readability issues with clear, actionable feedback."""
        formatted_issues = []
        
        if result.details and 'metrics' in result.details:
            metrics = result.details['metrics']
            formatted_issues.append("\n  Readability Scores:")
            formatted_issues.append(f"    • Flesch Reading Ease: {metrics['flesch_reading_ease']} (Aim for 50+; higher is easier to read)")
            formatted_issues.append(f"    • Grade Level: {metrics['flesch_kincaid_grade']} (Aim for 10 or lower; 12 acceptable for technical/legal)")
            formatted_issues.append(f"    • Gunning Fog Index: {metrics['gunning_fog_index']} (Aim for 12 or lower)")
            formatted_issues.append(f"    • Passive Voice: {metrics['passive_voice_percentage']}% (Aim for less than 10%; use active voice for clarity)")
        
        if result.issues:
            formatted_issues.append("\n  Identified Issues:")
            for issue in result.issues:
                if issue['type'] == 'jargon':
                    formatted_issues.append(
                        f"    • Replace '{issue['word']}' with '{issue['suggestion']}' in: \"{issue['sentence']}\""
                    )
                elif issue['type'] in ['readability_score', 'passive_voice']:
                    formatted_issues.append(f"    • {issue['message']}")
        
        return formatted_issues
    
def format_markdown_results(results: Dict[str, DocumentCheckResult], doc_type: str) -> str:
    """Format check results into a Markdown string for Gradio display."""
    output = []
    
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    output.extend([
        f"# Document Check Results - {current_time}",
        f"## Document Type: {doc_type}",
        "---\n"
    ])

    total_issues = sum(1 for r in results.values() if not r.success)
    
    if total_issues == 0:
        output.append("✅ **All checks passed successfully!**\n")
        return "\n".join(output)
    
    output.append(f"❗ Found issues in {total_issues} check categories\n")

    check_categories = {
        'heading_title_check': {'title': '📋 Required Headings', 'priority': 1},
        'heading_title_period_check': {'title': '🔍 Heading Period Usage', 'priority': 1},
        'terminology_check': {'title': '📖 Terminology Usage', 'priority': 1},
        'acronym_check': {'title': '📝 Acronym Definitions', 'priority': 1},
        'acronym_usage_check': {'title': '📎 Acronym Usage', 'priority': 1},
        'section_symbol_usage_check': {'title': '§ Section Symbol Usage', 'priority': 2},
        '508_compliance_check': {'title': '🕵️‍♂️ 508 Compliance', 'priority': 2},
        'cross_references_check': {'title': '🔗 Cross References', 'priority': 2},
        'hyperlink_check': {'title': '🔗 Hyperlinks', 'priority': 2},
        'date_formats_check': {'title': '📅 Date Formats', 'priority': 2},
        'placeholders_check': {'title': '🚩 Placeholder Content', 'priority': 2},
        'document_title_check': {'title': '📑 Document Title Format', 'priority': 2},
        'caption_check_table': {'title': '📊 Table Captions', 'priority': 3},
        'caption_check_figure': {'title': '🖼️ Figure Captions', 'priority': 3},
        'table_figure_reference_check': {'title': '🔗 Table/Figure References', 'priority': 3},
        'parentheses_check': {'title': '🔗 Parentheses Usage', 'priority': 4},
        'double_period_check': {'title': '⚡ Double Periods', 'priority': 4},
        'spacing_check': {'title': '⌨️ Spacing Issues', 'priority': 4},
        'phone_format_check': {'title': '🔗 Phone Number Format', 'priority': 4},
        'paragraph_length_check': {'title': '📏 Paragraph Length', 'priority': 5},
        'sentence_length_check': {'title': '📏 Sentence Length', 'priority': 5},
        
    }

    sorted_checks = sorted(
        [(name, result) for name, result in results.items()],
        key=lambda x: check_categories.get(x[0], {'priority': 999})['priority']
    )

    for check_name, result in sorted_checks:
        if not result.success:
            category = check_categories.get(check_name, {'title': check_name.replace('_', ' ').title()})
            
            output.append(f"### {category['title']}")
            
            if isinstance(result.issues, list):
                for issue in result.issues[:30]:
                    if isinstance(issue, dict):
                        for key, value in issue.items():
                            if isinstance(value, list):
                                output.extend([f"- {item}" for item in value])
                            else:
                                output.append(f"- {key}: {value}")
                    else:
                        output.append(f"- {issue}")
                
                if len(result.issues) > 50:
                    output.append(f"\n*...and {len(result.issues) - 50} more similar issues*")
            
            output.append("")

    return "\n".join(output)

def create_interface():
    """Create and configure the Gradio interface."""
    
    document_types = [
        "Advisory Circular",
        "Airworthiness Criteria",
        "Deviation Memo",
        "Exemption",
        "Federal Register Notice",
        "Order",
        "Policy Statement",
        "Rule",
        "Special Condition",
        "Technical Standard Order",
        "Other"
    ]
    
    template_types = ["Short AC template AC", "Long AC template AC"]

    def format_results_to_html(text_results):
        """Format text results as HTML."""
        html_output = []
        html_output.append("<html><head><style>")
        html_output.append("""
            body { font-family: Arial, sans-serif; margin: 20px; }
            .issue { margin: 10px 0; padding: 10px; border-left: 4px solid #ff9800; background-color: #fff3e0; }
            .success { color: #4caf50; }
            .error { color: #f44336; }
            .warning { color: #ff9800; }
            .info { color: #2196f3; }
        """)
        html_output.append("</style></head><body>")
        html_output.append(text_results.replace("\n", "<br>"))
        html_output.append("</body></html>")
        return "".join(html_output)

    def format_results_as_html(text_results):
        """Convert the text results into styled HTML."""
        if not text_results:
            return """
                <div class="p-4 text-gray-600">
                    Results will appear here after processing...
                </div>
            """
        
        # Remove ANSI color codes
        ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
        text_results = ansi_escape.sub('', text_results)
        
        # Split into sections while preserving the header
        sections = text_results.split('■')
        header = sections[0].strip()
        issues = sections[1:]
        
        # Extract the number of issues from the header text
        issues_count_match = re.search(r'Found (\d+) categories', header)
        issues_count = issues_count_match.group(1) if issues_count_match else len(issues)
        
        # Format header with title
        header_html = f"""
            <div class="max-w-4xl mx-auto p-4 bg-white rounded-lg shadow-sm mb-6">
                <h1 class="text-2xl font-bold mb-4">Document Check Results Summary</h1>
                <div class="text-lg text-amber-600">
                    Found {issues_count} categories of issues that need attention.
                </div>
            </div>
        """
        
        # Format each issue section
        issues_html = ""
        for section in issues:
            if not section.strip():
                continue
                
            parts = section.strip().split('\n', 1)
            if len(parts) < 2:
                continue
                
            title = parts[0].strip()
            content = parts[1].strip()
            
            # Special handling for readability metrics
            if "Readability Issues" in title:
                metrics_match = re.search(r'Readability Scores:(.*?)(?=Identified Issues:|$)', content, re.DOTALL)
                issues_match = re.search(r'Identified Issues:(.*?)(?=\Z)', content, re.DOTALL)
                
                metrics_html = ""
                if metrics_match:
                    metrics = metrics_match.group(1).strip().split('\n')
                    metrics_html = """
                        <div class="bg-blue-50 rounded-lg p-4 mb-4">
                            <h3 class="font-medium text-blue-800 mb-2">📊 Readability Metrics</h3>
                            <div class="grid grid-cols-1 md:grid-cols-2 gap-4">
                    """
                    for metric in metrics:
                        if metric.strip():
                            label, value = metric.strip('• ').split(':', 1)
                            metrics_html += f"""
                                <div class="flex flex-col">
                                    <span class="text-sm text-blue-600 font-medium">{label}</span>
                                    <span class="text-lg text-blue-900">{value}</span>
                                </div>
                            """
                    metrics_html += "</div></div>"

                issues_html_section = ""
                if issues_match:
                    issues_list = issues_match.group(1).strip().split('\n')
                    if issues_list:
                        issues_html_section = """
                            <div class="mt-4">
                                <h3 class="font-medium text-gray-800 mb-2">📝 Identified Issues:</h3>
                                <ul class="list-none space-y-2">
                        """
                        for issue in issues_list:
                            if issue.strip():
                                issues_html_section += f"""
                                    <li class="text-gray-600 ml-4">• {issue.strip('• ')}</li>
                                """
                        issues_html_section += "</ul></div>"

                # Combine the readability section
                issues_html += f"""
                    <div class="bg-white rounded-lg shadow-sm mb-6 overflow-hidden">
                        <div class="bg-gray-50 px-6 py-4 border-b">
                            <h2 class="text-lg font-semibold text-gray-800">{title}</h2>
                        </div>
                        <div class="px-6 py-4">
                            {metrics_html}
                            {issues_html_section}
                        </div>
                    </div>
                """
                continue

            # Extract description and solution
            description_parts = content.split('How to fix:', 1)
            description = description_parts[0].strip()
            solution = description_parts[1].split('Example Fix:', 1)[0].strip() if len(description_parts) > 1 else ""
            
            # Extract examples and issues
            examples_match = re.search(r'Example Fix:\s*❌[^✓]+✓[^•]+', content, re.MULTILINE | re.DOTALL)
            examples_html = ""
            if examples_match:
                examples_text = examples_match.group(0)
                incorrect = re.search(r'❌\s*Incorrect:\s*([^✓]+)', examples_text)
                correct = re.search(r'✓\s*Correct:\s*([^•\n]+)', examples_text)
                
                if incorrect and correct:
                    examples_html = f"""
                        <div class="mb-4">
                            <h3 class="font-medium text-gray-800 mb-2">Example Fix:</h3>
                            <div class="space-y-2 ml-4">
                                <div class="text-red-600">
                                    ❌ Incorrect:
                                </div>
                                <div class="text-red-600 ml-8">
                                    {incorrect.group(1).strip()}
                                </div>
                                <div class="text-green-600 mt-2">
                                    ✓ Correct:
                                </div>
                                <div class="text-green-600 ml-8">
                                    {correct.group(1).strip()}
                                </div>
                            </div>
                        </div>
                    """
            
            # Extract issues
            issues_match = re.findall(r'•\s*(.*?)(?=•|\Z)', content, re.DOTALL)
            issues_html_section = ""
            if issues_match:
                issues_html_section = """
                    <div class="mt-4">
                        <h3 class="font-medium text-gray-800 mb-2">Issues found in your document:</h3>
                        <ul class="list-none space-y-2">
                """
                for issue in issues_match[:30]:
                    clean_issue = issue.strip().lstrip('•').strip()
                    issues_html_section += f"""
                        <li class="text-gray-600 ml-4">• {clean_issue}</li>
                    """
                if len(issues_match) > 50:
                    issues_html_section += f"""
                        <li class="text-gray-500 italic ml-4">... and {len(issues_match) - 50} more similar issues.</li>
                    """
                issues_html_section += "</ul></div>"
            
            # Combine the section
            issues_html += f"""
                <div class="bg-white rounded-lg shadow-sm mb-6 overflow-hidden">
                    <div class="bg-gray-50 px-6 py-4 border-b">
                        <h2 class="text-lg font-semibold text-gray-800">{title}</h2>
                    </div>
                    
                    <div class="px-6 py-4">
                        <div class="text-gray-600 mb-4">
                            {description}
                        </div>
                        
                        <div class="bg-green-50 rounded p-4 mb-4">
                            <div class="text-green-800">
                                <span class="font-medium">How to fix: </span>
                                {solution}
                            </div>
                        </div>
                        
                        {examples_html}
                        {issues_html_section}
                    </div>
                </div>
            """
        
        # Add new CSS classes for readability metrics
        additional_styles = """
            .bg-blue-50 { background-color: #eff6ff; }
            .text-blue-600 { color: #2563eb; }
            .text-blue-800 { color: #1e40af; }
            .text-blue-900 { color: #1e3a8a; }
            .grid { display: grid; }
            .grid-cols-1 { grid-template-columns: repeat(1, minmax(0, 1fr)); }
            .gap-4 { gap: 1rem; }
            @media (min-width: 768px) {
                .md\\:grid-cols-2 { grid-template-columns: repeat(2, minmax(0, 1fr)); }
            }
        """
        
        # Add summary section before the final return
        summary_html = f"""
            <div class="bg-white rounded-lg shadow-sm mb-6 overflow-hidden">
                <div class="bg-gray-50 px-6 py-4 border-b">
                    <h2 class="text-lg font-semibold text-gray-800">📋 Next Steps</h2>
                </div>
                <div class="px-6 py-4">
                    <div class="space-y-4">
                        <div>
                            <h3 class="font-medium text-gray-800 mb-2">1. Review and Address Issues:</h3>
                            <ul class="list-none space-y-2 ml-4">
                                <li class="text-gray-600">• Review each issue category systematically</li>
                                <li class="text-gray-600">• Make corrections using the provided examples as guides</li>
                                <li class="text-gray-600">• Re-run document check to verify all issues are resolved</li>
                            </ul>
                        </div>
                        
                        <div>
                            <h3 class="font-medium text-gray-800 mb-2">2. Best Practices:</h3>
                            <ul class="list-none space-y-2 ml-4">
                                <li class="text-gray-600">• Use search/replace for consistent fixes</li>
                                <li class="text-gray-600">• Update your document template to prevent future issues</li>
                                <li class="text-gray-600">• Keep the style manuals and Orders handy while making corrections</li>
                            </ul>
                        </div>
                    </div>
                </div>
            </div>
        """

        # Update the final HTML to include the summary section
        full_html = f"""
        <div class="mx-auto p-4" style="font-family: system-ui, -apple-system, sans-serif;">
            <style>
                .text-2xl {{ font-size: 1.5rem; line-height: 2rem; }}
                .text-lg {{ font-size: 1.125rem; }}
                .text-sm {{ font-size: 0.875rem; }}
                .font-bold {{ font-weight: 700; }}
                .font-semibold {{ font-weight: 600; }}
                .font-medium {{ font-weight: 500; }}
                .text-gray-800 {{ color: #1f2937; }}
                .text-gray-600 {{ color: #4b5563; }}
                .text-gray-500 {{ color: #6b7280; }}
                .text-green-600 {{ color: #059669; }}
                .text-green-800 {{ color: #065f46; }}
                .text-red-600 {{ color: #dc2626; }}
                .text-amber-600 {{ color: #d97706; }}
                .bg-white {{ background-color: #ffffff; }}
                .bg-gray-50 {{ background-color: #f9fafb; }}
                .bg-green-50 {{ background-color: #ecfdf5; }}
                .rounded-lg {{ border-radius: 0.5rem; }}
                .shadow-sm {{ box-shadow: 0 1px 2px 0 rgba(0, 0, 0, 0.05); }}
                .mb-6 {{ margin-bottom: 1.5rem; }}
                .mb-4 {{ margin-bottom: 1rem; }}
                .mb-2 {{ margin-bottom: 0.5rem; }}
                .ml-4 {{ margin-left: 1rem; }}
                .ml-8 {{ margin-left: 2rem; }}
                .mt-2 {{ margin-top: 0.5rem; }}
                .mt-4 {{ margin-top: 1rem; }}
                .p-4 {{ padding: 1rem; }}
                .px-6 {{ padding-left: 1.5rem; padding-right: 1.5rem; }}
                .py-4 {{ padding-top: 1rem; padding-bottom: 1rem; }}
                .space-y-2 > * + * {{ margin-top: 0.5rem; }}
                .italic {{ font-style: italic; }}
                .border-b {{ border-bottom: 1px solid #e5e7eb; }}
                .overflow-hidden {{ overflow: hidden; }}
                .list-none {{ list-style-type: none; }}
                .space-y-4 > * + * {{ margin-top: 1rem; }}
                {additional_styles}
            </style>
            {header_html}
            {issues_html}
            {summary_html}
        </div>
        """
        
        return full_html

    # Define the function to read the README content
    def get_readme_content():
        readme_path = "README.md"
        try:
            with open(readme_path, "r", encoding="utf-8") as file:
                readme_content = file.read()
            return readme_content
        except Exception as e:
            logging.error(f"Error reading README.md: {str(e)}")
            return "Error loading help content."
    
    with gr.Blocks() as demo:
        with gr.Tabs():
            with gr.Tab("Document Checker"):
                gr.Markdown(
                    """
                # 📑 FAA Document Checker Tool

                This tool performs multiple **validation checks** on Word documents to ensure compliance with U.S. federal documentation standards. See the About tab for more information.

                ## How to Use

                1. Upload your Word document (`.docx` format).
                2. Select the document type.
                3. Click **Check Document**.

                > **Note:** Please ensure your document is clean (no track changes or comments). If your document has track changes or comments, you might get several false positives.
                """
                )
                
                with gr.Row():
                    with gr.Column(scale=1):
                        file_input = gr.File(
                            label="📎 Upload Word Document (.docx)",
                            file_types=[".docx"],
                            type="binary"
                        )
                        
                        doc_type = gr.Dropdown(
                            choices=document_types,
                            label="📋 Document Type",
                            value="Advisory Circular",
                            info="Select the type of document you're checking"
                        )
                        
                        template_type = gr.Radio(
                            choices=template_types,
                            label="📑 Template Type",
                            visible=False,
                            info="Only applicable for Advisory Circulars"
                        )
                        
                        submit_btn = gr.Button(
                            "🔍 Check Document",
                            variant="primary"
                        )
                    
                    with gr.Column(scale=2):
                        results = gr.HTML()
                
                # Download button
                # with gr.Row():
                #     download_btn = gr.Button(
                #         "⬇️ Download Results as PDF",
                #         variant="secondary",
                #         visible=False
                #     )
                #
                #     pdf_file = gr.File(
                #         label="Download PDF",
                #         visible=False,
                #         interactive=False,
                #         file_types=[".pdf"]
                #     )
                
                # Initialize document checker
                checker = FAADocumentChecker()

                def process_and_format(file_obj, doc_type_value, template_type_value):
                    """Process the uploaded file and format results."""
                    try:
                        # Run all checks
                        results_data = checker.run_all_document_checks(file_obj, doc_type_value, template_type_value)
                        
                        # Format results using DocumentCheckResultsFormatter
                        formatter = DocumentCheckResultsFormatter()
                        text_results = formatter.format_results(results_data, doc_type_value)
                        
                        # Convert to HTML
                        html_results = format_results_to_html(text_results)
                        
                        return html_results
                    except Exception as e:
                        return f"Error processing document: {str(e)}"
                
                # Update template type visibility based on document type
                def update_template_visibility(doc_type_value):
                    if doc_type_value == "Advisory Circular":
                        return gr.update(visible=True)
                    else:
                        return gr.update(visible=False)

                doc_type.change(
                    fn=update_template_visibility,
                    inputs=[doc_type],
                    outputs=[template_type]
                )

                # Handle document processing
                submit_btn.click(
                    fn=process_and_format,
                    inputs=[file_input, doc_type, template_type],
                    outputs=[results]  # Only output the results
                )

                # Function to generate PDF and provide it for download
                # def generate_pdf(html_content):
                #     try:
                #         # Use a temporary file to store the PDF
                #         with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                #             # Convert HTML to PDF using WeasyPrint
                #             HTML(string=html_content, base_url='.').write_pdf(tmp_pdf.name)
                #
                #         # Return the path to the PDF file
                #         return gr.update(value=tmp_pdf.name, visible=True)
                #     except Exception as e:
                #         logging.error(f"Error generating PDF with WeasyPrint: {str(e)}")
                #         traceback.print_exc()
                #         return gr.update(value=None, visible=False)


                # When the download button is clicked, generate the PDF
                # download_btn.click(
                #     fn=generate_pdf,
                #     inputs=[results],
                #     outputs=[pdf_file]
                # )

                
                gr.Markdown(
                    """
                    ### 📌 Important Notes
                    - This tool helps ensure compliance with federal documentation standards
                    - Results are based on current style guides and FAA requirements
                    - The tool provides suggestions but final editorial decisions rest with the document author
                    - For questions or feedback on the FAA documentation standards, contact the AIR-646 Senior Technical Writers
                    - For questions or feedback on the tool, contact Eric Putnam
                    - Results are not stored or saved
                    """
                )
            
            with gr.Tab("About"):
                readme_content = get_readme_content()
                gr.Markdown(readme_content)

    return demo

# Initialize and launch the interface
if __name__ == "__main__":
    # Create and launch the interface
    demo = create_interface()
    demo.launch(
        share=False,  # Set to True if you want to generate a public link
        server_name="0.0.0.0",  # Allows external access
        server_port=7860,  # Default Gradio port
        debug=True
    )

class PatternCache:
    """Cache for compiled regex patterns to avoid repeated compilation."""
    
    def __init__(self):
        self._cache: Dict[str, Pattern] = {}
        self._lock = threading.Lock()
    
    def get_pattern(self, pattern_str: str) -> Pattern:
        """
        Get a compiled pattern from cache or compile and cache it.
        
        Args:
            pattern_str: The regex pattern string to compile
            
        Returns:
            The compiled Pattern object
            
        Raises:
            ConfigurationError: If pattern compilation fails
        """
        with self._lock:
            if pattern_str not in self._cache:
                try:
                    self._cache[pattern_str] = re.compile(pattern_str)
                except re.error as e:
                    raise ConfigurationError(f"Invalid pattern '{pattern_str}': {e}")
            return self._cache[pattern_str]
    
    def clear(self) -> None:
        """Clear the pattern cache."""
        with self._lock:
            self._cache.clear()
