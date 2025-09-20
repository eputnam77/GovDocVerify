from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from govdocverify.utils import extract_docx_metadata


def test_extract_docx_metadata(tmp_path: Path) -> None:
    doc = Document()
    doc.core_properties.title = "My Title"
    doc.core_properties.author = "Alice"
    doc.core_properties.last_modified_by = "Bob"
    file_path = tmp_path / "meta.docx"
    doc.save(file_path)

    meta = extract_docx_metadata(str(file_path))
    assert meta["title"] == "My Title"
    assert meta["author"] == "Alice"
    assert meta["last_modified_by"] == "Bob"


def test_extract_docx_metadata_ignores_surrounding_whitespace(tmp_path: Path) -> None:
    doc = Document()
    doc.core_properties.title = "Whitespace"
    file_path = tmp_path / "meta.docx"
    doc.save(file_path)

    meta = extract_docx_metadata(f"  {file_path}  ")
    assert meta["title"] == "Whitespace"


def test_extract_docx_metadata_created_modified(tmp_path: Path) -> None:
    """Verify `created` and `modified` metadata are extracted."""
    doc = Document()
    doc.core_properties.created = datetime(2024, 1, 1, 12, 0, 0)
    doc.core_properties.modified = datetime(2024, 1, 2, 12, 0, 0)
    file_path = tmp_path / "meta2.docx"
    doc.save(file_path)

    meta = extract_docx_metadata(str(file_path))
    assert meta["created"].startswith("2024-01-01")
    assert meta["modified"].startswith("2024-01-02")


def test_extract_docx_metadata_timezone_normalized(tmp_path: Path) -> None:
    """MD-01: ensure timestamps are ISO-8601 UTC strings."""
    doc = Document()
    doc.core_properties.created = datetime(2024, 2, 1, 8, 0, tzinfo=timezone.utc)
    doc.core_properties.modified = datetime(2024, 2, 2, 9, 0, tzinfo=timezone.utc)
    file_path = tmp_path / "tz.docx"
    doc.save(file_path)

    meta = extract_docx_metadata(str(file_path))
    assert meta["created"] == "2024-02-01T08:00:00+00:00"
    assert meta["modified"] == "2024-02-02T09:00:00+00:00"


def test_extract_docx_metadata_missing_fields(tmp_path: Path) -> None:
    """MD-02: extractor handles missing metadata without crashing."""
    doc = Document()
    file_path = tmp_path / "blank.docx"
    doc.save(file_path)

    meta = extract_docx_metadata(str(file_path))
    assert "title" not in meta


def test_extract_docx_metadata_non_ascii(tmp_path: Path) -> None:
    """MD-05: UTF-8 metadata preserved during extraction."""
    doc = Document()
    doc.core_properties.title = "Résumé"
    doc.core_properties.author = "Jörg"
    file_path = tmp_path / "unicode.docx"
    doc.save(file_path)

    meta = extract_docx_metadata(str(file_path))
    assert meta["title"] == "Résumé"
    assert meta["author"] == "Jörg"


def test_extract_docx_metadata_matches_core_properties(tmp_path: Path) -> None:
    """MD-04: extracted metadata matches Word's core properties."""
    doc = Document()
    doc.core_properties.title = "Sample"
    doc.core_properties.author = "Author"
    doc.core_properties.last_modified_by = "Editor"
    file_path = tmp_path / "parity.docx"
    doc.save(file_path)

    meta = extract_docx_metadata(str(file_path))
    assert meta["title"] == doc.core_properties.title
    assert meta["author"] == doc.core_properties.author
    assert meta["last_modified_by"] == doc.core_properties.last_modified_by
