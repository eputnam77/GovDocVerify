# Run test: pytest -v
# pytest -v tests/test_accessibility_checks.py --log-cli-level=DEBUG

import logging
import unittest
from pathlib import Path
from unittest.mock import Mock

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from govdocverify.checks.accessibility_checks import AccessibilityChecks
from govdocverify.models import DocumentCheckResult, Severity
from govdocverify.utils.terminology_utils import TerminologyManager
from tests.test_base import TestBase

logger = logging.getLogger(__name__)


@pytest.fixture
def accessibility_checks():
    return AccessibilityChecks()


@pytest.fixture
def mock_document():
    doc = Mock(spec=Document)
    doc.inline_shapes = []
    return doc


@pytest.fixture
def mock_results():
    return DocumentCheckResult()


class TestAccessibilityChecks(TestBase):
    """Test cases for accessibility checking functionality."""

    @classmethod
    def setUpClass(cls):
        """Set up test class with shared resources."""
        super().setUpClass()
        cls.terminology_manager = TerminologyManager()
        cls.accessibility_checks = AccessibilityChecks(cls.terminology_manager)
        logger.debug("Initialized shared test resources")

    def setUp(self):
        """Set up individual test cases."""
        super().setUp()
        # Create a temporary test image
        self.test_image_path = Path("test_image.png")
        try:
            with open(self.test_image_path, "wb") as f:
                f.write(
                    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
                )
            logger.debug(f"Created test image at {self.test_image_path}")
        except Exception as e:
            logger.error(f"Failed to create test image: {e}")
            raise

    def tearDown(self):
        """Clean up after individual test cases."""
        super().tearDown()
        # Clean up test image
        try:
            if self.test_image_path.exists():
                self.test_image_path.unlink()
                logger.debug(f"Cleaned up test image at {self.test_image_path}")
        except Exception as e:
            logger.error(f"Failed to clean up test image: {e}")
            # Don't raise here to avoid masking test failures

    def _add_image_with_properties(
        self, doc: Document, image_path: str, name: str = None, descr: str = None, title: str = None
    ) -> None:
        """Helper method to add an image with properties to a document."""
        try:
            p = doc.add_paragraph()
            r = p.add_run()
            logger.debug(
                f"Adding image {image_path} with name={name}, descr={descr}, title={title}"
            )

            # Add the picture
            r.add_picture(image_path)

            # Get the inline shape that was just added
            inline_shapes = r._element.xpath(".//w:drawing//wp:inline")
            if not inline_shapes:
                raise ValueError("No inline shape found after adding picture")
            inline_shape = inline_shapes[0]
            logger.debug(f"Found inline shape: {inline_shape}")

            # Set properties
            docPr = inline_shape.xpath(".//wp:docPr")
            if not docPr:
                raise ValueError("No docPr element found in inline shape")
            docPr = docPr[0]

            if name:
                docPr.set("name", name)
                logger.debug(f"Set name to {name}")
            if descr:
                docPr.set("descr", descr)
                logger.debug(f"Set description to {descr}")
            if title:
                docPr.set("title", title)
                logger.debug(f"Set title to {title}")
        except Exception as e:
            logger.error(f"Failed to add image with properties: {e}")
            raise

    def test_invalid_image_file(self):
        """Test handling of invalid image files."""
        logger.debug("Starting test_invalid_image_file")
        doc = Document()

        # Create an invalid image file
        invalid_image = Path("invalid_image.png")
        try:
            with open(invalid_image, "wb") as f:
                f.write(b"invalid image data")

            with self.assertRaises(Exception):
                self._add_image_with_properties(doc, str(invalid_image))
        finally:
            if invalid_image.exists():
                invalid_image.unlink()
                logger.debug("Cleaned up invalid test image")

    def test_long_image_name(self):
        """Test handling of very long image names."""
        logger.debug("Starting test_long_image_name")
        doc = Document()

        long_name = "x" * 1000  # Very long name
        self._add_image_with_properties(doc, str(self.test_image_path), name=long_name)

        results = DocumentCheckResult()
        self.accessibility_checks._check_alt_text(doc, results)

        self.assertEqual(len(results.issues), 1)
        self.assertIn(long_name[:100], results.issues[0]["message"])  # Check truncated name

    def test_readability(self):
        """Test readability checking."""
        content = [
            "Short sentence.",
            (
                "This is a very long sentence that exceeds the recommended word count limit and "
                "should be broken down into multiple shorter sentences for better readability. "
                "This additional phrase ensures the sentence is long enough to trigger the check "
                "and will make the test pass as expected."
            ),
            "Another short sentence with a total word count of exactly twenty-five words here.",
        ]
        result = self.accessibility_checks.check_readability(content)
        self.assertFalse(result.success)
        self.assert_issue_contains(result, "Sentence too long")
        self.assert_issue_contains(result, "Word count exceeds")

    def test_section_508_compliance(self):
        """Test Section 508 compliance checking."""
        content = [
            "Document with proper structure.",
            "Content with appropriate formatting.",
            "Accessible document elements.",
        ]
        result = self.accessibility_checks.check_section_508_compliance(content)
        self.assertTrue(result.success)

    def test_complex_sentences(self):
        """Test handling of complex sentences."""
        content = [
            "The Federal Aviation Administration (FAA), which is responsible for overseeing "
            "civil aviation in the United States, including the regulation and oversight of "
            "aircraft operations, air traffic control, and the certification of personnel "
            "and aircraft, plays a crucial role in ensuring the safety and efficiency of "
            "the national airspace system."
        ]
        result = self.accessibility_checks.check_readability(content)
        self.assertFalse(result.success)
        self.assert_issue_contains(result, "Sentence too long")

    def test_heading_structure(self):
        """Test heading structure checking."""
        content = """
        # Main Heading
        ## Subheading
        ### Skip level heading
        # Another Main
        #### Too deep
        """
        result = self.accessibility_checks.check_heading_structure(content.split("\n"))
        self.assertFalse(result.success)
        self.assert_issue_contains(result, "Heading level jumps")

    def test_image_alt_text(self):
        """Test checking for image alt text."""
        content = """
        ![Descriptive alt text](image.png "Image title")
        ![Another image](photo.jpg "Photo description")
        """
        result = self.accessibility_checks.check_image_accessibility(content.split("\n"))
        self.assertTrue(result.success)

    def test_missing_alt_text(self):
        """Test detection of missing alt text."""
        content = """
        ![](image.png)
        """
        file_path = self.create_test_file(content, "test_accessibility.md")
        checker = AccessibilityChecks(terminology_manager=self.terminology_manager)
        result = checker.check_document(file_path, doc_type="MARKDOWN")
        self.assert_has_issues(result)
        self.assert_issue_contains(result, "Image at line 2 is missing alt text")

    def test_color_contrast(self):
        """Test color contrast checking."""
        content = """
        <span style="color: #777777">Low contrast text</span>
        <div style="background-color: #FFFF00; color: #FFFFFF">Poor contrast background</div>
        """
        results = DocumentCheckResult()
        self.accessibility_checks._check_color_contrast(content.split("\n"), results)
        self.assertFalse(results.success)
        self.assert_issue_contains(results, "Insufficient color contrast")

    def test_image_alt_text_filtering(self):
        """Test filtering of watermarks and table graphics from alt text checks."""
        logger.debug("Starting test_image_alt_text_filtering")
        doc = Document()

        # Add a regular image
        self._add_image_with_properties(doc, str(self.test_image_path), name="content_image")
        logger.debug("Added regular image")

        # Add a watermark
        self._add_image_with_properties(doc, str(self.test_image_path), name="watermark")
        logger.debug("Added watermark")

        # Add a table graphic
        self._add_image_with_properties(doc, str(self.test_image_path), name="table_border")
        logger.debug("Added table graphic")

        results = DocumentCheckResult()
        self.accessibility_checks._check_alt_text(doc, results)
        logger.debug(f"Check results: {results.issues}")

        # Should only flag the regular image for missing alt text
        self.assertEqual(len(results.issues), 1, "Expected exactly one issue for the regular image")
        logger.debug(f"Found {len(results.issues)} issues as expected")

        # Check message content
        message = results.issues[0]["message"]
        logger.debug(f"Checking message content: {message}")
        self.assertIn("content_image", message, "Message should contain the image name")
        self.assertIn(
            "missing alt text", message.lower(), "Message should indicate missing alt text"
        )
        logger.debug("Message content checks passed")

    def test_image_alt_text_with_name(self):
        """Test alt text check includes image name in context."""
        logger.debug("Starting test_image_alt_text_with_name")
        doc = Document()

        self._add_image_with_properties(doc, str(self.test_image_path), name="test_image")
        logger.debug("Added image with name")

        results = DocumentCheckResult()
        self.accessibility_checks._check_alt_text(doc, results)
        logger.debug(f"Check results: {results.issues}")

        self.assertEqual(len(results.issues), 1)
        self.assertIn("test_image", results.issues[0]["message"])
        logger.debug("Test completed successfully")

    def test_image_alt_text_with_alt(self):
        """Test that images with alt text are not flagged."""
        logger.debug("Starting test_image_alt_text_with_alt")
        doc = Document()

        self._add_image_with_properties(
            doc, str(self.test_image_path), name="test_image", descr="Test image description"
        )
        logger.debug("Added image with alt text")

        results = DocumentCheckResult()
        self.accessibility_checks._check_alt_text(doc, results)
        logger.debug(f"Check results: {results.issues}")

        self.assertEqual(len(results.issues), 0)
        logger.debug("Test completed successfully")

    def test_image_alt_text_with_title(self):
        """Test that images with title but no description are not flagged."""
        logger.debug("Starting test_image_alt_text_with_title")
        doc = Document()

        self._add_image_with_properties(
            doc, str(self.test_image_path), name="test_image", title="Test image title"
        )
        logger.debug("Added image with title")

        results = DocumentCheckResult()
        self.accessibility_checks._check_alt_text(doc, results)
        logger.debug(f"Check results: {results.issues}")

        self.assertEqual(len(results.issues), 0)
        logger.debug("Test completed successfully")

    def test_check_alt_text_with_document_missing_alt(self):
        """Test _check_alt_text with Document containing images missing alt text."""
        # Create mock shape with missing alt text
        mock_shape = Mock()
        mock_shape._inline = Mock()
        mock_shape._inline.docPr = {"name": "test_image"}
        mock_document = Mock(spec=Document)
        mock_document.inline_shapes = [mock_shape]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_alt_text(mock_document, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("missing alt text", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_alt_text_with_document_no_shapes(self):
        """Test _check_alt_text with Document containing no shapes."""
        mock_document = Mock(spec=Document)
        mock_document.inline_shapes = []
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_alt_text(mock_document, mock_results)

        self.assertTrue(mock_results.success)
        self.assertEqual(len(mock_results.issues), 0)

    def test_check_alt_text_with_empty_document(self):
        """Test _check_alt_text with empty document."""
        mock_document = Mock(spec=Document)
        mock_document.inline_shapes = []
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_alt_text(mock_document, mock_results)

        self.assertTrue(mock_results.success)
        self.assertEqual(len(mock_results.issues), 0)

    def test_check_alt_text_with_invalid_shape(self):
        """Test _check_alt_text with invalid shape object."""
        mock_shape = Mock()
        mock_shape._inline = None  # Invalid shape
        mock_document = Mock(spec=Document)
        mock_document.inline_shapes = [mock_shape]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_alt_text(mock_document, mock_results)

        self.assertTrue(mock_results.success)
        self.assertEqual(len(mock_results.issues), 0)

    def test_check_alt_text_with_invalid_type(self):
        """Test _check_alt_text with invalid content type."""
        invalid_content = 123  # Not Document or List[str]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_alt_text(invalid_content, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("Invalid content type", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_alt_text_with_none_content(self):
        """Test _check_alt_text with None content."""
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_alt_text(None, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("Invalid content type", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_alt_text_with_text_content(self):
        """Test _check_alt_text with text content (List[str])."""
        content = ["Some text", "![Missing Alt](image.jpg)", "![With Alt](image.jpg)"]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_alt_text(content, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("missing alt text", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_color_contrast_with_document(self):
        """Test _check_color_contrast with Document content."""
        mock_document = Mock(spec=Document)
        mock_document.paragraphs = [
            Mock(text="color: #000000; background-color: #FFFFFF"),  # Good contrast
            Mock(text="color: #000000; background-color: #111111"),  # Poor contrast
        ]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_color_contrast(mock_document, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("Insufficient color contrast ratio", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_color_contrast_with_empty_content(self):
        """Test _check_color_contrast with empty content."""
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_color_contrast([], mock_results)

        self.assertTrue(mock_results.success)
        self.assertEqual(len(mock_results.issues), 0)

    def test_check_color_contrast_with_invalid_colors(self):
        """Test _check_color_contrast with invalid color values."""
        content = [
            "color: #invalid; background-color: #FFFFFF",
            "color: #000000; background-color: #invalid",
        ]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_color_contrast(content, mock_results)

        self.assertTrue(mock_results.success)
        self.assertEqual(len(mock_results.issues), 0)

    def test_check_color_contrast_with_none_content(self):
        """Test _check_color_contrast with None content."""
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_color_contrast(None, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("Invalid content type", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_heading_hierarchy_with_empty_content(self):
        """Test _check_heading_hierarchy with empty content."""
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_heading_hierarchy([], mock_results)

        self.assertTrue(mock_results.success)
        self.assertEqual(len(mock_results.issues), 0)

    def test_check_heading_hierarchy_with_invalid_levels(self):
        """Test _check_heading_hierarchy with invalid heading levels."""
        headings = [("H1", "invalid"), ("H2", 2), ("H3", 3)]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_heading_hierarchy(headings, mock_results)

        # Updated expectation: should flag missing H1 and heading level skipped
        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 2)
        self.assertIn(
            "Document is missing a top-level heading (H1)", mock_results.issues[0]["message"]
        )
        self.assertIn("Heading level skipped", mock_results.issues[1]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)
        self.assertEqual(mock_results.issues[1]["severity"], Severity.ERROR)

    def test_check_heading_hierarchy_with_none_content(self):
        """Test _check_heading_hierarchy with None content."""
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_heading_hierarchy(None, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("Invalid content type", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_heading_hierarchy_with_skipped_levels(self):
        """Test _check_heading_hierarchy with skipped heading levels."""
        headings = [("H1", 1), ("H3", 3), ("H4", 4)]  # Skipped H2
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_heading_hierarchy(headings, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("Heading level skipped", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_heading_structure_with_document(self):
        """Test _check_heading_structure with Document content."""
        mock_document = Mock(spec=Document)
        mock_document.paragraphs = [
            Mock(style=Mock(name="Heading 2"), text="H2 Heading"),
            Mock(style=Mock(name="Heading 3"), text="H3 Heading"),
            Mock(style=Mock(name="Heading 4"), text="H4 Heading"),
        ]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_heading_structure(mock_document, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("missing a top-level heading", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_heading_structure_with_empty_content(self):
        """Test _check_heading_structure with empty content."""
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_heading_structure([], mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("missing a top-level heading", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_heading_structure_with_invalid_style(self):
        """Test _check_heading_structure with invalid heading style."""
        mock_paragraph = Mock()
        mock_paragraph.style = Mock(name="Invalid Style")
        mock_document = Mock(spec=Document)
        mock_document.paragraphs = [mock_paragraph]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_heading_structure(mock_document, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn(
            "Document is missing a top-level heading (H1)", mock_results.issues[0]["message"]
        )

    def test_check_heading_structure_with_none_content(self):
        """Test _check_heading_structure with None content."""
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_heading_structure(None, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn(
            "Invalid content type for heading structure check: None",
            mock_results.issues[0]["message"],
        )

    def test_check_hyperlinks_with_document(self):
        """Test _check_hyperlinks with Document content containing non-descriptive links."""
        mock_run = Mock()
        mock_run._element = Mock()
        mock_run._element.xpath.return_value = [True]
        mock_run.text = "click here"
        mock_paragraph = Mock()
        mock_paragraph.runs = [mock_run]
        mock_paragraph.text = "Some paragraph text"
        mock_document = Mock(spec=Document)
        mock_document.paragraphs = [mock_paragraph]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_hyperlinks(mock_document, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("Non-descriptive link text", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.WARNING)

    def test_unknown_paragraph_style_triggers_heading_issue(self):
        """VR-02: paragraphs with non-whitelisted styles are flagged."""
        doc = Document()
        custom_style = doc.styles.add_style("Custom", WD_STYLE_TYPE.PARAGRAPH)
        doc.add_paragraph("Some body text", style=custom_style)

        result = self.accessibility_checks.check_heading_structure(doc)
        assert not result.success
        assert any("top-level heading" in issue["message"] for issue in result.issues)

    def test_tables_without_headers_detected(self):
        """VR-08: table accessibility issues surface in 508 check."""
        checker = AccessibilityChecks()
        content = "This document mentions tables without headers in the appendix."
        result = checker.check_section_508_compliance(content)
        assert not result.success
        assert any("tables without headers" in issue["message"].lower() for issue in result.issues)

    def test_heading_outline_exports_to_json_tree(self):
        """VR-09: heading outline can be represented as a JSON tree."""
        doc = Document()
        doc.add_paragraph("1. INTRODUCTION.", style="Heading 1")
        doc.add_paragraph("1.1. DETAILS.", style="Heading 2")
        ac = AccessibilityChecks()
        headings = ac._extract_docx_headings(doc)

        def build_tree(headings):
            tree = []
            stack = []
            for level, text in headings:
                node = {"text": text, "children": []}
                while len(stack) >= level:
                    stack.pop()
                if stack:
                    stack[-1]["children"].append(node)
                else:
                    tree.append(node)
                stack.append(node)
            return tree

        outline = build_tree(headings)
        expected = [
            {"text": "1. INTRODUCTION.", "children": [{"text": "1.1. DETAILS.", "children": []}]}
        ]
        assert outline == expected

    def test_check_hyperlinks_with_empty_content(self):
        """Test _check_hyperlinks with empty content."""
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_hyperlinks([], mock_results)

        self.assertTrue(mock_results.success)
        self.assertEqual(len(mock_results.issues), 0)

    def test_check_hyperlinks_with_invalid_links(self):
        """Test _check_hyperlinks with invalid link formats."""
        content = [
            "[Invalid link",
            "Invalid link]",
            "[Invalid link]",
            "Invalid link](http://example.com)",
        ]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_hyperlinks(content, mock_results)

        self.assertTrue(mock_results.success)
        self.assertEqual(len(mock_results.issues), 0)

    def test_check_hyperlinks_with_malformed_urls(self):
        """Test _check_hyperlinks with malformed URLs."""
        content = ["[Click here](invalid url)", "[Learn more](http://)", "[Visit us](https://)"]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_hyperlinks(content, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(
            len(mock_results.issues), 2
        )  # Two non-descriptive links: "Click here" and "Learn more"
        self.assertIn("Click here", mock_results.issues[0]["message"])
        self.assertIn("Learn more", mock_results.issues[1]["message"])

    def test_check_hyperlinks_with_none_content(self):
        """Test _check_hyperlinks with None content."""
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_hyperlinks(None, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("Cannot check hyperlinks: content is None", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.ERROR)

    def test_check_hyperlinks_with_text_content(self):
        """Test _check_hyperlinks with text content containing non-descriptive links."""
        content = [
            "Some text with [click here](http://example.com)",
            "Good link: [Learn about accessibility](http://example.com)",
        ]
        mock_results = DocumentCheckResult()

        self.accessibility_checks._check_hyperlinks(content, mock_results)

        self.assertFalse(mock_results.success)
        self.assertEqual(len(mock_results.issues), 1)
        self.assertIn("Non-descriptive link text", mock_results.issues[0]["message"])
        self.assertEqual(mock_results.issues[0]["severity"], Severity.WARNING)


if __name__ == "__main__":
    unittest.main()
