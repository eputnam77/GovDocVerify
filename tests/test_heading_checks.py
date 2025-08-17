# python -m pytest tests/test_heading_checks.py -v
# pytest -v tests/test_heading_checks.py --log-cli-level=DEBUG

import pytest
from docx import Document

from govdocverify.checks.heading_checks import HeadingChecks
from govdocverify.models import Severity
from govdocverify.utils.terminology_utils import TerminologyManager


class TestHeadingChecks:
    @pytest.fixture(autouse=True)
    def setup(self):
        self.terminology_manager = TerminologyManager()
        self.heading_checks = HeadingChecks(self.terminology_manager)

    def test_validate_input(self):
        # Test valid input
        valid_doc = ["1. PURPOSE.", "2. BACKGROUND."]
        assert self.heading_checks.validate_input(valid_doc) is True

        # Test invalid input
        invalid_doc = ["1. PURPOSE.", 123]
        assert self.heading_checks.validate_input(invalid_doc) is False

    def test_check_heading_title_without_periods(self):
        """Test heading titles for document types that don't require periods."""
        doc = ["1. PURPOSE", "2. APPLICABILITY", "3. BACKGROUND", "4. DEFINITIONS"]
        result = self.heading_checks.check_heading_title(doc, "Advisory Circular")
        assert result.success is True
        # Accept INFO-level missing_optional_heading for CANCELLATION
        assert len(result.issues) == 1
        issue = result.issues[0]
        assert issue["type"] == "missing_optional_heading"
        assert issue["missing"] == "CANCELLATION"
        assert issue["severity"].name == "INFO"

    def test_check_heading_title_with_periods(self):
        """Test heading titles for document types that require periods."""
        doc = ["1. PURPOSE.", "2. APPLICABILITY.", "3. BACKGROUND.", "4. DEFINITIONS."]
        result = self.heading_checks.check_heading_title(doc, "ORDER")
        assert result.success is True
        assert len(result.issues) == 0

    def test_check_heading_period_without_periods(self):
        """Test heading periods for document types that don't require periods."""
        doc = ["1. PURPOSE", "2. BACKGROUND", "3. DEFINITIONS"]
        result = self.heading_checks.check_heading_period(doc, "Advisory Circular")
        # Now periods are required for Advisory Circular, so expect failure
        assert result.success is False
        assert all("missing required period" in issue["message"] for issue in result.issues)

    def test_check_heading_period_with_periods(self):
        """Test heading periods for document types that require periods."""
        doc = ["1. PURPOSE.", "2. BACKGROUND.", "3. DEFINITIONS."]
        result = self.heading_checks.check_heading_period(doc, "ORDER")
        assert result.success is True
        assert len(result.issues) == 0

    def test_check_heading_period_mixed(self):
        """Test heading periods with mixed usage (should fail)."""
        doc = ["1. PURPOSE", "2. BACKGROUND.", "3. DEFINITIONS"]
        result = self.heading_checks.check_heading_period(doc, "ORDER")
        assert result.success is False
        assert len(result.issues) > 0

    def test_check_heading_structure(self):
        doc = Document()
        doc.add_paragraph("1. PURPOSE.")
        doc.add_paragraph("2. BACKGROUND.")
        doc.add_paragraph("2.1. History.")
        doc.add_paragraph("2.2. Current Status.")
        doc.add_paragraph("3. DEFINITIONS.")

        issues = self.heading_checks.check_heading_structure(doc)
        assert len(issues) == 0

    def test_invalid_heading_structure(self):
        doc = Document()
        doc.add_paragraph("1. PURPOSE.")
        doc.add_paragraph("2. BACKGROUND.")
        doc.add_paragraph("2.1. History.")
        doc.add_paragraph("2.3. Current Status.")  # Skipped 2.2

        issues = self.heading_checks.check_heading_structure(doc)
        assert len(issues) > 0
        assert any("expected" in issue["message"] for issue in issues)

    def test_heading_case(self):
        """Test that headings are properly capitalized."""
        doc = ["1. Purpose.", "2. Background.", "3. Definitions."]
        result = self.heading_checks.check_heading_title(doc, "ORDER")
        assert result.success is False
        assert any("should be uppercase" in issue["message"] for issue in result.issues)
        assert all(
            issue["severity"] == Severity.WARNING
            for issue in result.issues
            if issue["type"] == "case_violation"
        )

    def test_heading_spacing(self):
        """Test spacing between headings and content."""
        doc = Document()
        doc.add_paragraph("1. PURPOSE.")
        doc.add_paragraph("This is some content.")
        doc.add_paragraph("2. BACKGROUND.")
        doc.add_paragraph("More content here.")
        doc.add_paragraph("3. DEFINITIONS.")
        doc.add_paragraph("Even more content.")

        issues = self.heading_checks.check_heading_structure(doc)
        assert len(issues) == 0  # No spacing issues

    def test_heading_with_content(self):
        """Test headings in context with actual content."""
        doc = Document()
        doc.add_paragraph("1. PURPOSE.")
        doc.add_paragraph("This document establishes requirements for...")
        doc.add_paragraph("2. BACKGROUND.")
        doc.add_paragraph("The Federal Aviation Administration...")
        doc.add_paragraph("3. DEFINITIONS.")
        doc.add_paragraph("For the purpose of this document...")

        issues = self.heading_checks.check_heading_structure(doc)
        assert len(issues) == 0  # No issues with headings and content

    def test_mixed_case_headings(self):
        """Test mixed case scenarios in headings."""
        doc = [
            "1. PURPOSE.",
            "2. Background.",
            "3. DEFINITIONS.",
            "4. Applicability.",
            "5. REQUIREMENTS.",
        ]
        result = self.heading_checks.check_heading_title(doc, "ORDER")
        # Warnings should cause overall failure
        assert result.success is False
        # Should have case violations and invalid word with warning severity
        assert any(
            issue["type"] == "case_violation" and issue["severity"] == Severity.WARNING
            for issue in result.issues
        )
        assert any(
            issue["type"] == "invalid_word" and issue["severity"] == Severity.WARNING
            for issue in result.issues
        )

    def test_document_type_matching(self):
        """Test document type matching with different case variations."""
        doc = ["1. PURPOSE", "2. BACKGROUND", "3. DEFINITIONS", "4. APPLICABILITY"]

        # Test with mixed case
        result = self.heading_checks.check_heading_title(doc, "Advisory Circular")
        assert result.success is True
        # Accept INFO-level missing_optional_heading for CANCELLATION
        assert len(result.issues) == 1
        issue = result.issues[0]
        assert issue["type"] == "missing_optional_heading"
        assert issue["missing"] == "CANCELLATION"
        assert issue["severity"].name == "INFO"
        assert (
            result.details["document_type"] == "ADVISORY_CIRCULAR"
        )  # Verify correct type was found

        # Test with lowercase
        result = self.heading_checks.check_heading_title(doc, "advisory circular")
        assert result.success is True
        assert len(result.issues) == 1
        issue = result.issues[0]
        assert issue["type"] == "missing_optional_heading"
        assert issue["missing"] == "CANCELLATION"
        assert issue["severity"].name == "INFO"
        assert result.details["document_type"] == "ADVISORY_CIRCULAR"

        # Test with uppercase
        result = self.heading_checks.check_heading_title(doc, "ADVISORY CIRCULAR")
        assert result.success is True
        assert len(result.issues) == 1
        issue = result.issues[0]
        assert issue["type"] == "missing_optional_heading"
        assert issue["missing"] == "CANCELLATION"
        assert issue["severity"].name == "INFO"
        assert result.details["document_type"] == "ADVISORY_CIRCULAR"

        # Test with extra spaces
        result = self.heading_checks.check_heading_title(doc, "  Advisory  Circular  ")
        assert result.success is True
        assert len(result.issues) == 1
        issue = result.issues[0]
        assert issue["type"] == "missing_optional_heading"
        assert issue["missing"] == "CANCELLATION"
        assert issue["severity"].name == "INFO"
        assert result.details["document_type"] == "ADVISORY_CIRCULAR"

        # Test with invalid document type
        result = self.heading_checks.check_heading_title(doc, "Invalid Type")
        assert result.success is True  # Should not fail, just return empty config
        assert len(result.issues) == 0
        assert result.details["document_type"] == "Invalid Type"  # Should preserve original type

    def test_heading_length_validation(self):
        """Test validation of heading length limits."""
        # Test with a heading that exceeds the maximum length
        long_heading_doc = [
            "1. THIS IS A VERY LONG HEADING THAT EXCEEDS THE MAXIMUM LENGTH LIMIT",
            "2. PURPOSE",
            "3. BACKGROUND",
            "4. DEFINITIONS",
            "5. APPLICABILITY",
        ]
        result = self.heading_checks.check_heading_title(long_heading_doc, "ORDER")
        # Length violations should trigger warnings, causing failure
        assert result.success is False
        assert any(
            issue["type"] == "length_violation" and issue["severity"] == Severity.WARNING
            for issue in result.issues
        )
        assert any("Shorten heading" in issue["suggestion"] for issue in result.issues)

        # Test with headings that are within the limit
        valid_length_doc = ["1. PURPOSE", "2. BACKGROUND", "3. DEFINITIONS", "4. APPLICABILITY"]
        result = self.heading_checks.check_heading_title(valid_length_doc, "ORDER")
        assert result.success is True
        assert len(result.issues) == 0

        # Test with a heading exactly at the limit using a valid heading word
        exact_length_doc = [
            "1. EFFECTIVE DATE",  # Exactly 13 characters
            "2. PURPOSE",
            "3. BACKGROUND",
            "4. DEFINITIONS",
            "5. APPLICABILITY",
        ]
        result = self.heading_checks.check_heading_title(exact_length_doc, "ORDER")
        assert result.success is True
        assert len(result.issues) == 0

        # Test with a heading just over the limit
        just_over_doc = [
            "1. THIS IS A VERY LONG HEADING THAT EXCEEDS THE MAXIMUM LENGTH LIMIT",  # 27 characters
            "2. PURPOSE",
            "3. BACKGROUND",
            "4. DEFINITIONS",
            "5. APPLICABILITY",
        ]
        result = self.heading_checks.check_heading_title(just_over_doc, "ORDER")
        assert result.success is False
        assert any(
            issue["type"] == "length_violation" and issue["severity"] == Severity.WARNING
            for issue in result.issues
        )
        assert any("Shorten heading" in issue["suggestion"] for issue in result.issues)

    def test_missing_cancellation_heading_info(self):
        """Test that missing 'CANCELLATION' heading in AC emits INFO, not error/warning."""
        doc = ["1. PURPOSE", "2. BACKGROUND", "3. DEFINITIONS", "4. APPLICABILITY"]
        result = self.heading_checks.check_heading_title(doc, "Advisory Circular")
        # Should succeed (no blocking errors/warnings)
        assert result.success is True
        # Should have an INFO-level issue for CANCELLATION
        info_issues = [i for i in result.issues if i.get("type") == "missing_optional_heading"]
        assert len(info_issues) == 1
        assert info_issues[0]["missing"] == "CANCELLATION"
        assert info_issues[0]["severity"].name == "INFO"
        assert "can be ignored" in info_issues[0]["message"]

    def test_ac_headings_require_period(self):
        doc = ["1. PURPOSE.", "2. BACKGROUND."]
        result = self.heading_checks.check_heading_period(doc, "Advisory Circular")
        assert result.success
        assert len(result.issues) == 0

    def test_ac_headings_missing_period(self):
        doc = ["1. PURPOSE", "2. BACKGROUND"]
        result = self.heading_checks.check_heading_period(doc, "Advisory Circular")
        assert not result.success
        assert any("missing required period" in issue["message"] for issue in result.issues)

    def test_ac_headings_extra_period(self):
        doc = ["1. PURPOSE..", "2. BACKGROUND."]
        result = self.heading_checks.check_heading_period(doc, "Advisory Circular")
        # Extra periods are not flagged as errors
        assert result.success
        assert len(result.issues) == 0

    def test_order_headings_require_period(self):
        doc = ["1. PURPOSE.", "2. BACKGROUND."]
        result = self.heading_checks.check_heading_period(doc, "ORDER")
        assert result.success
        assert len(result.issues) == 0

    def test_order_headings_missing_period(self):
        doc = ["1. PURPOSE", "2. BACKGROUND"]
        result = self.heading_checks.check_heading_period(doc, "ORDER")
        assert not result.success
        assert any("missing required period" in issue["message"] for issue in result.issues)

    def test_policy_statement_headings_require_period(self):
        doc = ["1. PURPOSE.", "2. BACKGROUND."]
        result = self.heading_checks.check_heading_period(doc, "POLICY_STATEMENT")
        assert result.success
        assert len(result.issues) == 0

    def test_policy_statement_headings_missing_period(self):
        doc = ["1. PURPOSE", "2. BACKGROUND"]
        result = self.heading_checks.check_heading_period(doc, "POLICY_STATEMENT")
        assert not result.success
        assert any("missing required period" in issue["message"] for issue in result.issues)

    def test_tso_headings_require_period(self):
        """Test that Technical Standard Orders require periods in headings."""
        doc = ["1. PURPOSE.", "2. BACKGROUND."]
        result = self.heading_checks.check_heading_period(doc, "Technical Standard Order")
        assert result.success
        assert len(result.issues) == 0

    def test_tso_headings_missing_period(self):
        """Test that Technical Standard Orders flag missing periods as errors."""
        doc = ["1. PURPOSE", "2. BACKGROUND"]
        result = self.heading_checks.check_heading_period(doc, "Technical Standard Order")
        assert not result.success
        assert any("missing required period" in issue["message"] for issue in result.issues)

    def test_tso_short_form_headings_require_period(self):
        """Test that TSO (short form) requires periods in headings."""
        doc = ["1. PURPOSE.", "2. BACKGROUND."]
        result = self.heading_checks.check_heading_period(doc, "TSO")
        assert result.success
        assert len(result.issues) == 0

    def test_tso_short_form_headings_missing_period(self):
        """Test that TSO (short form) flags missing periods as errors."""
        doc = ["1. PURPOSE", "2. BACKGROUND"]
        result = self.heading_checks.check_heading_period(doc, "TSO")
        assert not result.success
        assert any("missing required period" in issue["message"] for issue in result.issues)

    def test_run_checks_reports_skipped_heading_level_with_line_numbers(self):
        """VR-01: jumping from H1 to H3 triggers a high-severity issue."""

        doc = Document()
        doc.add_paragraph("1. INTRODUCTION.", style="Heading 1")
        doc.add_paragraph("1.1. DETAILS.", style="Heading 3")
        result = self.heading_checks.check_document(doc, "ORDER")
        assert not result.success
        assert len(result.issues) == 1
        issue = result.issues[0]
        assert "Missing heading H2" in issue["message"]
        assert issue["line_number"] == 2
        assert issue["severity"] == Severity.ERROR


if __name__ == "__main__":
    pytest.main()
