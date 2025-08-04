# pytest -v tests/test_formatting.py --log-cli-level=DEBUG
#
# NOTE: Uses the new `FormatChecks` and `FormattingChecker` helpers from
#       govdocverify.checks.format_checks.  No legacy `formatting_checks.py`.

from pathlib import Path

import pytest
from docx import Document

from govdocverify.checks.format_checks import FormatChecks, FormattingChecker
from govdocverify.models import Severity
from govdocverify.utils.terminology_utils import TerminologyManager


class TestFormattingChecks:
    # ------------------------------------------------------------------ #
    # Boilerplate                                                        #
    # ------------------------------------------------------------------ #

    @pytest.fixture(autouse=True)
    def setup(self):
        self.terminology_manager = TerminologyManager()
        self.format_checks = FormatChecks(self.terminology_manager)

    # ------------------------------------------------------------------ #
    # Baseline paragraph-level checks                                    #
    # ------------------------------------------------------------------ #

    def test_font_consistency(self):
        content = [
            "This is normal text.",
            "This is BOLD text.",
            "This is italic text.",
            "This is normal text again.",
        ]
        result = self.format_checks.check(content)
        assert not result["has_errors"]
        assert any("inconsistent font" in issue["message"].lower() for issue in result["warnings"])

    def test_spacing_consistency(self):
        content = [
            "This is a paragraph with single spacing.",
            "This is a paragraph with  double  spacing.",
            "This is a paragraph with   triple   spacing.",
            "This is a paragraph with single spacing again.",
        ]
        result = self.format_checks.check(content)
        assert not result["has_errors"]
        assert any("inconsistent spacing" in issue["message"].lower() for issue in result["warnings"])

    def test_margin_consistency(self):
        content = [
            "This is a paragraph with normal margins.",
            "    This is a paragraph with indented margins.",
            "This is a paragraph with normal margins again.",
            "        This is another paragraph with indented margins.",
        ]
        result = self.format_checks.check(content)
        assert not result["has_errors"]
        assert any("inconsistent margins" in issue["message"].lower() for issue in result["warnings"])

    def test_list_formatting(self):
        content = [
            "The following items are required:",
            "1. First item",
            "2. Second item",
            "a. Sub-item",
            "b. Another sub-item",
            "3. Third item",
        ]
        result = self.format_checks.check(content)
        assert not result["has_errors"]
        assert len(result["warnings"]) == 0

    def test_table_formatting(self):
        content = [
            "Table 1. Sample Table",
            "Column 1 | Column 2 | Column 3",
            "---------|----------|----------",
            "Data 1   | Data 2   | Data 3",
            "Data 4   | Data 5   | Data 6",
        ]
        result = self.format_checks.check(content)
        assert not result["has_errors"]
        assert len(result["warnings"]) == 0

    def test_heading_formatting(self):
        content = [
            "PURPOSE.",
            "This is the purpose section.",
            "BACKGROUND.",
            "This is the background section.",
            "DEFINITIONS.",
            "This is the definitions section.",
        ]
        result = self.format_checks.check(content)
        assert not result["has_errors"]
        assert len(result["warnings"]) == 0

    def test_reference_formatting(self):
        content = [
            "See paragraph 5.2.3 for more information.",
            "Refer to section 4.1.2 for details.",
            "As discussed in paragraph 3.4.5",
        ]
        result = self.format_checks.check(content)
        assert not result["has_errors"]
        assert any("inconsistent reference" in issue["message"].lower() for issue in result["warnings"])

    def test_figure_formatting(self):
        content = [
            "Figure 1. Sample Figure",
            "This is a figure caption.",
            "Figure 2. Another Sample Figure",
            "This is another figure caption.",
        ]
        result = self.format_checks.check(content)
        assert not result["has_errors"]
        assert len(result["warnings"]) == 0

    # ------------------------------------------------------------------ #
    # Fine-grained checker tests                                         #
    # ------------------------------------------------------------------ #

    def test_list_formatting_flags_number_and_bullet_issues(self):
        lines = [
            "1. First item",
            "2Second item",   # Missing period or space after number
            "•Third item",    # Missing space after bullet
        ]
        checker = FormattingChecker()
        result = checker.check_list_formatting(lines)
        assert not result.success
        messages = [issue["message"].lower() for issue in result.issues]
        assert any("list formatting" in m for m in messages)
        assert any("bullet spacing" in m for m in messages)
        assert {issue["line_number"] for issue in result.issues} == {2, 3}
        assert all(issue["severity"] == Severity.WARNING for issue in result.issues)

    def test_section_symbol_usage_detects_spacing(self):
        lines = [
            "See §123 for details",        # Missing space after symbol
            "Refer to §§123-456 for more", # Missing space after double symbol
        ]
        checker = FormattingChecker()
        result = checker.check_section_symbol_usage(lines)
        assert not result.success
        assert len(result.issues) == 2
        assert {issue["line_number"] for issue in result.issues} == {1, 2}
        assert all("section symbol" in issue["message"].lower() for issue in result.issues)
        assert all(issue["severity"] == Severity.WARNING for issue in result.issues)

    # ------------------------------------------------------------------ #
    # Caption schema checks                                              #
    # ------------------------------------------------------------------ #

    def test_caption_schema_requires_dash_for_order(self):
        doc = Document()
        doc.add_paragraph("Figure 1. Aircraft layout")
        checker = FormatChecks(self.terminology_manager)
        result = checker.check_document(doc, "Order")
        assert not result.success
        issue = result.issues[0]
        assert issue["line_number"] == 1
        assert issue["severity"] == Severity.ERROR
        assert "figure x-y" in issue["message"].lower()

    def test_caption_schema_disallows_dash_for_other_docs(self):
        doc = Document()
        doc.add_paragraph("Table 1-1. Schedule")
        checker = FormatChecks(self.terminology_manager)
        result = checker.check_document(doc, "Memo")
        assert not result.success
        issue = result.issues[0]
        assert issue["line_number"] == 1
        assert issue["severity"] == Severity.ERROR
        assert "table x" in issue["message"].lower()

    # ------------------------------------------------------------------ #
    # Edge-case checks (xfail until implemented)                         #
    # ------------------------------------------------------------------ #

    @pytest.mark.xfail(reason="Numbering continuity check not implemented")
    def test_list_numbering_continuity_flags_gaps(self):
        """VR-03: numbering gaps should be reported as formatting issues."""
        lines = [
            "1. First",
            "3. Third",  # Skips 2
        ]
        checker = FormattingChecker()
        result = checker.check_list_formatting(lines)
        assert not result.success
        messages = [issue["message"].lower() for issue in result.issues]
        assert any("list formatting" in m for m in messages)
        assert {issue["line_number"] for issue in result.issues} == {2}

    @pytest.mark.xfail(reason="Orphan bullet detection not implemented")
    def test_orphaned_bullet_is_detected(self):
        """VR-03: bullets without preceding list context should be flagged."""
        lines = [
            "Intro paragraph",
            "• Orphan bullet",
        ]
        checker = FormattingChecker()
        result = checker.check_list_formatting(lines)
        assert not result.success
        assert any("bullet" in issue["message"].lower() for issue in result.issues)

    def test_section_symbol_usage_allows_proper_spacing(self):
        """VR-06: correctly spaced section symbols should pass the check."""
        lines = ["See § 123 for details", "Refer to §§ 123-456 for more"]
        checker = FormattingChecker()
        result = checker.check_section_symbol_usage(lines)
        assert result.success
