# pytest -v tests/test_reference_checks.py --log-cli-level=DEBUG

import logging

import pytest
from docx import Document

from documentcheckertool.checks.reference_checks import (
    DocumentTitleFormatCheck,
    TableFigureReferenceCheck,
)
from documentcheckertool.models import DocumentCheckResult

# Configure logging for tests
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


@pytest.fixture
def reference_check():
    return TableFigureReferenceCheck()


@pytest.fixture
def mock_results():
    return DocumentCheckResult()


class TestTableFigureReferenceCheck:
    @pytest.fixture(autouse=True)
    def setup(self):
        logger.debug("Setting up test fixture")
        self.checker = TableFigureReferenceCheck()

    def test_empty_content(self):
        """Test handling of empty content."""
        logger.debug("Testing empty content handling")
        result = self.checker.check_text("")
        logger.debug(f"Empty content result: {result}")
        assert not result.success
        assert len(result.issues) == 1
        assert result.issues[0]["error"] == "Empty content"

    def test_none_content(self):
        """Test handling of None content."""
        logger.debug("Testing None content handling")
        result = self.checker.check(None)
        logger.debug(f"None content result: {result}")
        assert not result.success
        assert len(result.issues) == 1
        assert "Invalid document input" in result.issues[0]["error"]

    def test_invalid_content_type(self):
        """Test handling of invalid content type."""
        logger.debug("Testing invalid content type handling")
        result = self.checker.check(123)  # Integer instead of list
        logger.debug(f"Invalid content type result: {result}")
        assert not result.success
        assert len(result.issues) == 1
        assert "Invalid document input" in result.issues[0]["error"]

    def test_mixed_content_types(self):
        """Test handling of mixed content types in list."""
        logger.debug("Testing mixed content types handling")
        content = ["Table 1", 123, "Figure 2"]
        result = self.checker.check(content)
        logger.debug(f"Mixed content types result: {result}")
        assert not result.success
        assert len(result.issues) == 1
        assert "Invalid document input" in result.issues[0]["error"]

    def test_special_characters_in_references(self):
        """Test handling of special characters in references."""
        logger.debug("Testing special characters in references")
        content = [
            "Table 1.1 shows the results.",
            "As shown in Figure 2.3, the data indicates...",
            "The analysis in Table 3-4 demonstrates...",
        ]
        result = self.checker.check(content)
        logger.debug(f"Special characters result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_reference_in_parentheses(self):
        """Test handling of references in parentheses."""
        logger.debug("Testing references in parentheses")
        content = [
            "(Table 1) shows the results.",
            "As shown in (figure 2), the data indicates...",
            "The analysis in (Table 3) demonstrates...",
        ]
        result = self.checker.check(content)
        logger.debug(f"Parentheses references result: {result}")
        assert not result.success
        assert len(result.issues) == 2
        assert any("should be lowercase" in issue["issue"] for issue in result.issues)

    def test_reference_in_quotes(self):
        """Test handling of references in quotes."""
        logger.debug("Testing references in quotes")
        content = [
            '"Table 1" shows the results.',
            'As shown in "figure 2", the data indicates...',
            'The analysis in "Table 3" demonstrates...',
        ]
        result = self.checker.check(content)
        logger.debug(f"Quoted references result: {result}")
        assert not result.success
        assert len(result.issues) == 2
        assert any("should be lowercase" in issue["issue"] for issue in result.issues)

    def test_reference_in_list(self):
        """Test handling of references in list items."""
        logger.debug("Testing references in list items")
        content = [
            "• Table 1 shows the results.",
            "• As shown in figure 2, the data indicates...",
            "• The analysis in Table 3 demonstrates...",
        ]
        result = self.checker.check(content)
        logger.debug(f"List references result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_reference_in_table(self):
        """Test handling of references in table cells."""
        logger.debug("Testing references in table cells")
        content = [
            "| Table 1 | shows the results |",
            "| As shown in figure 2 | the data indicates |",
            "| The analysis in Table 3 | demonstrates |",
        ]
        result = self.checker.check(content)
        logger.debug(f"Table references result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_reference_in_code_block(self):
        """Test handling of references in code blocks."""
        logger.debug("Testing references in code blocks")
        content = [
            "```",
            "Table 1 shows the results.",
            "As shown in figure 2, the data indicates...",
            "The analysis in Table 3 demonstrates...",
            "```",
        ]
        result = self.checker.check(content)
        logger.debug(f"Code block references result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_reference_in_html(self):
        """Test handling of references in HTML tags."""
        logger.debug("Testing references in HTML tags")
        content = [
            "<p>Table 1 shows the results.</p>",
            "<div>As shown in figure 2, the data indicates...</div>",
            "<span>The analysis in Table 3 demonstrates...</span>",
        ]
        result = self.checker.check(content)
        logger.debug(f"HTML references result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_reference_in_markdown(self):
        """Test handling of references in markdown formatting."""
        logger.debug("Testing references in markdown formatting")
        content = [
            "**Table 1** shows the results.",
            "*As shown in figure 2*, the data indicates...",
            "`The analysis in Table 3` demonstrates...",
        ]
        result = self.checker.check(content)
        logger.debug(f"Markdown references result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_reference_with_multiple_spaces(self):
        """Test handling of references with multiple spaces."""
        logger.debug("Testing references with multiple spaces")
        content = [
            "Table  1 shows the results.",
            "As shown in figure  2, the data indicates...",
            "The analysis in Table  3 demonstrates...",
        ]
        result = self.checker.check(content)
        logger.debug(f"Multiple spaces result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_reference_with_tabs(self):
        """Test handling of references with tabs."""
        logger.debug("Testing references with tabs")
        content = [
            "Table\t1 shows the results.",
            "As shown in figure\t2, the data indicates...",
            "The analysis in Table\t3 demonstrates...",
        ]
        result = self.checker.check(content)
        logger.debug(f"Tab references result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_reference_with_newlines(self):
        """Test handling of references with newlines."""
        logger.debug("Testing references with newlines")
        content = [
            "Table\n1 shows the results.",
            "As shown in figure\n2, the data indicates...",
            "The analysis in Table\n3 demonstrates...",
        ]
        result = self.checker.check(content)
        logger.debug(f"Newline references result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_caption_skipping(self):
        """Test that caption lines are skipped."""
        logger.debug("Testing caption line skipping")
        content = [
            "Table 1. Sample Table",
            "This is a reference to table 1.",
            "Figure 2. Sample Figure",
            "This is a reference to figure 2.",
        ]
        logger.debug(f"Test content: {content}")
        result = self.checker.check(content, "GENERAL")
        logger.debug(f"Caption skipping result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_sentence_start_capitalization(self):
        """Test capitalization at sentence start."""
        logger.debug("Testing sentence start capitalization")
        content = [
            "table 1 shows the data.",
            "figure 2 illustrates the process.",
            "As shown in table 1, the data is clear.",
            "The process is shown in figure 2.",
        ]
        logger.debug(f"Test content: {content}")
        result = self.checker.check(content, "GENERAL")
        logger.debug(f"Sentence start capitalization result: {result}")
        assert not result.success
        assert len(result.issues) == 2
        assert any(
            "Table reference at sentence start should be capitalized" in i["issue"]
            for i in result.issues
        )
        assert any(
            "Figure reference at sentence start should be capitalized" in i["issue"]
            for i in result.issues
        )

    def test_inline_lowercase(self):
        """Test lowercase within sentences."""
        logger.debug("Testing inline lowercase references")
        content = [
            "The data is shown in Table 1.",
            "The process is illustrated in Figure 2.",
            "Table 1 shows the data.",
            "Figure 2 illustrates the process.",
        ]
        logger.debug(f"Test content: {content}")
        result = self.checker.check(content, "GENERAL")
        logger.debug(f"Inline lowercase result: {result}")
        assert not result.success
        assert len(result.issues) == 2
        assert any(
            "Table reference within sentence should be lowercase" in i["issue"]
            for i in result.issues
        )
        assert any(
            "Figure reference within sentence should be lowercase" in i["issue"]
            for i in result.issues
        )

    def test_mixed_references(self):
        """Test mixed correct and incorrect references."""
        logger.debug("Testing mixed references")
        content = [
            "Table 1 shows the data.",
            "The data is shown in table 1.",
            "Figure 2 illustrates the process.",
            "The process is shown in figure 2.",
        ]
        logger.debug(f"Test content: {content}")
        result = self.checker.check(content, "GENERAL")
        logger.debug(f"Mixed references result: {result}")
        assert result.success
        assert len(result.issues) == 0

    def test_reference_with_colon(self):
        """Test references after a colon (should be capitalized)."""
        logger.debug("Testing references after colon")
        content = ["The following data is shown: table 1", "The process is illustrated: figure 2"]
        logger.debug(f"Test content: {content}")
        result = self.checker.check(content, "GENERAL")
        logger.debug(f"Colon reference result: {result}")
        assert not result.success
        assert len(result.issues) == 2
        assert any(
            "Table reference at sentence start should be capitalized" in i["issue"]
            for i in result.issues
        )
        assert any(
            "Figure reference at sentence start should be capitalized" in i["issue"]
            for i in result.issues
        )

    def test_reference_with_semicolon(self):
        """Test references after a semicolon (should be capitalized)."""
        logger.debug("Testing references after semicolon")
        content = [
            "First point; table 1 shows the data",
            "Second point; figure 2 illustrates the process",
        ]
        logger.debug(f"Test content: {content}")
        result = self.checker.check(content, "GENERAL")
        logger.debug(f"Semicolon reference result: {result}")
        assert not result.success
        assert len(result.issues) == 2
        assert any(
            "Table reference at sentence start should be capitalized" in i["issue"]
            for i in result.issues
        )
        assert any(
            "Figure reference at sentence start should be capitalized" in i["issue"]
            for i in result.issues
        )

    def test_multiple_references(self):
        """Test multiple references in one sentence."""
        logger.debug("Testing multiple references in one sentence")
        content = [
            "The data in table 1 and Table 2 shows the process.",
            "The process in Figure 1 and figure 2 is illustrated.",
        ]
        logger.debug(f"Test content: {content}")
        result = self.checker.check(content, "GENERAL")
        logger.debug(f"Multiple references result: {result}")
        assert not result.success
        assert len(result.issues) == 2
        assert any(
            "Table reference within sentence should be lowercase" in i["issue"]
            for i in result.issues
        )
        assert any(
            "Figure reference within sentence should be lowercase" in i["issue"]
            for i in result.issues
        )

    def test_check_with_invalid_input(self, reference_check):
        """Test check method with invalid input type."""
        invalid_input = 123  # Not a list

        result = reference_check.check(invalid_input)

        assert not result.success
        assert len(result.issues) == 1
        assert "Invalid document input" in result.issues[0]["error"]

    def test_check_with_empty_list(self, reference_check):
        """Test check method with empty list."""
        result = reference_check.check([])

        assert result.success
        assert len(result.issues) == 0

    def test_check_with_valid_references(self, reference_check):
        """Test check method with valid table/figure references."""
        content = [
            "Table 1 shows the results.",
            "As shown in figure 2, the data indicates...",
            "The analysis in table 3 demonstrates...",
        ]

        result = reference_check.check(content)

        assert result.success
        assert len(result.issues) == 0

    def test_check_with_invalid_capitalization(self, reference_check):
        """Test check method with invalid capitalization in references."""
        content = [
            "table 1 shows the results.",  # Should be capitalized at start
            "As shown in Figure 2, the data indicates...",  # Should be lowercase in middle
        ]

        result = reference_check.check(content)

        assert not result.success
        assert len(result.issues) == 2
        assert any("should be capitalized" in issue["issue"] for issue in result.issues)
        assert any("should be lowercase" in issue["issue"] for issue in result.issues)

    def test_check_with_caption_lines(self, reference_check):
        """Test check method with caption lines that should be ignored."""
        content = [
            "Table 1. Results of the analysis",
            "The data in table 1 shows...",
            "Figure 2. Distribution of values",
            "As shown in figure 2...",
        ]

        result = reference_check.check(content)

        assert result.success
        assert len(result.issues) == 0

    def test_check_text_with_empty_content(self, reference_check):
        """Test check_text method with empty content."""
        result = reference_check.check_text([])
        assert result.success
        assert len(result.issues) == 0

    def test_check_text_with_valid_references(self, reference_check):
        """Test check_text method with valid references."""
        content = ["Table 1 shows the results.", "As shown in figure 2, the data indicates..."]

        result = reference_check.check_text(content)
        assert result.success
        assert len(result.issues) == 0

    def test_check_text_with_invalid_references(self, reference_check):
        """Test check_text method with invalid references."""
        content = ["table 1 shows the results.", "As shown in Figure 2, the data indicates..."]

        result = reference_check.check_text(content)
        assert not result.success
        assert len(result.issues) == 2
        assert any("should be capitalized" in issue["issue"] for issue in result.issues)
        assert any("should be lowercase" in issue["issue"] for issue in result.issues)

    def test_validate_input_with_invalid_type(self, reference_check):
        """Test validate_input method with invalid type."""
        assert not reference_check.validate_input(123)  # Not a list
        assert not reference_check.validate_input(["text", 123])  # List with non-string

    def test_validate_input_with_valid_input(self, reference_check):
        """Test validate_input method with valid input."""
        assert reference_check.validate_input(["text", "more text"])

    def test_check_with_mixed_content(self, reference_check):
        """Test check method with mixed content types."""
        content = [
            "Table 1. Introduction",
            "The results are shown in table 1.",
            "Figure 2. Methodology",
            "As demonstrated in Figure 2, the process...",
            "Table 3. Results",
            "The data in table 3 indicates...",
        ]

        result = reference_check.check(content)

        assert not result.success
        assert (
            len(result.issues) == 1
        )  # Only Figure 2 should be flagged as uppercase within sentence
        assert all("should be lowercase" in issue["issue"] for issue in result.issues)

    def test_check_with_complex_references(self, reference_check):
        """Test check method with complex reference patterns."""
        content = [
            "Tables 1-3 show the results.",
            "As shown in figures 2-4, the data indicates...",
            "The analysis in table 5 demonstrates...",
        ]

        result = reference_check.check(content)

        assert result.success
        assert len(result.issues) == 0

    @pytest.mark.parametrize(
        "text, violations",
        [
            ("Analyze table 1 of this AC, Table 1-1 of this AC.", 0),
            ("As shown in Table 3-2, the component values...", 0),
            ("Refer to table 3-2 for limits.", 0),
            ("The tests identified in Table 1-1 (see Section 3) are required.", 0),
            ('"Table 2-1" lists the data.', 1),
        ],
    )
    def test_table_figure_case_param(self, text, violations):
        """Parametrized test for table/figure reference case and context edge cases."""
        logger.debug(f"Param test: '{text}' expecting {violations} violations")
        result = self.checker.check(text.split("\n"))
        assert len(result.issues) == violations
        if violations:
            assert result.issues[0]["issue"].endswith("lowercase")


class TestDocumentTitleFormatting:
    """Test document title formatting checks."""

    def setup_method(self):
        """Set up test fixtures."""
        self.title_checker = DocumentTitleFormatCheck()

    def test_ac_title_no_formatting_should_fail(self):
        """Test that AC titles without formatting are flagged."""
        content = ["Use AC 33.91, Engine System and Component Tests, dated 7/25/2020."]
        result = self.title_checker.check_text(content, "Advisory Circular")

        assert not result.success
        assert len(result.issues) == 1
        assert "should be formatted in italics" in result.issues[0]["issue"]
        assert result.issues[0]["correct_format"] == "*Engine System and Component Tests*"

    def test_ac_title_with_quotes_should_fail(self):
        """Test that AC titles with quotes are flagged."""
        content = ['Use AC 33.91, "Engine System and Component Tests," dated July 25, 2020.']
        result = self.title_checker.check_text(content, "Advisory Circular")

        assert not result.success
        assert len(result.issues) == 1
        assert "should use italics, not quotation marks" in result.issues[0]["issue"]
        assert result.issues[0]["correct_format"] == "*Engine System and Component Tests*"

    def test_ac_title_with_italics_should_pass(self):
        """Test that AC titles with italics are correct."""
        content = ["Use AC 33.91, *Engine System and Component Tests*, dated 25 July 2020."]
        result = self.title_checker.check_text(content, "Advisory Circular")

        assert result.success
        assert len(result.issues) == 0

    def test_ac_title_with_word_italics_should_pass(self):
        """AC titles italicized in a DOCX file should be detected."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Use AC 33.91, ")
        italic_run = p.add_run("Engine System and Component Tests")
        italic_run.italic = True
        p.add_run(", dated 25 July 2020.")
        result = self.title_checker.check_document(doc, "Advisory Circular")

        assert result.success
        assert len(result.issues) == 0

    def test_non_ac_title_with_italics_should_fail(self):
        """Test that non-AC titles with italics are flagged."""
        content = [
            "Use Order 8900.1, *Flight Standards Information System*, dated January 1, 2020."
        ]
        result = self.title_checker.check_text(content, "Order")

        assert not result.success
        assert len(result.issues) == 1
        assert "should use quotation marks, not italics" in result.issues[0]["issue"]
        assert result.issues[0]["correct_format"] == '"Flight Standards Information System"'

    def test_non_ac_title_no_formatting_should_fail(self):
        """Test that non-AC titles without formatting are flagged."""
        content = ["Use Order 8900.1, Flight Standards Information System, dated January 1, 2020."]
        result = self.title_checker.check_text(content, "Order")

        assert not result.success
        assert len(result.issues) == 1
        assert "should be formatted in quotation marks" in result.issues[0]["issue"]
        assert result.issues[0]["correct_format"] == '"Flight Standards Information System"'

    def test_non_ac_title_with_quotes_should_pass(self):
        """Test that non-AC titles with quotes are correct."""
        content = [
            'Use Order 8900.1, "Flight Standards Information System," dated January 1, 2020.'
        ]
        result = self.title_checker.check_text(content, "Order")

        assert result.success
        assert len(result.issues) == 0

    def test_multiple_ac_references_mixed_formatting(self):
        """Test multiple AC references with mixed formatting."""
        content = [
            "Use AC 33.91, Engine System and Component Tests, dated 7/25/2020.",
            "Use AC 33.91, *Engine System and Component Tests*, dated 25 July 2020.",
            'Use AC 33.91, "Engine System and Component Tests," dated July 25, 2020.',
        ]
        result = self.title_checker.check_text(content, "Advisory Circular")

        assert not result.success
        assert len(result.issues) == 2  # Two incorrect formats

        # Check that the issues are for the unformatted and quoted versions
        issue_types = [issue["issue"] for issue in result.issues]
        assert any("should be formatted in italics" in issue for issue in issue_types)
        assert any("should use italics, not quotation marks" in issue for issue in issue_types)

    def test_no_ac_references_should_pass(self):
        """Test content without AC references."""
        content = ["This document contains no AC references.", "It should pass without any issues."]
        result = self.title_checker.check_text(content, "Advisory Circular")

        assert result.success
        assert len(result.issues) == 0

    def test_ac_reference_without_title_should_pass(self):
        """Test AC reference without a title."""
        content = ["See AC 33.91 for more information.", "AC 25.1309 is also relevant."]
        result = self.title_checker.check_text(content, "Advisory Circular")

        assert result.success
        assert len(result.issues) == 0

    def test_string_input_instead_of_list(self):
        """Test that string input is properly handled."""
        content = "Use AC 33.91, Engine System and Component Tests, dated 7/25/2020."
        result = self.title_checker.check_text(content, "Advisory Circular")

        assert not result.success
        assert len(result.issues) == 1
        assert "should be formatted in italics" in result.issues[0]["issue"]
