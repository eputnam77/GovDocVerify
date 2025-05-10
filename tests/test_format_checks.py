# python -m pytest tests/test_format_checks.py -v

import pytest
from documentcheckertool.checks.format_checks import FormatChecks, FormattingChecker
from documentcheckertool.utils.terminology_utils import TerminologyManager
from docx import Document
from documentcheckertool.models import DocumentCheckResult

class TestFormatChecks:
    @pytest.fixture(autouse=True)
    def setup(self):
        self.terminology_manager = TerminologyManager()
        self.format_checks = FormatChecks(self.terminology_manager)
        self.formatting_checker = FormattingChecker(self.terminology_manager)

    def test_check_date_formats(self):
        doc = Document()
        doc.add_paragraph("Date: 01/01/2023")  # Incorrect format
        doc.add_paragraph("Date: 2023-01-01")  # Correct format

        results = DocumentCheckResult()
        self.format_checks._check_date_formats([p.text for p in doc.paragraphs], results)
        assert len(results.issues) > 0
        assert any("Incorrect date format" in issue.message for issue in results.issues)

    def test_check_phone_numbers(self):
        doc = Document()
        doc.add_paragraph("Phone: (123) 456-7890")
        doc.add_paragraph("Phone: 123-456-7890")

        results = DocumentCheckResult()
        self.format_checks._check_phone_numbers([p.text for p in doc.paragraphs], results)
        assert len(results.issues) > 0
        assert any("Inconsistent phone number format" in issue.message for issue in results.issues)

    def test_check_placeholders(self):
        doc = Document()
        doc.add_paragraph("Lorem ipsum dolor sit amet")
        doc.add_paragraph("TODO: Add content here")

        results = DocumentCheckResult()
        self.format_checks._check_placeholders([p.text for p in doc.paragraphs], results)
        assert len(results.issues) > 0
        assert any("Placeholder text found" in issue.message for issue in results.issues)

    def test_check_dash_spacing(self):
        doc = Document()
        doc.add_paragraph("This is a test - with incorrect spacing")
        doc.add_paragraph("This is a test-without-spaces")

        results = DocumentCheckResult()
        self.format_checks._check_dash_spacing([p.text for p in doc.paragraphs], results)
        assert len(results.issues) > 0
        assert any("Remove spaces around dash" in issue.message for issue in results.issues)

    def test_formatting_checker(self):
        content = """
        This is a test document.
        It has some formatting issues:
        •Inconsistent bullet spacing
        •  Extra spaces
        "Mixed" quotation marks"
        """
        result = self.formatting_checker.check_text(content)
        assert len(result.issues) > 0
        assert any("Inconsistent bullet spacing" in issue.message for issue in result.issues)
        assert any("Extra spaces" in issue.message for issue in result.issues)
        assert any("Inconsistent quotation marks" in issue.message for issue in result.issues)