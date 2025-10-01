import pytest
from docx import Document

from govdocverify.checks.structure_checks import StructureChecks
from govdocverify.document_checker import FAADocumentChecker
from govdocverify.utils.terminology_utils import TerminologyManager


class TestStructureChecks:
    @pytest.fixture(autouse=True)
    def setup(self):
        self.structure_checks = StructureChecks(TerminologyManager())
        self.document_checker = FAADocumentChecker()

    def test_paragraph_length(self):
        long_para = " ".join([f"Sentence {i}." for i in range(7)])
        result = self.document_checker.check_paragraph_length([long_para])
        assert any("Paragraph" in issue["message"] for issue in result.issues)

    def test_sentence_length(self):
        long_sentence = "word " * 40
        result = self.document_checker.check_sentence_length(long_sentence)
        assert any("Sentence" in issue["message"] for issue in result.issues)

    def test_section_balance(self):
        doc = Document()
        doc.add_paragraph("SECTION 1. PURPOSE.", style="Heading 1")
        doc.add_paragraph("Short section content")
        doc.add_paragraph("SECTION 2. BACKGROUND.", style="Heading 1")
        for _ in range(50):
            doc.add_paragraph("Long section content")
        result = self.structure_checks.check_document(doc, "internal_review")
        assert any("section" in issue["message"].lower() for issue in result.issues)

    def test_list_formatting(self):
        doc = Document()
        doc.add_paragraph("Items:")
        doc.add_paragraph("• First item")
        doc.add_paragraph("- Second item")
        doc.add_paragraph("* Third item")
        result = self.structure_checks.check_document(doc, "internal_review")
        assert any("list" in issue["message"].lower() for issue in result.issues)

    def test_parentheses(self):
        doc = Document()
        doc.add_paragraph("This line has (unmatched parentheses.")
        result = self.structure_checks.check_document(doc, "internal_review")
        assert any("parentheses" in issue["message"].lower() for issue in result.issues)

    def test_watermark_missing(self):
        doc = Document()
        doc.add_paragraph("Body text")
        result = self.structure_checks.check_document(doc, "internal_review")
        assert any("watermark" in issue["message"].lower() for issue in result.issues)

    def test_required_ac_paragraphs(self):
        doc = Document()
        doc.add_paragraph("1. PURPOSE.", style="Heading 1")
        result = self.structure_checks.check_document(doc, "Advisory Circular")
        assert any("Advisory Circular" in issue["message"] for issue in result.issues)

    def test_footnote_sequence_gap_detected(self):
        doc = Document()
        doc.add_paragraph("Intro text [1].")
        doc.add_paragraph("Later reference [3].")

        result = self.structure_checks.check_document(doc, "internal_review")

        footnote_issues = [
            issue for issue in result.issues if "footnote" in issue["message"].lower()
        ]
        assert any("gap" in issue["message"].lower() for issue in footnote_issues)

    def test_footnote_duplicate_detected(self):
        doc = Document()
        doc.add_paragraph("Intro text [1].")
        doc.add_paragraph("Repeated reference [1].")

        result = self.structure_checks.check_document(doc, "internal_review")

        footnote_issues = [
            issue for issue in result.issues if "footnote" in issue["message"].lower()
        ]
        assert any("duplic" in issue["message"].lower() for issue in footnote_issues)

    def test_footnote_out_of_order_detected(self):
        doc = Document()
        doc.add_paragraph("Intro text [1].")
        doc.add_paragraph("Skipped footnote [3].")
        doc.add_paragraph("Then earlier number [2].")

        result = self.structure_checks.check_document(doc, "internal_review")

        footnote_issues = [
            issue for issue in result.issues if "footnote" in issue["message"].lower()
        ]
        assert any("order" in issue["message"].lower() for issue in footnote_issues)

    def test_footnote_reset_allowed_in_appendix(self):
        doc = Document()
        doc.add_paragraph("Intro text [1].")
        doc.add_paragraph("Another reference [2].")
        doc.add_paragraph("Appendix A", style="Heading 1")
        doc.add_paragraph("Appendix reference [1].")
        doc.add_paragraph("Appendix follow up [2].")

        result = self.structure_checks.check_document(doc, "internal_review")

        footnote_issues = [
            issue for issue in result.issues if "footnote" in issue["message"].lower()
        ]
        assert not footnote_issues

    def test_footnote_sequence_passes_when_ordered(self):
        doc = Document()
        doc.add_paragraph("Intro text [1].")
        doc.add_paragraph("Continuation [2].")
        doc.add_paragraph("More references [3].")

        result = self.structure_checks.check_document(doc, "internal_review")

        footnote_issues = [
            issue for issue in result.issues if "footnote" in issue["message"].lower()
        ]
        assert not footnote_issues
