# pytest -v tests/test_formats.py --log-cli-level=DEBUG

# This file consolidates all format-checking tests for both string and DOCX inputs.
# Shared test data ensures DRY, maintainable coverage of all public APIs.

import logging
import os
import tempfile

import pytest
from docx import Document

from govdocverify.checks.format_checks import FormatChecks, FormattingChecker
from govdocverify.models import DocumentCheckResult, Severity
from govdocverify.utils.terminology_utils import TerminologyManager

logger = logging.getLogger(__name__)


class TestFormatChecks:
    @pytest.fixture(autouse=True)
    def setup(self):
        """Set up test fixtures."""
        self.terminology_manager = TerminologyManager()
        self.format_checks = FormatChecks(self.terminology_manager)
        self.formatting_checker = FormattingChecker(self.terminology_manager)
        logger.debug("Initialized test with FormatChecks")
        logger.debug(f"Available Severity values: {[s.value for s in Severity]}")

    def create_test_docx(self, content: list, filename: str) -> str:
        """Create a temporary DOCX file with given content."""
        doc = Document()
        for paragraph in content:
            doc.add_paragraph(paragraph)

        temp_path = os.path.join(tempfile.gettempdir(), filename)
        doc.save(temp_path)
        return temp_path

    def test_font_consistency(self):
        """Test font consistency check."""
        content = [
            "This is normal text.",
            "This is BOLD text.",
            "This is italic text.",
            "This is normal text again.",
        ]
        doc_path = self.create_test_docx(content, "font_consistency.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_spacing_consistency(self):
        """Test spacing consistency check."""
        content = [
            "This is a paragraph with single spacing.",
            "This is a paragraph with  double  spacing.",
            "This is a paragraph with   triple   spacing.",
            "This is a paragraph with single spacing again.",
        ]
        # Use FormattingChecker for spacing checks
        result = self.formatting_checker.check_text("\n".join(content))
        assert not result.success
        # The checker uses 'Remove extra spaces' as the message
        assert any("Remove extra spaces" in issue["message"] for issue in result.issues)

    def test_margin_consistency(self):
        """Test margin consistency check."""
        content = [
            "This is a paragraph with normal margins.",
            "    This is a paragraph with indented margins.",
            "This is a paragraph with normal margins again.",
            "        This is another paragraph with indented margins.",
        ]
        doc_path = self.create_test_docx(content, "margin_consistency.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_list_formatting(self):
        """Test list formatting check."""
        content = [
            "The following items are required:",
            "1. First item",
            "2. Second item",
            "a. Sub-item",
            "b. Another sub-item",
            "3. Third item",
        ]
        doc_path = self.create_test_docx(content, "list_formatting.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_table_formatting(self):
        """Test table formatting check."""
        content = [
            "Table 1. Sample Table",
            "Column 1 | Column 2 | Column 3",
            "---------|----------|----------",
            "Data 1   | Data 2   | Data 3",
            "Data 4   | Data 5   | Data 6",
        ]
        doc_path = self.create_test_docx(content, "table_formatting.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_heading_formatting(self):
        """Test heading formatting check."""
        content = [
            "PURPOSE.",
            "This is the purpose section.",
            "BACKGROUND.",
            "This is the background section.",
            "DEFINITIONS.",
            "This is the definitions section.",
        ]
        doc_path = self.create_test_docx(content, "heading_formatting.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_reference_formatting(self):
        """Test reference formatting check."""
        content = [
            "See paragraph 5.2.3 for more information.",
            "Refer to section 4.1.2 for details.",
            "As discussed in paragraph 3.4.5",
        ]
        doc_path = self.create_test_docx(content, "reference_formatting.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_figure_formatting(self):
        """Test figure formatting check."""
        content = [
            "Figure 1. Sample Figure",
            "This is a figure caption.",
            "Figure 2. Another Sample Figure",
            "This is another figure caption.",
        ]
        doc_path = self.create_test_docx(content, "figure_formatting.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_phone_number_format_check_valid(self):
        """Test phone number format check with valid formats."""
        content = [
            "Contact us at (202) 267-1000.",
            "For emergencies, call (800) 555-1212.",
            "The office number is (703) 123-4567.",
        ]
        doc_path = self.create_test_docx(content, "valid_phone_numbers.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_phone_number_format_check_invalid(self):
        """Test phone number format check with invalid formats."""
        content = [
            "Contact us at 202-267-1000.",
            "For emergencies, call 800.555.1212.",
            "The office number is 7031234567.",
        ]
        doc_path = self.create_test_docx(content, "invalid_phone_numbers.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.severity == Severity.WARNING
        assert len(result.issues) > 0
        assert all(issue["severity"] == Severity.WARNING for issue in result.issues)

    def test_date_format_check_valid(self):
        """Test date format check with valid formats."""
        content = [
            "The document was issued on January 1, 2023.",
            "The meeting is scheduled for December 31, 2023.",
            "The deadline is March 15, 2023.",
        ]
        doc_path = self.create_test_docx(content, "valid_dates.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_date_format_check_invalid(self):
        """Test date format check with invalid formats."""
        content = [
            "The document was issued on 1/1/2023.",
            "The meeting is scheduled for 12-31-2023.",
            "The deadline is 03.15.2023.",
        ]
        doc_path = self.create_test_docx(content, "invalid_dates.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert not result.success
        assert result.severity == Severity.ERROR
        assert len(result.issues) > 0
        assert all(issue["severity"] == Severity.ERROR for issue in result.issues)

    def test_placeholder_check_valid(self):
        """Test placeholder check with valid text."""
        content = [
            "This is a document without placeholders.",
            "This is another sentence without placeholders.",
            "This is a third sentence without placeholders.",
        ]
        doc_path = self.create_test_docx(content, "valid_placeholders.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_placeholder_check_invalid(self):
        """Test placeholder check with invalid text."""
        content = [
            "This is a document with [PLACEHOLDER].",
            "This is another sentence with <PLACEHOLDER>.",
            "This is a third sentence with {PLACEHOLDER}.",
        ]
        doc_path = self.create_test_docx(content, "invalid_placeholders.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert not result.success
        assert result.severity == Severity.ERROR
        assert len(result.issues) > 0
        assert all(issue["severity"] == Severity.ERROR for issue in result.issues)

    def test_phone_number_format_usage(self):
        """Test phone number format usage check."""
        content = ["Call (123) 456-7890 for support", "Contact us at 123-456-7890"]
        doc_path = self.create_test_docx(content, "phone_number_usage.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.severity == Severity.WARNING
        assert len(result.issues) > 0
        assert all(issue["severity"] == Severity.WARNING for issue in result.issues)

    def test_phone_number_format_usage_inconsistent(self):
        """Test phone number format usage check with inconsistent formats."""
        content = ["Call (123) 456-7890 for support", "Contact us at 123.456.7890"]
        doc_path = self.create_test_docx(content, "phone_number_usage_inconsistent.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.severity == Severity.WARNING
        assert len(result.issues) > 0
        assert all(issue["severity"] == Severity.WARNING for issue in result.issues)

    def test_date_format_usage(self):
        """Test date format usage check."""
        content = ["The meeting is on January 1, 2023", "Deadline: 01/01/2023"]
        doc_path = self.create_test_docx(content, "date_format_usage.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert not result.success
        assert result.severity == Severity.ERROR
        assert len(result.issues) > 0
        assert all(issue["severity"] == Severity.ERROR for issue in result.issues)

    def test_date_format_usage_inconsistent(self):
        """Test date format usage check with inconsistent formats."""
        content = ["The meeting is on January 1, 2023", "Deadline: 1/1/23"]
        doc_path = self.create_test_docx(content, "date_format_usage_inconsistent.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert not result.success
        assert result.severity == Severity.ERROR
        assert len(result.issues) > 0
        assert all(issue["severity"] == Severity.ERROR for issue in result.issues)

    def test_placeholder_usage(self):
        """Test placeholder usage check."""
        content = ["This is a complete sentence.", "No placeholders here."]
        doc_path = self.create_test_docx(content, "placeholder_usage.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success

    def test_placeholder_usage_with_placeholders(self):
        """Test placeholder usage check with placeholders."""
        content = ["This is a complete sentence.", "TBD: Add more content here."]
        doc_path = self.create_test_docx(content, "placeholder_usage_with_placeholders.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert not result.success
        assert result.severity == Severity.ERROR
        assert len(result.issues) > 0
        assert all(issue["severity"] == Severity.ERROR for issue in result.issues)

    def test_caption_format_check_valid_ac(self):
        """Test caption format check with valid formats for AC/Order."""
        content = [
            "Table 5-1. Sample Table",
            "This is a table caption.",
            "Figure 3-2. Sample Figure",
            "This is a figure caption.",
        ]
        doc_path = self.create_test_docx(content, "valid_captions_ac.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "Advisory Circular", result)
        assert result.success
        assert len(result.issues) == 0

    # Helper for robust caption error checking
    def _caption_issue_matches(self, issue):
        msg = issue.get("message", "")
        if isinstance(msg, dict):
            # Check for incorrect_caption key or Table/Figure/Incorrect Caption in any value
            return "incorrect_caption" in msg or any(
                isinstance(v, str)
                and (
                    "table" in v.lower()
                    or "figure" in v.lower()
                    or "incorrect caption" in v.lower()
                )
                for v in msg.values()
            )
        elif isinstance(msg, str):
            return (
                "incorrect caption" in msg.lower()
                or "table" in msg.lower()
                or "figure" in msg.lower()
            )
        return False

    def test_caption_format_check_invalid_ac(self):
        """Test caption format check with invalid formats for AC/Order."""
        content = [
            "Table 5. Sample Table",  # Missing hyphen
            "This is a table caption.",
            "Figure 3. Sample Figure",  # Missing hyphen
            "This is a figure caption.",
        ]
        doc_path = self.create_test_docx(content, "invalid_captions_ac.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "Advisory Circular", result)
        assert not result.success
        assert len(result.issues) == 2
        assert all(issue["severity"] == Severity.ERROR for issue in result.issues)
        # Robustly check for caption error structure in the message
        assert all(self._caption_issue_matches(issue) for issue in result.issues)

    def test_caption_format_check_valid_other(self):
        """Test caption format check with valid formats for other document types."""
        content = [
            "Table 5. Sample Table",
            "This is a table caption.",
            "Figure 3. Sample Figure",
            "This is a figure caption.",
        ]
        doc_path = self.create_test_docx(content, "valid_captions_other.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "Other Document", result)
        assert result.success
        assert len(result.issues) == 0

    def test_caption_format_check_invalid_other(self):
        """Test caption format check with invalid formats for other document types."""
        content = [
            "Table 5-1. Sample Table",  # Has hyphen
            "This is a table caption.",
            "Figure 3-2. Sample Figure",  # Has hyphen
            "This is a figure caption.",
        ]
        doc_path = self.create_test_docx(content, "invalid_captions_other.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "Other Document", result)
        assert not result.success
        assert len(result.issues) == 2
        assert all(issue["severity"] == Severity.ERROR for issue in result.issues)
        # Robustly check for caption error structure in the message
        assert all(self._caption_issue_matches(issue) for issue in result.issues)

    def test_date_format_check_tso_reference(self):
        """Test that TSO references are not flagged as date format errors
        (regression test for FAA technical references).
        """
        content = [
            "The equipment meets TSO-C16b requirements.",
            "Refer to TSO-C123 for details.",
            "This is compliant with TSO-C2d.",
            "TSO-C151a is referenced in the document.",
        ]
        doc_path = self.create_test_docx(content, "tso_reference.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success
        assert len(result.issues) == 0

    def test_date_format_check_faa_technical_references(self):
        """Test that FAA technical references (AC, AD) are not flagged as date format errors."""
        content = [
            "See AC 25-1 for more information.",
            "Refer to AD 2020-12-05 for compliance.",
            "SFAR 88 is applicable.",
            "Order 8900.1 is referenced.",
            "Notice 8900.2 was issued.",
            "Policy 8900.3 applies.",
            "Memo 8900.4 is included.",
        ]
        doc_path = self.create_test_docx(content, "faa_technical_references.docx")
        result = DocumentCheckResult()
        self.format_checks.run_checks(Document(doc_path), "ORDER", result)
        assert result.success
        assert len(result.issues) == 0

    def test_date_format_check_tso_reference_string(self):
        """Test that TSO references are not flagged as date format errors in string-based checks."""
        content = "The equipment meets TSO-C16b requirements. Refer to TSO-C123 for details."
        result = self.formatting_checker.check_text(content)
        assert result.success
        assert len(result.issues) == 0

    def test_date_format_check_faa_technical_references_string(self):
        """
        Test that FAA references are not flagged as date format errors in string-based checks.
        """
        content = (
            "See AC 25-1 for more information. Refer to AD 2020-12-05 for compliance. "
            "SFAR 88 is applicable. Order 8900.1 is referenced. Notice 8900.2 was issued. "
            "Policy 8900.3 applies. Memo 8900.4 is included."
        )
        result = self.formatting_checker.check_text(content)
        assert result.success
        assert len(result.issues) == 0


# --- Shared test data for consolidation ---
#
# DATE_CASES: Only lines with MM/DD/YYYY format (not skipped by technical reference patterns)
# should be flagged.
# PHONE_CASES: If more than one phone number style is present, all lines with phone numbers
# are flagged.
# PLACEHOLDER_CASES: Each line with a placeholder is flagged.
DATE_CASES = [
    (["Date: 01/01/2023", "Date: May 11, 2025"], 1),  # Only the first line is flagged
    (["Date: 12/31/2023", "Date: December 31, 2023"], 1),  # Only the first line is flagged
    (["Date: May 11, 2025", "Date: December 31, 2023"], 0),  # No flagged lines
]

PHONE_CASES = [
    (["Phone: (123) 456-7890", "Phone: 123-456-7890"], 2),  # Two styles present, both lines flagged
    (["Phone: 123.456.7890", "Phone: 1234567890"], 2),  # Two styles present, both lines flagged
    (["Phone: 123-456-7890", "Phone: 123-456-7890"], 0),  # Only one style, no issues
]

PLACEHOLDER_CASES = [
    (["Lorem ipsum dolor sit amet", "TODO: Add content here"], 1),
    (["FIXME: Update this", "DRAFT: Review needed"], 2),
    (["Regular content", "More content"], 0),
]


def make_docx(lines, path):
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(path)
    return path


@pytest.mark.parametrize("lines,expect_errors", DATE_CASES)
def test_date_formats(managers, tmp_path, lines, expect_errors):
    """
    Test both text-level and docx-level date format checks.
    - Only lines with MM/DD/YYYY format (not skipped by technical reference patterns)
      should be flagged.
    - The expected error count should match the number of such lines.
    """
    fmt, txt = managers
    # text-level
    result = txt.check_text("\n".join(lines))
    assert len(result.issues) == expect_errors

    # docx-level
    docx_path = make_docx(lines, tmp_path / "tmp.docx")
    doc_result = DocumentCheckResult()
    fmt.run_checks(Document(docx_path), "ORDER", doc_result)
    assert len(doc_result.issues) == expect_errors


@pytest.mark.parametrize("lines,expect_errors", PHONE_CASES)
def test_phone_formats(managers, tmp_path, lines, expect_errors):
    """
    Test both text-level and docx-level phone number format checks.
    - If more than one phone number style is present, all lines with phone numbers are flagged.
    - If only one style is present, no issues are reported.
    """
    fmt, txt = managers
    # text-level
    result = txt.check_text("\n".join(lines))
    assert len(result.issues) == expect_errors

    # docx-level
    docx_path = make_docx(lines, tmp_path / "tmp.docx")
    doc_result = DocumentCheckResult()
    fmt.run_checks(Document(docx_path), "ORDER", doc_result)
    assert len(doc_result.issues) == expect_errors


@pytest.mark.parametrize("lines,expect_errors", PLACEHOLDER_CASES)
def test_placeholder_formats(managers, tmp_path, lines, expect_errors):
    """
    Test both text-level and docx-level placeholder checks.
    - Each line with a placeholder is flagged.
    """
    fmt, txt = managers
    # text-level
    result = txt.check_text("\n".join(lines))
    assert len(result.issues) == expect_errors

    # docx-level
    docx_path = make_docx(lines, tmp_path / "tmp.docx")
    doc_result = DocumentCheckResult()
    fmt.run_checks(Document(docx_path), "ORDER", doc_result)
    assert len(doc_result.issues) == expect_errors


if __name__ == "__main__":
    pytest.main()
