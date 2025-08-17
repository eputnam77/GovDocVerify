from docx import Document

from govdocverify.document_checker import FAADocumentChecker


def test_run_all_checks(tmp_path):
    doc = Document()
    doc.add_paragraph("1. INTRODUCTION", style="Heading 1")
    doc.add_paragraph(
        "This paragraph intentionally uses complex wording to lower readability scores."
    )
    path = tmp_path / "integration.docx"
    doc.save(path)

    checker = FAADocumentChecker()
    result = checker.run_all_document_checks(str(path), doc_type="internal_review")

    categories = [cat for _, cat in checker._get_check_modules()]
    assert set(categories).issubset(result.per_check_results.keys())
    assert result.issues and all("severity" in issue for issue in result.issues)
