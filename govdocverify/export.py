"""Helpers for exporting check results to various formats."""

from __future__ import annotations

import json
from typing import Any

from docx import Document

try:  # Optional dependency for PDF export
    from fpdf import FPDF
except ImportError:  # pragma: no cover - dependency might be missing
    FPDF = None


def save_results_as_docx(results: dict[str, Any], path: str) -> None:
    """Save results as a DOCX file."""
    doc = Document()
    doc.add_heading("Document Check Results", level=1)
    doc.add_paragraph(json.dumps(results, indent=2, ensure_ascii=False))
    doc.save(path)


def save_results_as_pdf(results: dict[str, Any], path: str) -> None:
    """Save results as a PDF file."""
    if FPDF is None:  # pragma: no cover - fallback when fpdf is unavailable
        with open(path, "wb") as fh:
            fh.write(b"%PDF-1.4\n% Unsupported minimal PDF\n")
        return

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    for line in json.dumps(results, indent=2, ensure_ascii=False).splitlines():
        pdf.cell(0, 10, txt=line, ln=True)
    pdf.output(path)
