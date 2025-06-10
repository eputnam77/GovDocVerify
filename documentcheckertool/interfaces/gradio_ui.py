import logging
import os
import tempfile
from importlib.metadata import PackageNotFoundError, version
from pprint import pformat

import gradio as gr

from documentcheckertool.document_checker import FAADocumentChecker
from documentcheckertool.models import (
    DocumentType,
    VisibilitySettings,
)
from documentcheckertool.utils.formatting import (
    FormatStyle,
    ResultFormatter,
)
from documentcheckertool.utils.security import validate_file

logger = logging.getLogger(__name__)
logger.debug("[UI DEBUG] gradio_ui.py module loaded")
try:
    GRADIO_VERSION = version("gradio")
except PackageNotFoundError:
    GRADIO_VERSION = "unknown"


def _get_custom_css():
    """Get custom CSS styles for the interface."""
    return """
    #document-checker-container {
        max-width: 1200px;
        margin: 0 auto;
        padding: 20px;
        font-family: system-ui, -apple-system, sans-serif;
    }

    /* Enhanced file upload styling */
    .upload-box {
        border: 2px dashed #e5e7eb !important;
        border-radius: 8px !important;
        background: #f8fafc !important;
        transition: all 0.3s ease !important;
        min-height: 150px !important;
        position: relative !important;
    }

    .upload-box:hover {
        border-color: #2563eb !important;
        background: #eff6ff !important;
    }

    /* Fix for file upload container */
    .file-upload {
        padding: 20px !important;
        border: none !important;
        background: transparent !important;
        height: 100% !important;
    }

    .file-upload > div {
        height: 100% !important;
        border: none !important;
        background: transparent !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }

    /* Additional upload styling */
    .file-upload label span {
        font-size: 1rem !important;
        color: #4b5563 !important;
    }

    .file-upload .upload-text {
        text-align: center !important;
        width: 100% !important;
    }

    .gr-form {
        background: white;
        border-radius: 8px;
        padding: 20px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }

    .gr-button-primary {
        background: #2563eb !important;
        border: none !important;
        color: white !important;
    }

    .gr-button-primary:hover {
        background: #1d4ed8 !important;
    }

    .gr-form > div {
        border: 1px solid #e5e7eb;
        border-radius: 8px;
        padding: 15px;
        margin-bottom: 15px;
    }

    .gr-box {
        border: 1px solid #e5e7eb;
        border-radius: 8px;
        background: white;
    }

    .gr-padded {
        padding: 15px;
    }

    .markdown-body {
        padding: 20px;
        background: white;
        border-radius: 8px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }

    /* Visibility controls styling */
    .visibility-controls {
        background: #f8f9fa;
        border-radius: 8px;
        padding: 15px;
        margin-bottom: 20px;
    }

    .visibility-controls h3 {
        color: #0056b3;
        margin-bottom: 15px;
    }

    .visibility-controls .checkbox-group {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 10px;
    }
    """


def get_readme_content():
    """Get README content for the About tab."""
    readme_path = "README.md"
    try:
        with open(readme_path, "r", encoding="utf-8") as file:
            readme_content = file.read()
        return readme_content
    except Exception as e:
        logging.error(f"Error reading README.md: {str(e)}")
        return "Error loading help content."


def create_interface():
    """Create and configure the Gradio interface."""
    # Initialize visibility settings
    VisibilitySettings()

    # Get CSS and template configuration
    custom_css = _get_custom_css()
    template_types = ["Short AC template AC", "Long AC template AC"]

    # Create the main interface
    demo = _create_gradio_blocks(custom_css)

    # Create UI components
    ui_components = _create_ui_components(template_types, demo)

    # Set up event handlers
    _setup_event_handlers(ui_components)

    return demo


def _create_gradio_blocks(custom_css):
    """Create the main Gradio Blocks interface."""
    return gr.Blocks(
        css=custom_css,
        theme=gr.themes.Soft(
            primary_hue="blue",
            secondary_hue="blue",
            neutral_hue="slate",
            spacing_size="sm",
            radius_size="lg",
            font=["system-ui", "sans-serif"],
        ),
    )


def _create_ui_components(template_types, demo):
    """Create all UI components and return them for event handling."""
    with demo:
        with gr.Row(elem_id="document-checker-container"):
            with gr.Column():
                with gr.Tabs():
                    with gr.Tab("Document Checker"):
                        _create_header()
                        components = _create_main_form(template_types)
                        _create_footer()
                    with gr.Tab("About"):
                        gr.Markdown(get_readme_content(), elem_classes="markdown-body")
    return components


def _create_header():
    """Create the header section with instructions."""
    gr.Markdown(
        """
        # 📑 FAA Document Checker Tool
        This tool performs multiple **validation checks** on Word documents to
        ensure compliance with U.S. federal documentation standards. See the
        tab for more information.
        ## How to Use
        1. Upload your Word document (`.docx` format).
        2. Select the document type.
        3. Click **Check Document**.
        > **Note:** Please ensure your document is clean (no track changes or
        comments). If your document contains track changes and comments, you
        get several false positives.
        """,
        elem_classes="markdown-body",
    )


def _create_main_form(template_types):
    """Create the main form with input controls and results area."""
    with gr.Group(elem_classes="gr-form"):
        with gr.Row():
            with gr.Column(scale=1):
                # Input controls
                file_input = gr.File(
                    label="📎 Upload Word Document (.docx)",
                    file_types=[".docx"],
                    type="binary",
                    elem_classes="file-upload upload-box",
                    interactive=True,
                )

                doc_type = gr.Dropdown(
                    label="Document Type",
                    choices=DocumentType.values(),
                    value=DocumentType.ADVISORY_CIRCULAR.value,
                    info="Select the type of document you're checking",
                    elem_classes="gr-box gr-padded",
                )

                template_type = gr.Radio(
                    choices=template_types,
                    label="📑 Template Type",
                    visible=False,
                    info="Only applicable for Advisory Circulars",
                    elem_classes="gr-box gr-padded",
                )

                group_by = gr.Radio(
                    choices=["category", "severity"],
                    value="category",
                    label="Group Results By",
                    info="Choose how to group the results: by functional "
                    "category or by severity.",
                )

                # Visibility controls
                visibility_components = _create_visibility_controls()

                submit_btn = gr.Button(
                    "🔍 Check Document",
                    variant="primary",
                    elem_classes="gr-button-primary",
                )

            with gr.Column(scale=2):
                # Results area
                results = gr.HTML(elem_classes="results-container")
                status_box = gr.Markdown("", elem_id="status-box")
                with gr.Row():
                    download_docx = gr.Button("📄 Download Report (DOCX)", visible=False)
                    download_pdf = gr.Button("📑 Download Report (PDF)", visible=False)
                report_file = gr.File(label="Download Report", visible=False)

    # Return all components needed for event handling
    return {
        "file_input": file_input,
        "doc_type": doc_type,
        "template_type": template_type,
        "group_by": group_by,
        "submit_btn": submit_btn,
        "results": results,
        "status_box": status_box,
        "download_docx": download_docx,
        "download_pdf": download_pdf,
        "report_file": report_file,
        **visibility_components,
    }


def _create_visibility_controls():
    """Create visibility control checkboxes."""
    with gr.Group(elem_classes="visibility-controls"):
        gr.Markdown("### 📊 Visibility Controls")
        with gr.Row():
            with gr.Column():
                show_readability = gr.Checkbox(
                    label="Readability Metrics",
                    value=True,
                    info="Show readability metrics and scores",
                )
                show_paragraph_length = gr.Checkbox(
                    label="Paragraph & Sentence Length",
                    value=True,
                    info="Show paragraph and sentence length checks",
                )
                show_terminology = gr.Checkbox(
                    label="Terminology Checks",
                    value=True,
                    info="Show terminology and style checks",
                )
                show_headings = gr.Checkbox(
                    label="Heading Checks",
                    value=True,
                    info="Show heading format and structure checks",
                )
            with gr.Column():
                show_structure = gr.Checkbox(
                    label="Structure Checks",
                    value=True,
                    info="Show document structure checks",
                )
                show_format = gr.Checkbox(
                    label="Format Checks",
                    value=True,
                    info="Show formatting and style checks",
                )
                show_accessibility = gr.Checkbox(
                    label="Accessibility Checks",
                    value=True,
                    info="Show accessibility compliance checks",
                )
                show_document_status = gr.Checkbox(
                    label="Document Status Checks",
                    value=True,
                    info="Show document status and watermark checks.",
                )

    return {
        "show_readability": show_readability,
        "show_paragraph_length": show_paragraph_length,
        "show_terminology": show_terminology,
        "show_headings": show_headings,
        "show_structure": show_structure,
        "show_format": show_format,
        "show_accessibility": show_accessibility,
        "show_document_status": show_document_status,
    }


def _create_footer():
    """Create the footer section with important notes."""
    gr.Markdown(
        """
        ### 📌 Important Notes
        - This tool helps you comply with federal documentation standards.
        - Results are based on current style guides and FAA requirements.
        - Final editorial decisions are the responsibility of the author.
        - For FAA documentation questions, contact your senior technical writer.
        - For tool questions or feedback, contact Eric Putnam.
        - Results are not stored or saved.
        """
    )


def _setup_event_handlers(components):
    """Set up all event handlers for the interface."""

    # Template visibility handler
    def update_template_visibility(doc_type_value):
        return gr.update(visible=doc_type_value == "Advisory Circular")

    components["doc_type"].change(
        fn=update_template_visibility,
        inputs=[components["doc_type"]],
        outputs=[components["template_type"]],
    )

    # Main processing handler
    components["submit_btn"].click(
        fn=_create_process_function(),
        inputs=[
            components["file_input"],
            components["doc_type"],
            components["template_type"],
            components["group_by"],
            components["show_readability"],
            components["show_paragraph_length"],
            components["show_terminology"],
            components["show_headings"],
            components["show_structure"],
            components["show_format"],
            components["show_accessibility"],
            components["show_document_status"],
        ],
        outputs=[
            components["results"],
            components["download_docx"],
            components["download_pdf"],
            components["report_file"],
            components["status_box"],
        ],
    )

    # Download handlers
    components["download_docx"].click(
        fn=lambda: generate_report_file(None, None, "docx"),
        inputs=[],
        outputs=[components["report_file"]],
    )

    components["download_pdf"].click(
        fn=lambda: generate_report_file(None, None, "pdf"),
        inputs=[],
        outputs=[components["report_file"]],
    )


def _create_process_function():
    """Create the main document processing function."""

    def process_and_format(
        file_obj,
        doc_type_value,
        template_type_value,
        group_by_value,
        show_readability_value,
        show_paragraph_length_value,
        show_terminology_value,
        show_headings_value,
        show_structure_value,
        show_format_value,
        show_accessibility_value,
        show_document_status_value,
    ):
        logger.debug("[UI DEBUG] process_and_format called")
        status = "Checking document..."
        try:
            if not file_obj:
                return (
                    "Please upload a document file.",
                    gr.update(visible=False),
                    gr.update(visible=False),
                    None,
                    "",
                )
            logger.info("Starting document processing...")

            # Create temporary file for validation
            temp_file_path = _create_temp_file(file_obj)

            try:
                validate_file(temp_file_path)
                checker = FAADocumentChecker()
                result_obj = checker.run_all_document_checks(
                    document_path=temp_file_path, doc_type=doc_type_value
                )
                results_dict = getattr(result_obj, "per_check_results", None)
                if not results_dict:
                    # fallback for legacy
                    results_dict = {
                        "all": {
                            "all": {
                                "success": result_obj.success,
                                "issues": result_obj.issues,
                                "details": getattr(result_obj, "details", {}),
                            }
                        }
                    }

                # Process and format results
                return _process_results(
                    results_dict,
                    result_obj,
                    doc_type_value,
                    group_by_value,
                    show_readability_value,
                    show_paragraph_length_value,
                    show_terminology_value,
                    show_headings_value,
                    show_structure_value,
                    show_format_value,
                    show_accessibility_value,
                    show_document_status_value,
                    status,
                )

            finally:
                try:
                    os.unlink(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file: {str(e)}")

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}", exc_info=True)
            status = f"Error: {str(e)}"
            return (
                format_error_message(str(e)),
                gr.update(visible=True),
                gr.update(visible=False),
                None,
                status,
            )

    return process_and_format


def _process_results(
    results_dict,
    result_obj,
    doc_type_value,
    group_by_value,
    show_readability_value,
    show_paragraph_length_value,
    show_terminology_value,
    show_headings_value,
    show_structure_value,
    show_format_value,
    show_accessibility_value,
    show_document_status_value,
    status,
):
    """Process and format the check results."""
    # Count issues in results_dict
    total_issues, issues_by_category = _count_issues_in_results(results_dict)

    # Build visibility settings
    visibility_settings = _build_visibility_settings(
        show_readability_value,
        show_paragraph_length_value,
        show_terminology_value,
        show_headings_value,
        show_structure_value,
        show_format_value,
        show_accessibility_value,
        show_document_status_value,
    )

    # Get selected categories and filter results
    selected_categories = _get_selected_categories(visibility_settings)
    filtered_results = _filter_results_by_visibility(results_dict, selected_categories)

    # Format results
    formatter = ResultFormatter(style=FormatStyle.HTML)
    formatted_results = formatter.format_results(
        filtered_results, doc_type_value, group_by=group_by_value
    )

    if formatted_results is None:
        logger.error(
            f"ResultFormatter.format_results returned None. "
            f"Inputs: results_dict={pformat(results_dict)}, "
            f"doc_type_value={doc_type_value}, "
            f"group_by_value={group_by_value}"
        )
        return (
            format_error_message("Internal error: Could not format results."),
            gr.update(visible=True),
            gr.update(visible=False),
            None,
            status,
        )

    # Log debug information
    _log_debug_info(
        total_issues,
        issues_by_category,
        visibility_settings,
        selected_categories,
        filtered_results,
        formatted_results,
    )

    logger.debug(f"Formatted HTML results: {formatted_results[:500]}")

    # Store for download handlers
    global _last_results
    _last_results = {
        "results": results_dict,
        "filtered_results": filtered_results,
        "visibility": visibility_settings.to_dict(),
        "summary": getattr(result_obj, "summary", {}),
        "formatted_results": formatted_results,
    }
    status = "Check complete."
    return (
        formatted_results,
        gr.update(visible=True),
        gr.update(visible=True),
        None,
        status,
    )


# Helper functions for the main processing
def format_error_message(error: str) -> str:
    """Format error messages for display in the UI."""
    return (
        f'<div style="color: #721c24; background-color: #f8d7da; '
        f"padding: 20px; "
        'border-radius: 8px; margin: 20px 0;">'
        f'<h3 style="color: #721c24; margin-top: 0;">'
        "Error Processing Document</h3>"
        f'<p style="margin-bottom: 0;">{error}</p>'
        "</div>"
    )


def _create_temp_file(file_obj):
    """Create temporary file from uploaded file object."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        if isinstance(file_obj, bytes):
            temp_file.write(file_obj)
        else:
            with open(file_obj.name, "rb") as f:
                temp_file.write(f.read())
        temp_file_path = temp_file.name
        logger.info(f"Created temporary file: {temp_file_path}")
        return temp_file_path


def _count_issues_in_results(results_dict):
    """Count total issues and issues by category."""
    total_issues = 0
    issues_by_category = {}
    for category, category_results in results_dict.items():
        cat_issues = 0
        if isinstance(category_results, dict):
            for result in category_results.values():
                if isinstance(result, dict) and "issues" in result:
                    cat_issues += len(result["issues"])
                elif hasattr(result, "issues"):
                    cat_issues += len(result.issues)
        issues_by_category[category] = cat_issues
        total_issues += cat_issues
    return total_issues, issues_by_category


def _build_visibility_settings(
    show_readability_value,
    show_paragraph_length_value,
    show_terminology_value,
    show_headings_value,
    show_structure_value,
    show_format_value,
    show_accessibility_value,
    show_document_status_value,
):
    """Build visibility settings from UI values."""
    return VisibilitySettings(
        show_readability=show_readability_value,
        show_paragraph_length=show_paragraph_length_value,
        show_terminology=show_terminology_value,
        show_headings=show_headings_value,
        show_structure=show_structure_value,
        show_format=show_format_value,
        show_accessibility=show_accessibility_value,
        show_document_status=show_document_status_value,
    )


def _get_selected_categories(visibility_settings):
    """Get selected categories based on visibility settings."""
    visibility_to_categories = {
        "show_readability": ["readability"],
        "show_paragraph_length": ["paragraph_length", "sentence_length"],
        "show_terminology": ["terminology"],
        "show_headings": ["heading"],
        "show_structure": ["structure"],
        "show_format": ["format"],
        "show_accessibility": ["accessibility"],
        "show_document_status": ["document_status"],
    }

    selected_categories = set()
    for setting_name, categories in visibility_to_categories.items():
        if getattr(visibility_settings, setting_name):
            selected_categories.update(categories)
    return selected_categories


def _filter_results_by_visibility(results_dict, selected_categories):
    """Filter results based on selected categories."""
    filtered_results = {}
    for category, category_results in results_dict.items():
        if category in selected_categories:
            filtered_results[category] = category_results
    return filtered_results


def _log_debug_info(
    total_issues,
    issues_by_category,
    visibility_settings,
    selected_categories,
    filtered_results,
    formatted_results,
):
    """Log debug information for UI display."""
    logger.debug(f"[UI DEBUG] Total issues in results_dict: {total_issues}")
    logger.debug(f"[UI DEBUG] Issues by category: {issues_by_category}")
    logger.debug(f"[UI DEBUG] Visibility settings: {visibility_settings}")
    logger.debug(f"[UI DEBUG] Selected categories for display: {selected_categories}")
    logger.debug(f"[UI DEBUG] Filtered categories: {list(filtered_results.keys())}")

    # Check if the UI will display 'All checks passed'
    if "All checks passed successfully" in formatted_results:
        if total_issues > 0:
            logger.warning(
                "[UI DEBUG] Issues found in results_dict, but UI will "
                "display 'All checks passed'. Possible bug in the "
                "formatter or results structure."
            )
        else:
            logger.debug("[UI DEBUG] No issues found; 'All checks passed'.")
    else:
        logger.debug("[UI DEBUG] UI will display issues.")


# Report generation functions
def _extract_report_data():
    """Extract data needed for report generation."""
    if not _last_results:
        logger.error("No results available for report generation")
        return None

    return {
        "results_dict": _last_results["results"],
        "visibility_settings": VisibilitySettings.from_dict(_last_results["visibility"]),
        "summary": _last_results["summary"],
        "formatted_results": _last_results.get("formatted_results"),
    }


def _filter_results_for_export(results_dict, visibility_settings):
    """Filter results based on visibility settings for export."""
    filtered_results = {}
    for category, category_results in results_dict.items():
        if getattr(visibility_settings, f"show_{category}", True):
            filtered_results[category] = category_results
    return filtered_results


def _count_export_issues(filtered_results):
    """Count total issues to be exported."""
    total_issues = sum(
        len(result["issues"])
        for cat in filtered_results.values()
        for result in cat.values()
        if "issues" in result and result["issues"]
    )
    logger.info(f"Exporting {total_issues} issues")
    if not filtered_results:
        logger.warning("No filtered results found for export.")
    return total_issues


# Report generation functions moved to module level
def _generate_docx_report(filepath, summary, filtered_results, total_issues):
    """Generate DOCX format report."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Add title
        title = doc.add_heading("Document Check Results", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add summary section
        doc.add_heading("Summary", level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f'Found {summary["total"]} issues that need attention:').bold = True

        # Add category summaries
        for category, count in summary["by_category"].items():
            if count > 0:
                doc.add_paragraph(
                    f'{category.replace("_", " ").title()}: {count} issues',
                    style="List Bullet",
                )

        # Add detailed results
        if total_issues == 0:
            doc.add_paragraph("No issues found in this document.", style="Normal")
        else:
            _add_docx_detailed_results(doc, filtered_results)

        doc.save(filepath)
        logger.info(f"DOCX report saved to: {filepath}")
        return filepath

    except ImportError:
        logger.error("python-docx not installed. Please install.")
        return None
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        return None


def _add_docx_detailed_results(doc, filtered_results):
    """Add detailed results to DOCX document."""
    for category, category_results in filtered_results.items():
        if category_results:
            doc.add_heading(category.replace("_", " ").title(), level=1)

            for check_name, result in category_results.items():
                if result.get("issues"):
                    doc.add_heading(check_name.replace("_", " ").title(), level=2)

                    for issue in result["issues"]:
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(issue["message"])
                        if issue.get("line_number"):
                            p.add_run(f' (Line {issue["line_number"]})').italic = True


def _generate_pdf_report(filepath, summary, filtered_results):
    """Generate PDF format report."""
    try:
        import pdfkit

        html_content = _build_pdf_html_content(summary, filtered_results)
        pdfkit.from_string(html_content, filepath)
        logger.info(f"PDF report saved to: {filepath}")
        return filepath

    except ImportError:
        logger.error("pdfkit not installed. Please install it.")
        return None
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        return None


def _build_pdf_html_content(summary, filtered_results):
    """Build HTML content for PDF generation."""
    html_content = f"""
    <html>
    <head>
        <style>
body {{ font-family: Arial, sans-serif; margin: 40px; }}
h1 {{ color: #0056b3; text-align: center; }}
h2 {{
    color: #0056b3;
    border-bottom: 2px solid #0056b3;
    padding-bottom: 10px;
}}
.summary {{
    background: #f8f9fa;
    padding: 20px;
    border-radius: 8px;
    margin-bottom: 30px;
}}
.category {{ margin-bottom: 30px; }}
.issue {{ margin: 10px 0; padding-left: 20px; }}
        </style>
    </head>
    <body>
        <h1>Document Check Results</h1>

        <div class="summary">
            <h2>Summary</h2>
            <p>Found {summary['total']} issues:</p>
            <ul>
    """

    for category, count in summary["by_category"].items():
        if count > 0:
            html_content += f"<li>{category.replace('_', ' ').title()}: {count} issues</li>"

    html_content += """
            </ul>
        </div>
    """

    for category, category_results in filtered_results.items():
        if category_results:
            html_content += f"""
            <div class="category">
                <h2>{category.replace('_', ' ').title()}</h2>
            """

            for check_name, result in category_results.items():
                if result.get("issues"):
                    display_name = check_name.replace("_", " ").title()
                    html_content += f"<h3>{display_name}</h3>\n<ul>\n"

                    for issue in result["issues"]:
                        line_info = (
                            f" (Line {issue['line_number']})" if issue.get("line_number") else ""
                        )
                        html_content += f'<li class="issue">{issue["message"]}{line_info}</li>\n'

                    html_content += "</ul>"

            html_content += "</div>"

    html_content += """
    </body>
    </html>
    """
    return html_content


def generate_report_file(results_data, doc_type_value, format="html"):
    """Generate downloadable report file."""
    try:
        # Extract data needed for report generation
        report_data = _extract_report_data()
        if not report_data:
            return None

        results_dict = report_data["results_dict"]
        visibility_settings = report_data["visibility_settings"]
        summary = report_data["summary"]
        formatted_results = report_data["formatted_results"]

        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{format}") as temp_file:
            filepath = temp_file.name

        # Filter results based on visibility settings
        filtered_results = _filter_results_for_export(results_dict, visibility_settings)

        # Count issues to be exported
        total_issues = _count_export_issues(filtered_results)

        # Generate report based on format
        if format == "docx":
            return _generate_docx_report(filepath, summary, filtered_results, total_issues)
        elif format == "pdf":
            return _generate_pdf_report(filepath, summary, filtered_results)
        else:
            # Default to HTML
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(formatted_results)
                logger.info(f"HTML report saved to: {filepath}")
                return filepath

    except Exception as e:
        logger.error(f"Error generating report: {str(e)}", exc_info=True)
        return None


# Global variable to store the last results
_last_results = None
